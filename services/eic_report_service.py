from __future__ import annotations

from datetime import datetime
from io import BytesIO
import hashlib
import json
from pathlib import Path
from typing import Any
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from core.i18n import t
from services.experiment_report_context_service import get_report_context, report_context_completion
from services import eic_service
from services.ai_usage_service import log_ai_usage_event, with_provider_chain
from services.authorization_service import CAPABILITY_VIEW_TECHNICAL_ARTIFACTS
from services.eic_display_service import (
    get_business_action_display,
    get_component_display_name,
    get_component_text_display,
    get_component_type_display,
    get_evidence_display,
    get_integrity_status_display,
    get_model_comparison_column_display,
    get_model_comparison_value_display,
    get_model_name_display,
    get_model_result_status_display,
    get_maturity_display,
    get_report_type_display,
    get_run_status_display,
)
from services.eic_report_charts import build_report_charts
from services.ml_experiment_service import (
    EXPERIMENT_PARADIGM_UNSUPERVISED,
    FINAL_VALIDATED_RUN_STATES,
    get_experiment_paradigm,
    list_run_artifacts,
)


REPORT_ROOT = Path("reports") / "ml_architecture" / "eic_reports"
REPORT_TYPES = ("experiment_docx", "executive_docx", "academic_docx", "technical_docx")
EXPERIMENT_REPORT_TEMPLATE_VERSION = 6
PAGE_WIDTH_PORTRAIT = 6.9
PAGE_WIDTH_LANDSCAPE = 9.4
TABLE_WIDTH_GUTTER = 0.18
CLASSIO_BLUE = RGBColor(31, 58, 95)
CLASSIO_TEAL = RGBColor(23, 128, 126)
CLASSIO_GOLD = RGBColor(183, 121, 31)
CLASSIO_RED = RGBColor(197, 48, 48)
MUTED = RGBColor(90, 98, 111)
TABLE_HEADER_FILL = "EAF0F7"
NOTE_FILL = "F6F8FB"


_COPY: dict[str, dict[str, str]] = {
    "en": {
        "validated": "Validated evidence",
        "confidential": "Validated internal report",
        "cover_kicker": "Classio Educational Intelligence Center",
        "cover_status": "Validation status: {value}",
        "cover_generated": "Generated on {value}",
        "cover_run": "Validated run: {value}",
        "cover_fingerprint": "Dataset fingerprint: {value}",
        "appendix": "Appendix",
        "source_note": "Source: validated stored run artifacts only.",
        "figure_note": "Interpretation uses stored metrics and predictions; no model retraining was performed.",
        "yes": "Yes",
        "no": "No",
        "not_available": "Not available",
        "executive_report_name": "Executive Business Report",
        "academic_report_name": "Findings Interpretation Report",
        "technical_report_name": "Technical Data Science Report",
        "meta_field": "Field",
        "meta_value": "Value",
        "summary_snapshot": "Validated evidence snapshot",
        "what_was_evaluated": "What was evaluated",
        "what_evidence_showed": "What the evidence showed",
        "replacement_question": "Can the heuristic system be replaced",
        "leadership_action": "What leadership should do now",
        "missing_evidence": "What evidence is still missing",
        "portfolio_summary": "Intelligence portfolio summary",
        "data_feedback_health": "Data and feedback health",
        "business_risks": "Business risks",
        "prioritized_actions": "Prioritized actions",
        "roadmap": "Roadmap and next review milestone",
        "run_metadata": "Concise run metadata",
        "exec_findings": "Leadership findings",
        "exec_intro": "This report translates the latest validated EIC evidence into business decisions for Classio leadership.",
        "exec_replace_hold": "The current heuristic assignment and recommendation logic should remain in place. The validated supervised experiment is informative, but it does not yet establish a production-grade replacement.",
        "exec_missing": "Evidence is still missing across teacher diversity, telemetry completeness, and repeated validated runs with broader operational coverage.",
        "academic_abstract": "Executive abstract",
        "company_context": "Company and EdTech context",
        "problem_statement": "Problem statement",
        "solution_statement": "Solution statement",
        "smart_objective": "SMART objectives",
        "supervised_formulation": "Supervised-learning formulation",
        "dataset_sources": "Dataset and data sources",
        "data_preparation": "Data preparation",
        "feature_selection": "Variable and feature selection",
        "target_construction": "Target construction",
        "methodology": "Experimental methodology",
        "models_evaluated": "Models evaluated",
        "evaluation_metrics": "Evaluation metrics",
        "results": "Results",
        "comparative_analysis": "Comparative analysis",
        "conclusions": "Conclusions",
        "implementation": "Business implementation implications",
        "limitations": "Limitations",
        "future_work": "Future work",
        "references": "References",
        "technical_metadata": "Document control and metadata",
        "experiment_definition": "Experiment definition",
        "source_logic": "Source tables and extraction logic",
        "dataset_accounting": "Dataset accounting",
        "label_reconciliation": "Label construction and reconciliation",
        "leakage_controls": "Leakage controls",
        "feature_health": "Feature schema and feature health",
        "preprocessing": "Preprocessing pipelines",
        "cv_section": "Chronological split and cross-validation",
        "model_configuration": "Model configurations",
        "baseline_results": "Baseline and comparison results",
        "threshold_analysis": "Threshold analysis",
        "roc_analysis": "ROC analysis",
        "pr_analysis": "Precision-recall analysis",
        "calibration_analysis": "Calibration analysis",
        "error_analysis": "Error analysis",
        "uncertainty": "Uncertainty and confidence intervals",
        "runtime": "Runtime and resource performance",
        "integrity": "Integrity-review results",
        "reproducibility": "Reproducibility",
        "deployment": "Deployment and shadow-testing recommendation",
        "artifact_manifest": "Artifact manifest",
        "section_note": "Interpretation",
        "academic_conclusion": "The comparison did not establish a statistically robust overall winner. Random Forest led several stored evaluation criteria, but the small holdout and single-teacher sample prevent a production-level conclusion.",
        "technical_conclusion": "The validated review concludes that this run is suitable as a technical evidence package, but not as a deployment decision.",
        "artifact_name": "Artifact",
        "artifact_purpose": "Purpose",
        "artifact_format": "Format",
        "artifact_checksum": "Checksum",
        "artifact_availability": "Availability",
        "artifact_filename": "Stored filename",
        "system": "System",
        "current_approach": "Current approach",
        "business_use": "Business use",
        "evidence": "Evidence",
        "recommended_action": "Recommended action",
        "priority": "Priority",
        "component": "Component",
        "impact": "Business impact",
        "owner": "Owner",
        "review_trigger": "Review trigger",
        "model": "Model",
        "result": "Result",
        "interpretation": "Interpretation",
        "caption_portfolio": "Portfolio table summarizing the systems currently discussed in the admin intelligence workspace.",
        "caption_actions": "Prioritized decisions table aligned to the admin leadership workflow.",
        "caption_models": "Readable comparison of stored supervised models using the validated run artifacts.",
        "caption_feature_health": "Feature-health table using the stored audit and review output.",
        "caption_manifest": "Artifact manifest summarizing the technical evidence package without exposing repository paths.",
        "visual_qa_passed": "Prepared for visual QA render",
    },
    "es": {
        "validated": "Evidencia validada",
        "confidential": "Informe interno validado",
        "cover_kicker": "Centro de Inteligencia Educativa de Classio",
        "cover_status": "Estado de validación: {value}",
        "cover_generated": "Generado el {value}",
        "cover_run": "Ejecución validada: {value}",
        "cover_fingerprint": "Huella del conjunto de datos: {value}",
        "appendix": "Apéndice",
        "source_note": "Fuente: solo artefactos validados y almacenados.",
        "figure_note": "La interpretación usa métricas y predicciones almacenadas; no se volvió a entrenar ningún modelo.",
        "yes": "Sí",
        "no": "No",
        "not_available": "No disponible",
        "executive_report_name": "Informe Ejecutivo de Negocio",
        "academic_report_name": "Informe de Interpretación de Hallazgos",
        "technical_report_name": "Informe Técnico de Ciencia de Datos",
        "meta_field": "Campo",
        "meta_value": "Valor",
        "summary_snapshot": "Resumen de evidencia validada",
        "what_was_evaluated": "Qué se evaluó",
        "what_evidence_showed": "Qué mostró la evidencia",
        "replacement_question": "Si el sistema heurístico puede sustituirse",
        "leadership_action": "Qué debe hacer ahora el liderazgo",
        "missing_evidence": "Qué evidencia sigue faltando",
        "portfolio_summary": "Resumen del portafolio de inteligencia",
        "data_feedback_health": "Salud de datos y retroalimentación",
        "business_risks": "Riesgos de negocio",
        "prioritized_actions": "Acciones priorizadas",
        "roadmap": "Hoja de ruta y próximo hito de revisión",
        "run_metadata": "Metadatos resumidos de la ejecución",
        "exec_findings": "Hallazgos para liderazgo",
        "exec_intro": "Este informe traduce la evidencia EIC validada más reciente en decisiones de negocio para el liderazgo de Classio.",
        "exec_replace_hold": "La lógica heurística actual para asignaciones y recomendaciones debe mantenerse. El experimento supervisado validado aporta señales útiles, pero no demuestra aún un reemplazo apto para producción.",
        "exec_missing": "Aún falta evidencia en diversidad docente, completitud de telemetría y repeticiones validadas con mayor cobertura operativa.",
        "academic_abstract": "Resumen ejecutivo",
        "company_context": "Contexto de la empresa y EdTech",
        "problem_statement": "Planteamiento del problema",
        "solution_statement": "Planteamiento de la solución",
        "smart_objective": "Objetivos SMART",
        "supervised_formulation": "Formulación de aprendizaje supervisado",
        "dataset_sources": "Conjunto de datos y fuentes",
        "data_preparation": "Preparación de datos",
        "feature_selection": "Selección de variables y características",
        "target_construction": "Construcción del objetivo",
        "methodology": "Metodología experimental",
        "models_evaluated": "Modelos evaluados",
        "evaluation_metrics": "Métricas de evaluación",
        "results": "Resultados",
        "comparative_analysis": "Análisis comparativo",
        "conclusions": "Conclusiones",
        "implementation": "Implicaciones de implementación",
        "limitations": "Limitaciones",
        "future_work": "Trabajo futuro",
        "references": "Referencias",
        "technical_metadata": "Control documental y metadatos",
        "experiment_definition": "Definición del experimento",
        "source_logic": "Tablas fuente y lógica de extracción",
        "dataset_accounting": "Contabilidad del conjunto de datos",
        "label_reconciliation": "Construcción y conciliación de etiquetas",
        "leakage_controls": "Controles contra fuga de información",
        "feature_health": "Esquema y salud de variables",
        "preprocessing": "Pipelines de preprocesamiento",
        "cv_section": "Corte cronológico y validación cruzada",
        "model_configuration": "Configuraciones de modelo",
        "baseline_results": "Resultados base y comparación",
        "threshold_analysis": "Análisis de umbrales",
        "roc_analysis": "Análisis ROC",
        "pr_analysis": "Análisis precision-recall",
        "calibration_analysis": "Análisis de calibración",
        "error_analysis": "Análisis de errores",
        "uncertainty": "Incertidumbre e intervalos de confianza",
        "runtime": "Rendimiento temporal y recursos",
        "integrity": "Resultados de la revisión de integridad",
        "reproducibility": "Reproducibilidad",
        "deployment": "Recomendación de despliegue y shadow testing",
        "artifact_manifest": "Manifiesto de artefactos",
        "section_note": "Interpretación",
        "academic_conclusion": "La comparación no estableció un ganador global estadísticamente robusto. Random Forest lideró varios criterios almacenados, pero el holdout pequeño y la muestra de un solo docente impiden una conclusión apta para producción.",
        "technical_conclusion": "La revisión validada concluye que esta ejecución sirve como paquete de evidencia técnica, pero no como decisión de despliegue.",
        "artifact_name": "Artefacto",
        "artifact_purpose": "Propósito",
        "artifact_format": "Formato",
        "artifact_checksum": "Checksum",
        "artifact_availability": "Disponibilidad",
        "artifact_filename": "Nombre almacenado",
        "system": "Sistema",
        "current_approach": "Enfoque actual",
        "business_use": "Uso de negocio",
        "evidence": "Evidencia",
        "recommended_action": "Acción recomendada",
        "priority": "Prioridad",
        "component": "Componente",
        "impact": "Impacto de negocio",
        "owner": "Responsable",
        "review_trigger": "Disparador de revisión",
        "model": "Modelo",
        "result": "Resultado",
        "interpretation": "Interpretación",
        "caption_portfolio": "Tabla del portafolio con los sistemas discutidos en el espacio de inteligencia del panel de administración.",
        "caption_actions": "Tabla de decisiones priorizadas alineada con el flujo ejecutivo del panel.",
        "caption_models": "Comparación legible de los modelos supervisados almacenados en la ejecución validada.",
        "caption_feature_health": "Tabla de salud de variables basada en la auditoría y la revisión almacenadas.",
        "caption_manifest": "Manifiesto de artefactos del paquete técnico sin exponer rutas del repositorio.",
        "visual_qa_passed": "Preparado para la revisión visual renderizada",
    },
    "tr": {
        "validated": "Doğrulanmış kanıt",
        "confidential": "Doğrulanmış dahili rapor",
        "cover_kicker": "Classio Eğitsel Zekâ Merkezi",
        "cover_status": "Doğrulama durumu: {value}",
        "cover_generated": "{value} tarihinde üretildi",
        "cover_run": "Doğrulanmış çalışma: {value}",
        "cover_fingerprint": "Veri kümesi parmak izi: {value}",
        "appendix": "Ek",
        "source_note": "Kaynak: yalnızca doğrulanmış ve saklanan artefaktlar.",
        "figure_note": "Yorum, saklanan metriklere ve tahminlere dayanır; model yeniden eğitilmedi.",
        "yes": "Evet",
        "no": "Hayır",
        "not_available": "Mevcut değil",
        "executive_report_name": "Yönetici İş Raporu",
        "academic_report_name": "Bulgular Yorumlama Raporu",
        "technical_report_name": "Teknik Veri Bilimi Raporu",
        "meta_field": "Alan",
        "meta_value": "Değer",
        "summary_snapshot": "Doğrulanmış kanıt özeti",
        "what_was_evaluated": "Ne değerlendirildi",
        "what_evidence_showed": "Kanıt ne gösterdi",
        "replacement_question": "Sezgisel sistemin yerini alıp alamayacağı",
        "leadership_action": "Liderliğin şimdi ne yapması gerektiği",
        "missing_evidence": "Hangi kanıtların hâlâ eksik olduğu",
        "portfolio_summary": "Zekâ portföyü özeti",
        "data_feedback_health": "Veri ve geri bildirim sağlığı",
        "business_risks": "İş riskleri",
        "prioritized_actions": "Öncelikli eylemler",
        "roadmap": "Yol haritası ve sonraki inceleme kilometre taşı",
        "run_metadata": "Kısa çalışma meta verisi",
        "exec_findings": "Liderlik bulguları",
        "exec_intro": "Bu rapor, en son doğrulanmış EIC kanıtını Classio liderliği için iş kararlarına dönüştürür.",
        "exec_replace_hold": "Mevcut sezgisel atama ve öneri mantığı korunmalıdır. Doğrulanmış denetimli deney yararlı bir sinyal sunuyor, ancak üretim düzeyinde bir ikameyi henüz kanıtlamıyor.",
        "exec_missing": "Öğretmen çeşitliliği, telemetri bütünlüğü ve daha geniş operasyonel kapsamda tekrarlanan doğrulanmış çalışmalar açısından kanıt hâlâ eksik.",
        "academic_abstract": "Yönetici özeti",
        "company_context": "Şirket ve EdTech bağlamı",
        "problem_statement": "Problemin tanımı",
        "solution_statement": "Çözüm yaklaşımı",
        "smart_objective": "SMART hedefleri",
        "supervised_formulation": "Denetimli öğrenme formülasyonu",
        "dataset_sources": "Veri kümesi ve kaynaklar",
        "data_preparation": "Veri hazırlığı",
        "feature_selection": "Değişken ve özellik seçimi",
        "target_construction": "Hedef oluşturma",
        "methodology": "Deneysel metodoloji",
        "models_evaluated": "Değerlendirilen modeller",
        "evaluation_metrics": "Değerlendirme metrikleri",
        "results": "Sonuçlar",
        "comparative_analysis": "Karşılaştırmalı analiz",
        "conclusions": "Sonuçlar",
        "implementation": "İş uygulaması etkileri",
        "limitations": "Sınırlamalar",
        "future_work": "Gelecek çalışmalar",
        "references": "Kaynaklar",
        "technical_metadata": "Belge kontrolü ve meta veriler",
        "experiment_definition": "Deney tanımı",
        "source_logic": "Kaynak tablolar ve çıkarım mantığı",
        "dataset_accounting": "Veri kümesi muhasebesi",
        "label_reconciliation": "Etiket oluşturma ve uzlaştırma",
        "leakage_controls": "Sızıntı kontrolleri",
        "feature_health": "Özellik şeması ve sağlık durumu",
        "preprocessing": "Ön işleme hatları",
        "cv_section": "Kronolojik ayrım ve çapraz doğrulama",
        "model_configuration": "Model yapılandırmaları",
        "baseline_results": "Temel ve karşılaştırma sonuçları",
        "threshold_analysis": "Eşik analizi",
        "roc_analysis": "ROC analizi",
        "pr_analysis": "Precision-recall analizi",
        "calibration_analysis": "Kalibrasyon analizi",
        "error_analysis": "Hata analizi",
        "uncertainty": "Belirsizlik ve güven aralıkları",
        "runtime": "Çalışma süresi ve kaynak performansı",
        "integrity": "Bütünlük inceleme sonuçları",
        "reproducibility": "Yeniden üretilebilirlik",
        "deployment": "Yayına alma ve shadow test önerisi",
        "artifact_manifest": "Artefakt manifestosu",
        "section_note": "Yorum",
        "academic_conclusion": "Karşılaştırma istatistiksel olarak güçlü bir genel kazanan ortaya koymadı. Random Forest birkaç kayıtlı ölçütte öne çıktı, ancak küçük holdout ve tek öğretmen örneklemi üretim düzeyinde sonuca izin vermiyor.",
        "technical_conclusion": "Doğrulanmış inceleme, bu çalışmanın teknik kanıt paketi olarak kullanılabileceğini ancak dağıtım kararı olarak kullanılamayacağını gösteriyor.",
        "artifact_name": "Artefakt",
        "artifact_purpose": "Amaç",
        "artifact_format": "Biçim",
        "artifact_checksum": "Sağlama",
        "artifact_availability": "Kullanılabilirlik",
        "artifact_filename": "Saklanan dosya adı",
        "system": "Sistem",
        "current_approach": "Mevcut yaklaşım",
        "business_use": "İş kullanımı",
        "evidence": "Kanıt",
        "recommended_action": "Önerilen eylem",
        "priority": "Öncelik",
        "component": "Bileşen",
        "impact": "İş etkisi",
        "owner": "Sorumlu",
        "review_trigger": "İnceleme tetikleyicisi",
        "model": "Model",
        "result": "Sonuç",
        "interpretation": "Yorum",
        "caption_portfolio": "Yönetici panelindeki zekâ çalışma alanında tartışılan sistemleri özetleyen portföy tablosu.",
        "caption_actions": "Yönetici iş akışıyla hizalanmış öncelikli kararlar tablosu.",
        "caption_models": "Doğrulanmış çalışma artefaktlarındaki denetimli modellerin okunabilir karşılaştırması.",
        "caption_feature_health": "Saklanan denetim ve gözden geçirme çıktısından üretilen özellik sağlığı tablosu.",
        "caption_manifest": "Depo yollarını göstermeden teknik kanıt paketini özetleyen artefakt manifestosu.",
        "visual_qa_passed": "Görsel render incelemesi için hazırlandı",
    },
}


def _clean_text(value: Any) -> str:
    return " ".join(str(value or "").split()).strip()


def _translated_business_question(detail: dict[str, Any], lang: str) -> str:
    experiment_id = _clean_text(detail.get("experiment_id") or "")
    if experiment_id in {
        "assigned_resource_open_within_7d",
        "student_recommendation_open_within_7d",
        "resource_affinity_unsupervised_discovery",
    }:
        return get_component_text_display(
            experiment_id,
            "business_question",
            detail.get("business_question"),
            lang=lang,
        )
    return _clean_text(detail.get("business_question"))


def _translate_known_evidence_text(value: Any, lang: str) -> str:
    text = _clean_text(value)
    if not text or lang == "en":
        return text
    replacements = {
        "es": {
            "This run was generated by the current Phase 3.6 pipeline, so historical Phase 3.5 audit-count reconciliation is not required for validation.": "Esta ejecución fue generada por la canalización actual de la Fase 3.6, por lo que la conciliación histórica de conteos de auditoría de la Fase 3.5 no es necesaria para la validación.",
            "No identifier mismatch was found.": "No se detectó ningún desajuste de identificadores.",
            "Missing values are produced by the feature-construction rule that only uses strictly earlier mature history.": "Los valores faltantes se producen por la regla de construcción de variables que solo utiliza historial maduro estrictamente anterior.",
            "For early assignments there is no mature prior history yet; resource-level sparsity is especially severe for prior_resource_open_rate.": "En las asignaciones tempranas todavía no existe historial maduro previo; la dispersión a nivel de recurso es especialmente alta para prior_resource_open_rate.",
            "Learning-program topics are curricular anchors. Worksheets, exams, videos, lesson plans, and learning programs are candidate resources tested for semantic alignment with those anchors.": "Los temas del learning program son anclas curriculares. Worksheets, exámenes, videos, lesson plans y learning programs se tratan como recursos candidatos cuya alineación semántica se evalúa contra esas anclas.",
            "No human-labeled semantic relevance labels are used in this unsupervised phase.": "En esta fase no supervisada no se usan etiquetas humanas de relevancia semántica.",
            "Small or imbalanced resource catalogs can make clustering metrics unstable.": "Los catálogos de recursos pequeños o desbalanceados pueden volver inestables las métricas de agrupamiento.",
            "Program topics are curricular anchors, not standalone resources; candidate-resource alignment must be reviewed separately from topic-to-topic similarity.": "Los program topics son anclas o subunidades curriculares, no recursos completos; la alineación de recursos candidatos debe revisarse por separado de la similitud entre temas.",
            "Business rules must still filter by subject, language, teacher, student, level, and archive state.": "Las reglas de negocio todavía deben filtrar por materia, idioma, profesor, estudiante, nivel y estado de archivo.",
            "Cross-language contamination is calculated only from non-empty known language values. Missing language metadata limits interpretation.": "La contaminación entre idiomas se calcula solo con valores de idioma conocidos y no vacíos. La metadata de idioma faltante limita la interpretación.",
            "Archived rows are excluded before model development.": "Las filas archivadas se excluyen antes del desarrollo del modelo.",
            "Rows without a stable resource id are excluded.": "Las filas sin un identificador estable de recurso se excluyen.",
            "Rows with fewer than three profile tokens are excluded.": "Las filas con menos de tres tokens de perfil se excluyen.",
            "Duplicate rows with the same resource key and profile hash are excluded after the first occurrence.": "Las filas duplicadas con la misma clave de recurso y el mismo hash de perfil se excluyen después de la primera aparición.",
            "The primary extractor does not select worksheet_json or exam_data. Content is included only through bounded sanitized excerpts when the lightweight excerpt view exists.": "El extractor principal no selecciona worksheet_json ni exam_data. El contenido solo se incluye mediante extractos sanitizados y acotados cuando existe la vista liviana de extractos.",
            "If the excerpt view is unavailable, worksheets and exams remain included through metadata-only profiles and the run records a warning.": "Si la vista de extractos no está disponible, worksheets y exámenes permanecen incluidos mediante perfiles basados solo en metadata, y la ejecución registra una advertencia.",
            "Each included row becomes one canonical text profile.": "Cada fila incluida se convierte en un perfil textual canónico.",
            "content_excerpt is included only when a bounded sanitized excerpt is available; images, media, URLs, answer keys, correct answers, and solutions are excluded.": "content_excerpt solo se incluye cuando existe un extracto sanitizado y acotado; se excluyen imágenes, medios, URLs, answer keys, respuestas correctas y soluciones.",
            "program_topic rows are assigned resource_role=curricular_anchor.": "Las filas program_topic reciben resource_role=curricular_anchor.",
            "worksheet, exam, video, lesson_plan, and program rows are assigned resource_role=candidate_resource.": "Las filas worksheet, exam, video, lesson_plan y program reciben resource_role=candidate_resource.",
            "Lesson-plan topics are extracted from plan_json keys related to topic, title, objective, focus, vocabulary, success, and assessment, plus the fallback topic field.": "Los temas de lesson plans se extraen desde claves de plan_json relacionadas con tema, título, objetivo, foco, vocabulario, éxito y evaluación, además del campo topic como respaldo.",
            "Program topics inherit parent learning-program metadata when the parent program is available, so topic anchors carry subject, level, and stage context.": "Los program topics heredan metadata del learning program padre cuando está disponible, por lo que las anclas de tema conservan contexto de materia, nivel y etapa.",
            "Build a pandas DataFrame of frozen resource profiles.": "Construir un DataFrame de pandas con los perfiles de recursos congelados.",
            "Vectorize profile_text with scikit-learn TfidfVectorizer using unigrams and bigrams.": "Vectorizar profile_text con TfidfVectorizer de scikit-learn usando unigramas y bigramas.",
            "When matrix shape permits, reduce TF-IDF features with TruncatedSVD.": "Cuando la forma de la matriz lo permite, reducir las variables TF-IDF con TruncatedSVD.",
            "Normalize vectors with scikit-learn Normalizer using L2 normalization.": "Normalizar vectores con Normalizer de scikit-learn usando normalización L2.",
            "Train candidate KMeans, AgglomerativeClustering, and DBSCAN configurations.": "Entrenar configuraciones candidatas de KMeans, AgglomerativeClustering y DBSCAN.",
            "Evaluate each successful configuration with Silhouette, Calinski-Harabasz, Davies-Bouldin, noise ratio, cluster-size diagnostics, and cross-subject/cross-language contamination.": "Evaluar cada configuración exitosa con Silhouette, Calinski-Harabasz, Davies-Bouldin, proporción de ruido, diagnósticos de tamaño de clúster y contaminación entre materias/idiomas.",
            "Select the winner by the transparent balanced selection_score formula.": "Seleccionar el ganador con la fórmula transparente y balanceada de selection_score.",
            "Generate pairwise semantic neighbors with cosine similarity.": "Generar vecinos semánticos pareados mediante similitud coseno.",
            "Generate the human-review sample from program_topic anchors to candidate resources.": "Generar la muestra de revisión humana desde anclas program_topic hacia recursos candidatos.",
            "Persist the fitted vectorizer, SVD, normalizer, vector matrix, ordered resource keys, frozen dataset, model comparison, cluster assignments, and audit files.": "Persistir el vectorizador ajustado, SVD, normalizador, matriz de vectores, claves de recursos ordenadas, dataset congelado, comparación de modelos, asignaciones de clúster y archivos de auditoría.",
            "Random seed is fixed at 20260726.": "La semilla aleatoria se fija en 20260726.",
            "The run stores a dataset fingerprint built from profile hashes.": "La ejecución guarda una huella del dataset construida a partir de hashes de perfil.",
            "The run stores a configuration hash for the model grid and preprocessing settings.": "La ejecución guarda un hash de configuración para la grilla de modelos y los ajustes de preprocesamiento.",
            "The fitted representation is serialized so runtime scoring does not refit TF-IDF from live data.": "La representación ajustada se serializa para que el scoring en runtime no reajuste TF-IDF con datos vivos.",
            "highest selection_score": "mayor selection_score",
            "highest silhouette_score as tie-breaker": "mayor silhouette_score como desempate",
            "lowest cross_subject_contamination_rate as tie-breaker": "menor cross_subject_contamination_rate como desempate",
            "Share of clusters with exactly two rows, used as a fragmentation warning.": "Proporción de clústeres con exactamente dos filas, usada como advertencia de fragmentación.",
            "Share of clusters with three or fewer rows, used as a broader fragmentation warning.": "Proporción de clústeres con tres filas o menos, usada como advertencia más amplia de fragmentación.",
            "Share of evaluated clusters containing more than one normalized subject among rows with known subject metadata.": "Proporción de clústeres evaluados que contienen más de una materia normalizada entre filas con metadata de materia conocida.",
            "Share of rows assigned to DBSCAN noise label -1.": "Proporción de filas asignadas a la etiqueta de ruido -1 de DBSCAN.",
            "Primary unsupervised cohesion/separation metric calculated with cosine distance on non-noise observations.": "Métrica no supervisada principal de cohesión/separación, calculada con distancia coseno sobre observaciones que no son ruido.",
            "Silhouette remains the positive base because the experiment is unsupervised and has no human relevance labels yet.": "Silhouette se mantiene como base positiva porque el experimento es no supervisado y aún no tiene etiquetas humanas de relevancia.",
            "Cross-subject contamination receives the largest penalty because Classio's recommendation rules treat subject boundaries as business-critical.": "La contaminación entre materias recibe la penalización más alta porque las reglas de recomendación de Classio tratan los límites de materia como críticos para el negocio.",
            "Noise receives a moderate penalty because a model that leaves many rows unassigned is less useful for candidate generation.": "El ruido recibe una penalización moderada porque un modelo que deja muchas filas sin asignar es menos útil para generar candidatos.",
            "Tiny-cluster penalties discourage models that look clean only because they fragment the catalog into very small groups.": "Las penalizaciones por clústeres diminutos desincentivan modelos que parecen limpios solo porque fragmentan el catálogo en grupos muy pequeños.",
            "The weights are heuristic and transparent; they are intended for exploratory model triage, not as proof of recommendation impact.": "Los pesos son heurísticos y transparentes; sirven para triage exploratorio de modelos, no como prueba de impacto en recomendaciones.",
            "The highest Silhouette model can over-reward small, isolated, or noisy clusters. The balanced score prefers a model that is coherent enough while still preserving coverage and useful cluster structure for downstream heuristics.": "El modelo con mayor Silhouette puede sobrepremiar clústeres pequeños, aislados o ruidosos. El score balanceado prefiere un modelo suficientemente coherente que preserve cobertura y una estructura de clúster útil para las heurísticas posteriores.",
            "The winner is the top balanced selection-score candidate. The metric leader is reported separately because a higher Silhouette score alone can over-reward fragmented, noisy, or business-contaminated cluster structures.": "El ganador es el candidato con mayor score balanceado de selección. El líder métrico se informa por separado porque un Silhouette más alto por sí solo puede sobrepremiar estructuras fragmentadas, ruidosas o contaminadas desde el punto de vista de negocio.",
            "subject, language, level, resource_type, and resource_role are audited.": "Se auditan materia, idioma, nivel, resource_type y resource_role.",
            "Category values are case-folded.": "Los valores categóricos se normalizan con casefold.",
            "Slash and pipe separators are converted to spaces.": "Los separadores slash y pipe se convierten en espacios.",
            "Known aliases are mapped before metrics are calculated.": "Los alias conocidos se mapean antes de calcular las métricas.",
            "Accent-insensitive fallback keys are used when an alias exists without diacritics.": "Se usan claves de respaldo insensibles a acentos cuando existe un alias sin diacríticos.",
            "Empty metadata values are excluded from contamination denominators.": "Los valores de metadata vacíos se excluyen de los denominadores de contaminación.",
            "All text fields are converted to strings, collapsed to single spaces, and stripped of leading/trailing whitespace.": "Todos los campos de texto se convierten a strings, los espacios múltiples se compactan en uno y se eliminan espacios iniciales/finales.",
            "Unicode text is normalized with NFKC for category comparison.": "El texto Unicode se normaliza con NFKC para comparar categorías.",
            "Dash variants are normalized before category comparison.": "Las variantes de guion se normalizan antes de comparar categorías.",
            "Canonical profile text is vectorized with TF-IDF using unigrams and bigrams.": "El texto canónico del perfil se vectoriza con TF-IDF usando unigramas y bigramas.",
            "When TF-IDF shape allows it, Truncated SVD reduces the sparse matrix to dense latent dimensions.": "Cuando la forma TF-IDF lo permite, Truncated SVD reduce la matriz dispersa a dimensiones latentes densas.",
            "All vectors are L2-normalized before cosine similarity, clustering quality calculations, and runtime scoring.": "Todos los vectores se normalizan con L2 antes de calcular similitud coseno, calidad de clústeres y scoring en runtime.",
            "Optional resource_affinity_content_excerpts view: resource_type, resource_id, content_excerpt, source, and character count. This view must return sanitized bounded text only; the experiment does not fetch full worksheet_json or exam_data as its primary path.": "Vista opcional resource_affinity_content_excerpts: resource_type, resource_id, content_excerpt, source y conteo de caracteres. Esta vista debe devolver solo texto sanitizado y acotado; el experimento no obtiene worksheet_json ni exam_data completos como ruta principal.",
            "quick_exams table: title, subject, topic, learner stage, level, exam length, exercise types, status, visibility, and creation timestamp. Full exam_data is not fetched by the experiment.": "Tabla quick_exams: título, materia, tema, etapa del estudiante, nivel, duración del examen, tipos de ejercicios, estado, visibilidad y fecha de creación. El experimento no obtiene exam_data completo.",
            "learning_program_topics table: program_id, unit/topic order, title, subtopic, lesson focus, lesson purpose, objectives, success criteria, can-do statements, suggested worksheet/exam types, homework idea, teacher notes, student summary, estimated lessons, and timestamps.": "Tabla learning_program_topics: program_id, orden de unidad/tema, título, subtema, foco de lección, propósito, objetivos, criterios de éxito, can-do statements, tipos sugeridos de worksheet/examen, idea de tarea, notas del profesor, resumen para estudiante, lecciones estimadas y timestamps.",
            "learning_programs table: title, subject/custom subject, learner stage, level/band, overview, status, visibility, unit/topic counts, sequence order, timestamps, and program_data content.": "Tabla learning_programs: título, materia/materia personalizada, etapa del estudiante, nivel/banda, resumen, estado, visibilidad, conteos de unidades/temas, orden de secuencia, timestamps y contenido de program_data.",
            "lesson_plans table: title, subject, topic, learner stage, level/band, lesson purpose, source/planner metadata, language fields, status, visibility, creation/update timestamps, and plan_json content.": "Tabla lesson_plans: título, materia, tema, etapa del estudiante, nivel/banda, propósito de la lección, metadata de fuente/planificador, campos de idioma, estado, visibilidad, timestamps de creación/actualización y contenido de plan_json.",
            "videos table: title, subject/custom subject, topic, description, learner stage, level/band, status, visibility, creation/update timestamps.": "Tabla videos: título, materia/materia personalizada, tema, descripción, etapa del estudiante, nivel/banda, estado, visibilidad y timestamps de creación/actualización.",
            "worksheets table: title, subject, topic, learner stage, level/band, worksheet type, language fields, status, visibility, and creation timestamp. Full worksheet_json is not fetched by the experiment.": "Tabla worksheets: título, materia, tema, etapa del estudiante, nivel/banda, tipo de worksheet, campos de idioma, estado, visibilidad y fecha de creación. El experimento no obtiene worksheet_json completo.",
        },
        "tr": {
            "This run was generated by the current Phase 3.6 pipeline, so historical Phase 3.5 audit-count reconciliation is not required for validation.": "Bu çalışma mevcut Phase 3.6 hattı tarafından üretildiği için, doğrulama açısından geçmiş Phase 3.5 denetim sayımı uzlaştırması gerekli değildir.",
            "No identifier mismatch was found.": "Kimlik eşleşmesinde bir tutarsızlık bulunmadı.",
            "Missing values are produced by the feature-construction rule that only uses strictly earlier mature history.": "Eksik değerler, yalnızca daha önceki olgun geçmişi kullanan özellik oluşturma kuralından kaynaklanır.",
            "For early assignments there is no mature prior history yet; resource-level sparsity is especially severe for prior_resource_open_rate.": "Erken atamalarda henüz olgun bir önceki geçmiş yoktur; kaynak düzeyindeki seyreklik özellikle prior_resource_open_rate için yüksektir.",
        },
    }
    translated = text
    for source, target in replacements.get(lang, {}).items():
        translated = translated.replace(source, target)
    return translated


_AFFINITY_DATA_SOURCE_LABELS = {
    "es": {
        "bounded_content_excerpts": "Extractos acotados de contenido",
        "exams": "Exámenes",
        "learning_program_topics": "Temas del learning program",
        "learning_programs": "Learning programs",
        "lesson_plans": "Lesson plans",
        "videos": "Videos",
        "worksheets": "Worksheets",
    },
    "tr": {},
}

_AFFINITY_COMPONENT_LABELS = {
    "es": {
        "cluster_size_2_rate": "Tasa de clústeres de tamaño 2",
        "cluster_size_le_3_rate": "Tasa de clústeres de tamaño 3 o menos",
        "cross_subject_contamination_rate": "Contaminación entre materias",
        "noise_ratio": "Proporción de ruido",
        "silhouette_score": "Silhouette score",
    },
    "tr": {},
}

_AFFINITY_METRIC_LABELS = {
    "es": {
        "balanced_selection_score": "score balanceado de selección",
        "selection_score": "selection_score",
    },
    "tr": {},
}


def _localize_known_text(value: Any, lang: str) -> str:
    return _translate_known_evidence_text(value, lang)


def _localized_join(values: Any, lang: str, *, separator: str = " ") -> str:
    return separator.join(_localize_known_text(item, lang) for item in list(values or []) if _clean_text(item))


def _localized_bullets(values: Any, lang: str) -> list[str]:
    return _unique_nonempty([_localize_known_text(item, lang) for item in list(values or [])])


def _localized_json_map(values: dict[str, Any], lang: str, labels: dict[str, dict[str, str]] | None = None) -> str:
    localized: dict[str, Any] = {}
    lang_labels = (labels or {}).get(lang, {})
    for key, value in (values or {}).items():
        display_key = lang_labels.get(str(key), _humanize_identifier(key))
        localized[display_key] = _localize_known_text(value, lang)
    return json.dumps(localized, ensure_ascii=False, sort_keys=True)


def _date_connector(lang: str) -> str:
    if lang == "es":
        return "a"
    if lang == "tr":
        return "-"
    return "to"


def _copy(lang: str, key: str, **kwargs) -> str:
    template = _COPY[_lang(lang)].get(key, _COPY["en"].get(key, key))
    return template.format(**kwargs)


_REPORT_PHRASES: dict[str, dict[str, str]] = {
    "en": {
        "chronological_cutoff": "Chronological cutoff",
        "source_rows": "Source rows",
        "included_rows": "Included rows",
        "positive_labels": "Positive labels",
        "negative_labels": "Negative labels",
        "right_censored": "Right-censored",
        "teachers_represented": "Teachers represented",
        "students_represented": "Students represented",
        "resources_represented": "Resources represented",
        "date_range": "Date range",
        "development_rows": "Development rows",
        "holdout_rows": "Holdout rows",
        "stored_review": "Stored review",
        "past_only": "Past-only",
        "full_model_comparison": "Full model comparison",
        "discrimination_ranking": "Discrimination and ranking",
        "calibration_thresholding": "Calibration and thresholding",
        "runtime_status": "Runtime and status",
        "feature_col": "Feature",
        "source_col": "Source",
        "availability_col": "Availability",
        "prediction_time_availability_col": "Prediction-time availability",
        "overall_missing_col": "Overall missing",
        "dev_missing_col": "Dev missing",
        "holdout_missing_col": "Holdout missing",
        "included_col": "Included",
        "explanation_col": "Explanation",
        "kind_col": "Kind",
        "status_col": "Status",
        "selected_hyperparameters": "Selected hyperparameters",
        "stored_supporting_artifact": "Stored supporting artifact",
        "validated_run_summary": "Validated run summary",
        "dataset_accounting_purpose": "Dataset accounting and counts",
        "stored_model_metrics": "Stored model metrics",
        "stored_holdout_probabilities": "Stored holdout probabilities",
        "feature_availability_audit": "Feature availability audit",
        "label_reconciliation_review": "Label reconciliation review",
        "integrity_narrative": "Integrity narrative",
        "findings_markdown_baseline": "Findings interpretation markdown baseline",
        "technical_markdown_baseline": "Technical markdown baseline",
    },
    "es": {
        "chronological_cutoff": "Corte cronológico",
        "source_rows": "Filas de origen",
        "included_rows": "Filas incluidas",
        "positive_labels": "Etiquetas positivas",
        "negative_labels": "Etiquetas negativas",
        "right_censored": "Censuradas por ventana",
        "teachers_represented": "Docentes representados",
        "students_represented": "Estudiantes representados",
        "resources_represented": "Recursos representados",
        "date_range": "Rango de fechas",
        "development_rows": "Filas de desarrollo",
        "holdout_rows": "Filas de holdout",
        "stored_review": "Revisión almacenada",
        "past_only": "Solo pasado",
        "full_model_comparison": "Comparación completa de modelos",
        "discrimination_ranking": "Discriminación y ranking",
        "calibration_thresholding": "Calibración y umbrales",
        "runtime_status": "Rendimiento y estado",
        "feature_col": "Variable",
        "source_col": "Fuente",
        "availability_col": "Disponibilidad",
        "prediction_time_availability_col": "Disponibilidad al momento de predicción",
        "overall_missing_col": "Falta total",
        "dev_missing_col": "Falta en desarrollo",
        "holdout_missing_col": "Falta en holdout",
        "included_col": "Incluida",
        "explanation_col": "Explicación",
        "kind_col": "Tipo",
        "status_col": "Estado",
        "selected_hyperparameters": "Hiperparámetros seleccionados",
        "stored_supporting_artifact": "Artefacto de soporte almacenado",
        "validated_run_summary": "Resumen de ejecución validada",
        "dataset_accounting_purpose": "Contabilidad y conteos del conjunto de datos",
        "stored_model_metrics": "Métricas de modelo almacenadas",
        "stored_holdout_probabilities": "Probabilidades holdout almacenadas",
        "feature_availability_audit": "Auditoría de disponibilidad de variables",
        "label_reconciliation_review": "Revisión de conciliación de etiquetas",
        "integrity_narrative": "Narrativa de integridad",
        "findings_markdown_baseline": "Base markdown de interpretación de hallazgos",
        "technical_markdown_baseline": "Base markdown técnica",
    },
    "tr": {
        "chronological_cutoff": "Kronolojik kesim",
        "source_rows": "Kaynak satırlar",
        "included_rows": "Dahil edilen satırlar",
        "positive_labels": "Pozitif etiketler",
        "negative_labels": "Negatif etiketler",
        "right_censored": "Sağ sansürlü",
        "teachers_represented": "Temsil edilen öğretmenler",
        "students_represented": "Temsil edilen öğrenciler",
        "resources_represented": "Temsil edilen kaynaklar",
        "date_range": "Tarih aralığı",
        "development_rows": "Geliştirme satırları",
        "holdout_rows": "Holdout satırları",
        "stored_review": "Saklanan inceleme",
        "past_only": "Yalnızca geçmiş",
        "full_model_comparison": "Tam model karşılaştırması",
        "discrimination_ranking": "Ayrıştırma ve sıralama",
        "calibration_thresholding": "Kalibrasyon ve eşikleme",
        "runtime_status": "Çalışma süresi ve durum",
        "feature_col": "Özellik",
        "source_col": "Kaynak",
        "availability_col": "Kullanılabilirlik",
        "prediction_time_availability_col": "Tahmin anı kullanılabilirliği",
        "overall_missing_col": "Genel eksik",
        "dev_missing_col": "Geliştirme eksik",
        "holdout_missing_col": "Holdout eksik",
        "included_col": "Dahil",
        "explanation_col": "Açıklama",
        "kind_col": "Tür",
        "status_col": "Durum",
        "selected_hyperparameters": "Seçilen hiperparametreler",
        "stored_supporting_artifact": "Saklanan destekleyici artefakt",
        "validated_run_summary": "Doğrulanmış çalışma özeti",
        "dataset_accounting_purpose": "Veri kümesi sayımı ve adetler",
        "stored_model_metrics": "Saklanan model metrikleri",
        "stored_holdout_probabilities": "Saklanan holdout olasılıkları",
        "feature_availability_audit": "Özellik kullanılabilirlik denetimi",
        "label_reconciliation_review": "Etiket uzlaştırma incelemesi",
        "integrity_narrative": "Bütünlük anlatısı",
        "findings_markdown_baseline": "Bulgular yorumlama markdown temeli",
        "technical_markdown_baseline": "Teknik markdown temeli",
    },
}


def _phrase(lang: str, key: str) -> str:
    safe_lang = _lang(lang)
    return _REPORT_PHRASES.get(safe_lang, {}).get(key, _REPORT_PHRASES["en"].get(key, key))


def _now_text(lang: str) -> str:
    stamp = datetime.now().astimezone()
    if lang == "es":
        return stamp.strftime("%d/%m/%Y %H:%M")
    if lang == "tr":
        return stamp.strftime("%d.%m.%Y %H:%M")
    return stamp.strftime("%Y-%m-%d %H:%M")


def _format_file_timestamp(path: Path, lang: str) -> str:
    try:
        stamp = datetime.fromtimestamp(path.stat().st_mtime).astimezone()
    except Exception:
        return ""
    if lang == "es":
        return stamp.strftime("%d/%m/%Y %H:%M:%S")
    if lang == "tr":
        return stamp.strftime("%d.%m.%Y %H:%M:%S")
    return stamp.strftime("%Y-%m-%d %H:%M:%S")


def _lang(lang: str | None) -> str:
    safe = _clean_text(lang).lower()
    return safe if safe in {"en", "es", "tr"} else "en"


def _report_dir(run_id: str, lang: str) -> Path:
    return REPORT_ROOT / _clean_text(run_id) / _lang(lang)


def _report_filename(report_type: str, run_id: str) -> str:
    mapping = {
        "experiment_docx": f"classio_experiment_report_{run_id}_v{EXPERIMENT_REPORT_TEMPLATE_VERSION}.docx",
        "executive_docx": f"classio_eic_executive_report_{run_id}.docx",
        "academic_docx": f"classio_eic_academic_report_{run_id}.docx",
        "technical_docx": f"classio_eic_technical_report_{run_id}.docx",
    }
    return mapping[report_type]


def _read_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}
    return payload if isinstance(payload, dict) else {}


def _artifact_path_map(run_id: str) -> dict[str, Path]:
    rows = list_run_artifacts(run_id)
    mapping = {
        str(row.get("artifact_type") or ""): Path(str(row.get("storage_path") or ""))
        for row in rows
        if str(row.get("artifact_type") or "").strip() and str(row.get("storage_path") or "").strip()
    }
    if "findings_interpretation_report_md" not in mapping and "academic_report_md" in mapping:
        mapping["findings_interpretation_report_md"] = mapping["academic_report_md"]
    run_summary_path = mapping.get("run_summary_json")
    run_dir = run_summary_path.parent if run_summary_path and run_summary_path.exists() else None
    if run_dir:
        inferred = {
            "cluster_assignments_csv": run_dir / "resource_affinity_cluster_assignments.csv",
            "frozen_dataset_csv": run_dir / "resource_affinity_dataset_frozen.csv",
            "holdout_predictions_csv": run_dir / "resource_affinity_pairwise_neighbors.csv",
            "anchor_resource_candidates_csv": run_dir / "resource_affinity_program_topic_resource_candidates.csv",
            "exclusion_audit_csv": run_dir / "resource_affinity_exclusion_audit.csv",
            "category_normalization_audit_csv": run_dir / "resource_affinity_category_normalization_audit.csv",
            "human_review_sample_csv": run_dir / "resource_affinity_human_review_sample.csv",
            "experiment_config_json": run_dir / "resource_affinity_experiment_config.json",
            "representation_manifest_json": run_dir / "resource_affinity_representation_manifest.json",
        }
        for artifact_type, path in inferred.items():
            if artifact_type not in mapping and path.exists():
                mapping[artifact_type] = path
    return mapping


def _set_font(run, size: float, *, bold: bool = False, color: RGBColor | None = None, name: str = "Arial", italic: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _set_doc_defaults(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.space_after = Pt(7)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15
    style_map = {
        "Title": (22, CLASSIO_BLUE),
        "Subtitle": (11, MUTED),
        "Heading 1": (16, CLASSIO_BLUE),
        "Heading 2": (13, CLASSIO_BLUE),
        "Heading 3": (11.5, CLASSIO_TEAL),
    }
    for style_name, (size, color) in style_map.items():
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_after = Pt(6)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _apply_page_furniture(doc: Document, report_name: str, run_id: str, lang: str) -> None:
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hr = hp.add_run(f"Classio · {report_name} · {run_id}")
        _set_font(hr, 8.5, color=MUTED)
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(_copy(lang, "confidential") + " · ")
        _set_font(fr, 8.2, color=MUTED)
        _add_page_number(fp)


def _cover(doc: Document, title: str, subtitle: str, status: str, meta_lines: list[str], lang: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(78)
    run = p.add_run(_copy(lang, "cover_kicker"))
    _set_font(run, 11.5, bold=True, color=CLASSIO_GOLD)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(16)
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(title)
    _set_font(r2, 24, bold=True, color=CLASSIO_BLUE)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(16)
    r3 = p3.add_run(subtitle)
    _set_font(r3, 11, color=MUTED)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_after = Pt(16)
    r4 = p4.add_run(_copy(lang, "cover_status", value=status))
    _set_font(r4, 10.5, bold=True, color=CLASSIO_TEAL)

    for line in meta_lines:
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_after = Pt(4)
        mr = meta.add_run(line)
        _set_font(mr, 10, color=MUTED)
    doc.add_page_break()


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_paragraph(style=f"Heading {level}")
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run(text)
    _set_font(run, {1: 16, 2: 13, 3: 11.5}.get(level, 10.5), bold=True, color=CLASSIO_BLUE if level < 3 else CLASSIO_TEAL)


def _add_paragraph(doc: Document, text: str, *, italic: bool = False, color: RGBColor | None = None) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_font(run, 10.5, color=color or RGBColor(40, 40, 40), italic=italic)


def _add_bullets(doc: Document, rows: list[str]) -> None:
    for row in rows:
        if not _clean_text(row):
            continue
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(str(row))
        _set_font(run, 10.3, color=RGBColor(40, 40, 40))


def _unique_nonempty(rows: list[Any]) -> list[str]:
    seen: set[str] = set()
    values: list[str] = []
    for row in rows:
        text = _clean_text(row)
        if not text or text in seen:
            continue
        seen.add(text)
        values.append(text)
    return values


def _add_note_box(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_table_width(table, 6.7)
    cell = table.rows[0].cells[0]
    _set_cell_width(cell, 6.7)
    _shade_cell(cell, NOTE_FILL)
    _set_cell_margins(cell, top=110, bottom=110, left=140, right=140)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title)
    _set_font(r1, 10.2, bold=True, color=CLASSIO_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    _set_font(r2, 10, color=RGBColor(50, 50, 50))
    doc.add_paragraph("")


def _short_fingerprint(value: str) -> str:
    safe = _clean_text(value)
    return safe[:12] + "..." if len(safe) > 15 else safe


def _float(value: Any, default: float = 0.0) -> float:
    try:
        return float(value)
    except Exception:
        return default


def _int(value: Any, default: int = 0) -> int:
    try:
        return int(value)
    except Exception:
        return default


def _pct(value: Any) -> str:
    return f"{round(_float(value) * 100.0, 1)}%"


def _humanize_identifier(value: Any) -> str:
    safe = _clean_text(value)
    if not safe:
        return _copy("en", "not_available")
    tokens = safe.replace(".", " ").replace("_", " ").replace("-", " ").split()
    words: list[str] = []
    for token in tokens:
        lower = token.lower()
        if lower in {"id", "json", "csv", "md", "pdf", "docx", "roc", "auc", "f1"}:
            words.append(token.upper())
        elif lower.startswith("v") and lower[1:].isdigit():
            words.append(token.upper())
        elif token.isupper() and len(token) <= 5:
            words.append(token)
        else:
            words.append(token.capitalize())
    return " ".join(words)


def _display_scalar(value: Any) -> str:
    if value in (None, "", "None"):
        return "n/a"
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, float):
        return str(round(value, 4))
    return _humanize_identifier(value) if isinstance(value, str) and ("_" in value or "." in value or "-" in value) else str(value)


def _display_jsonish(value: Any, *, max_chars: int = 160) -> str:
    payload = value
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return "{}"
        try:
            payload = json.loads(text)
        except Exception:
            return text[:max_chars]
    if isinstance(payload, dict):
        pieces = [f"{_humanize_identifier(key)}: {_display_scalar(item)}" for key, item in payload.items()]
        rendered = "; ".join(pieces) if pieces else "{}"
        return rendered[:max_chars]
    if isinstance(payload, list):
        rendered = ", ".join(_display_scalar(item) for item in payload)
        return rendered[:max_chars]
    return str(payload)[:max_chars]


def _humanize_list(values: list[Any]) -> str:
    return ", ".join(_humanize_identifier(value) for value in values if _clean_text(value))


def _artifact_display_name(artifact_type: str, path: Path | None = None) -> str:
    filename = str((path or Path("")).name)
    if filename.startswith("resource_affinity_"):
        affinity_mapping = {
            "label_audit_csv": "Profile Audit CSV",
            "holdout_predictions_csv": "Pairwise Semantic Neighbors CSV",
            "anchor_resource_candidates_csv": "Program Topic Resource Candidates CSV",
            "label_reconciliation_csv": "Unsupervised Validation Reconciliation CSV",
            "resource_affinity_program_topic_resource_candidates.csv": "Program Topic Resource Candidates CSV",
            "resource_affinity_exclusion_audit.csv": "Exclusion Audit CSV",
            "resource_affinity_category_normalization_audit.csv": "Category Normalization Audit CSV",
            "resource_affinity_human_review_sample.csv": "Human Review Sample CSV",
            "resource_affinity_experiment_config.json": "Experiment Configuration JSON",
            "resource_affinity_representation_manifest.json": "Fitted Representation Manifest JSON",
        }
        specific = affinity_mapping.get(_clean_text(artifact_type)) or affinity_mapping.get(filename)
        if specific:
            return specific
    mapping = {
        "academic_report_md": "Findings Interpretation Report Markdown",
        "findings_interpretation_report_md": "Findings Interpretation Report Markdown",
        "dataset_summary_json": "Dataset Summary JSON",
        "feature_audit_csv": "Feature Audit CSV",
        "frozen_dataset_csv": "Frozen Dataset CSV",
        "cluster_assignments_csv": "Cluster Assignments CSV",
        "holdout_predictions_csv": "Holdout Predictions CSV",
        "integrity_review_md": "Integrity Review Markdown",
        "integrity_report_md": "Integrity Report Markdown",
        "label_audit_csv": "Label Audit CSV",
        "label_reconciliation_csv": "Label Reconciliation CSV",
        "model_comparison_csv": "Model Comparison CSV",
        "run_summary_json": "Run Summary JSON",
        "technical_report_md": "Technical Report Markdown",
    }
    return mapping.get(_clean_text(artifact_type), _humanize_identifier(artifact_type))


def _hash_file(path: Path) -> str:
    if not path.exists():
        return "missing"
    digest = hashlib.sha256(path.read_bytes()).hexdigest()
    return digest[:10]


def _set_table_width(table, width_in: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_in * 1440)))


def _section_content_width(doc: Document) -> float:
    section = doc.sections[-1]
    usable = section.page_width - section.left_margin - section.right_margin
    return max(4.5, usable / 914400.0)


def _fit_table_widths(widths: list[float], max_width: float) -> list[float]:
    if not widths:
        return widths
    total = sum(widths)
    if total <= 0:
        return widths
    target = max(3.5, max_width - TABLE_WIDTH_GUTTER)
    if total <= target:
        return widths
    scale = target / total
    return [round(width * scale, 4) for width in widths]


def _set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))


def _set_cell_margins(cell, *, top: int = 80, bottom: int = 80, left: int = 120, right: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tag = tr_pr.find(qn("w:tblHeader"))
    if tag is None:
        tag = OxmlElement("w:tblHeader")
        tr_pr.append(tag)
    tag.set(qn("w:val"), "true")


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tag = tr_pr.find(qn("w:cantSplit"))
    if tag is None:
        tag = OxmlElement("w:cantSplit")
        tr_pr.append(tag)


def _configure_table(table, widths: list[float], *, header_fill: str = TABLE_HEADER_FILL) -> list[float]:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_table_width(table, sum(widths))
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        _prevent_row_split(row)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            _set_cell_width(cell, width)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing = 1.05
    if table.rows:
        _mark_header_row(table.rows[0])
        for cell in table.rows[0].cells:
            _shade_cell(cell, header_fill)
    return widths


def _add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    *,
    caption: str = "",
    numeric_cols: set[int] | None = None,
) -> None:
    numeric_cols = numeric_cols or set()
    fitted_widths = _fit_table_widths(widths, _section_content_width(doc))
    header_font_size = 9.6 if len(headers) <= 5 else 9.1
    body_font_size = 9.2 if len(headers) <= 5 else 8.8
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(header))
        _set_font(r, header_font_size, bold=True, color=CLASSIO_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            _set_font(r, body_font_size, color=RGBColor(45, 45, 45))
    _configure_table(table, fitted_widths)
    if caption:
        _add_caption(doc, caption)
    doc.add_paragraph("")


def _add_caption(doc: Document, caption: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(caption)
    _set_font(run, 8.8, italic=True, color=MUTED)


def _add_picture_with_caption(doc: Document, image_path: Path, caption: str, *, width: float = 6.5) -> None:
    if not image_path.exists():
        return
    doc.add_picture(str(image_path), width=Inches(width))
    _add_caption(doc, caption)


def _add_landscape_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def _add_portrait_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def _verified_docx_bytes(path: Path) -> bytes:
    data = path.read_bytes()
    with ZipFile(BytesIO(data), "r") as handle:
        if "word/document.xml" not in handle.namelist():
            raise BadZipFile("Missing word/document.xml")
    return data


def _validated_run_detail(run_id: str) -> dict[str, Any]:
    detail = eic_service.get_experiment_business_detail(run_id, cache_bust=f"report-{run_id}")
    if detail and _clean_text(detail.get("run_status")) in FINAL_VALIDATED_RUN_STATES:
        return detail
    artifacts = _artifact_path_map(run_id)
    run_summary = _read_json(artifacts.get("run_summary_json", Path("__missing__")))
    dataset = run_summary.get("dataset") or {}
    evaluation = run_summary.get("evaluation") or {}
    review = run_summary.get("review") or {}
    final_verdict = _clean_text(review.get("final_verdict") or "VALIDATED_EXPLORATORY_RUN")
    if not run_summary or final_verdict not in FINAL_VALIDATED_RUN_STATES:
        return {}
    best = evaluation.get("best_model") or {}
    return {
        "run_id": run_id,
        "experiment_id": _clean_text(evaluation.get("experiment_id")) or "resource_affinity_unsupervised_discovery",
        "run_status": final_verdict,
        "integrity_status": "approved_exploratory",
        "maturity_verdict": _clean_text(evaluation.get("maturity_verdict")),
        "evidence_level": _clean_text(evaluation.get("overall_evidence_strength")),
        "evidence_verdict": _clean_text(evaluation.get("overall_evidence_strength")),
        "business_question": _clean_text(evaluation.get("target_definition")),
        "recommended_business_action": "move_to_shadow_testing" if _clean_text(evaluation.get("maturity_verdict")) == "CANDIDATE_FOR_SHADOW_TESTING" else "continue_collecting_data",
        "included_row_count": int(dataset.get("included_row_count") or 0),
        "positive_label_count": 0,
        "negative_label_count": 0,
        "teachers_represented": 0,
        "students_represented": 0,
        "resources_represented": int(dataset.get("candidate_resource_count") or dataset.get("resource_count") or 0),
        "primary_metric_leader": _clean_text(evaluation.get("winner") or best.get("model_name")),
        "primary_metric": _clean_text(evaluation.get("primary_metric")),
        "dataset_fingerprint": _clean_text(dataset.get("data_fingerprint")),
        "model_results": {
            "models_compared": [],
            "overall_evidence_conclusion": _clean_text(evaluation.get("overall_evidence_strength")),
        },
        "limitations": list(evaluation.get("limitations") or []),
    }


def _report_base_context(run_id: str, lang: str) -> dict[str, Any]:
    detail = _validated_run_detail(run_id)
    if not detail:
        return {}
    experiment_id = str(detail.get("experiment_id") or "")
    try:
        academic = eic_service.get_academic_evidence_summary(run_id, cache_bust=f"academic-{run_id}")
    except Exception:
        academic = {}
    try:
        telemetry = eic_service.get_business_telemetry_health(cache_bust=f"telemetry-{run_id}")
    except Exception:
        telemetry = {}
    try:
        portfolio = eic_service.get_intelligence_component_portfolio(cache_bust=f"portfolio-{run_id}")
    except Exception:
        portfolio = []
    try:
        decisions = eic_service.get_prioritized_intelligence_decisions(cache_bust=f"decisions-{run_id}")
    except Exception:
        decisions = []
    try:
        latest_summary = eic_service.get_intelligence_business_summary(cache_bust=f"summary-{run_id}")
    except Exception:
        latest_summary = {}
    artifacts = _artifact_path_map(run_id)
    run_summary = _read_json(artifacts.get("run_summary_json", Path("__missing__")))
    dataset_summary = _read_json(artifacts.get("dataset_summary_json", Path("__missing__")))
    report_context = get_report_context(run_id, experiment_id, lang)
    return {
        "detail": detail,
        "academic": academic,
        "telemetry": telemetry,
        "portfolio": portfolio,
        "decisions": decisions,
        "latest_summary": latest_summary,
        "run_summary": run_summary,
        "dataset_summary": dataset_summary,
        "report_context": report_context,
        "report_context_completion": report_context_completion(report_context),
        "lang": lang,
    }


def _generate_report_ai_narrative(context: dict[str, Any], *, report_kind: str, lang: str) -> dict[str, Any]:
    try:
        from helpers.lesson_planner import (
            _extract_json_object_from_text,
            _generate_with_gemini,
            _generate_with_openai,
            _generate_with_openrouter,
            get_ai_provider_order,
        )
    except Exception:
        return {}

    detail = context.get("detail") or {}
    academic = context.get("academic") or {}
    run_summary = context.get("run_summary") or {}
    dataset_summary = context.get("dataset_summary") or {}
    evaluation = run_summary.get("evaluation") or {}
    review = run_summary.get("review") or {}

    language_name = {"en": "English", "es": "Spanish", "tr": "Turkish"}.get(_lang(lang), "English")
    system_prompt = (
        "You are Classio's internal experiment reporting analyst. "
        "Return exactly one valid JSON object and nothing else. "
        "Do not use markdown. Do not use code fences. "
        "Write all narrative text in the requested report language. "
        "Be precise, evidence-grounded, concise, and professional. "
        "Do not invent metrics, counts, or causal claims. "
        "If evidence is weak, say so clearly."
    )
    prompt_payload = {
        "report_kind": report_kind,
        "report_language": language_name,
        "experiment": str(detail.get("experiment_id") or ""),
        "run_id": str(detail.get("run_id") or ""),
        "run_status": str(detail.get("run_status") or ""),
        "integrity_status": str(detail.get("integrity_status") or ""),
        "maturity_verdict": str(detail.get("maturity_verdict") or ""),
        "evidence_verdict": str(detail.get("evidence_level") or detail.get("evidence_verdict") or ""),
        "business_question": str(detail.get("business_question") or ""),
        "recommended_business_action": str(detail.get("recommended_business_action") or ""),
        "included_rows": int(detail.get("included_row_count") or dataset_summary.get("included_row_count") or 0),
        "positive_labels": int(detail.get("positive_label_count") or dataset_summary.get("positive_count") or 0),
        "negative_labels": int(detail.get("negative_label_count") or dataset_summary.get("negative_count") or 0),
        "teachers_represented": int(detail.get("teachers_represented") or dataset_summary.get("teacher_count") or 0),
        "students_represented": int(detail.get("students_represented") or dataset_summary.get("student_count") or 0),
        "resources_represented": int(detail.get("resources_represented") or dataset_summary.get("resource_count") or 0),
        "primary_metric_leader": str(evaluation.get("primary_metric_leader") or detail.get("primary_metric_leader") or ""),
        "best_thresholded_classifier": str(evaluation.get("best_thresholded_classifier") or ""),
        "best_precision_recall_ranking": str(evaluation.get("best_precision_recall_ranking") or ""),
        "calibration_leader": str(evaluation.get("calibration_leader") or ""),
        "limitations": _unique_nonempty(list(academic.get("limitations") or []) + list(detail.get("limitations") or []) + list((review.get("label_reconciliation") or {}).get("limitations") or [])),
        "future_work": list(academic.get("future_improvements") or []),
        "validation_notes": str(detail.get("validation_notes") or ""),
        "report_context": context.get("report_context") or {},
    }
    user_prompt = (
        "Create a JSON object with these keys: "
        "analysis_paragraph, conclusion_paragraph, implementation_paragraph, decision_summary_paragraph, non_proof_paragraph, main_limitation_text, proposed_next_action_text, context_rewrites, limitations, future_work. "
        "The context_rewrites field must be an object. Rewrite only the user-provided context values into polished report language for these keys when present: "
        "business_problem, decision_supported, expected_value, product_impact, success_definition, minimum_evidence_required, risks, main_limitation, evidence_non_proof, recommended_next_action, next_review_trigger, next_review_date, responsible_person_or_team, meeting_notes. "
        "The limitations and future_work fields must be arrays of short strings. "
        "Do not invent missing business context. If a context value is missing, leave that key empty in context_rewrites. "
        "Base every statement on this experiment context only:\n"
        + json.dumps(prompt_payload, ensure_ascii=False, indent=2)
    )

    provider_order = get_ai_provider_order()
    log_ai_usage_event(
        "experiment_report_ai",
        "requested",
        with_provider_chain({
            "report_kind": report_kind,
            "language": lang,
            "experiment_id": str(detail.get("experiment_id") or ""),
            "run_id": str(detail.get("run_id") or ""),
        }, provider_order),
    )
    errors: list[str] = []
    for provider in provider_order:
        try:
            if provider == "gemini":
                raw_text = _generate_with_gemini(system_prompt, user_prompt)
            elif provider == "openrouter":
                raw_text = _generate_with_openrouter(system_prompt, user_prompt)
            else:
                raw_text = _generate_with_openai(system_prompt, user_prompt)
            parsed = _extract_json_object_from_text(raw_text)
            if not isinstance(parsed, dict):
                raise ValueError("AI narrative payload was not an object.")
            log_ai_usage_event(
                "experiment_report_ai",
                "success",
                {
                    "report_kind": report_kind,
                    "language": lang,
                    "experiment_id": str(detail.get("experiment_id") or ""),
                    "run_id": str(detail.get("run_id") or ""),
                    "provider": provider,
                },
            )
            return {
                "analysis_paragraph": _clean_text(parsed.get("analysis_paragraph")),
                "conclusion_paragraph": _clean_text(parsed.get("conclusion_paragraph")),
                "implementation_paragraph": _clean_text(parsed.get("implementation_paragraph")),
                "decision_summary_paragraph": _clean_text(parsed.get("decision_summary_paragraph")),
                "non_proof_paragraph": _clean_text(parsed.get("non_proof_paragraph")),
                "main_limitation_text": _clean_text(parsed.get("main_limitation_text")),
                "proposed_next_action_text": _clean_text(parsed.get("proposed_next_action_text")),
                "context_rewrites": dict(parsed.get("context_rewrites") or {}),
                "limitations": _unique_nonempty(list(parsed.get("limitations") or [])),
                "future_work": _unique_nonempty(list(parsed.get("future_work") or [])),
                "_ai_used": True,
                "_provider": provider,
            }
        except Exception as exc:
            errors.append(f"{provider}: {exc}")
    log_ai_usage_event(
        "experiment_report_ai",
        "failed",
        with_provider_chain({
            "report_kind": report_kind,
            "language": lang,
            "experiment_id": str(detail.get("experiment_id") or ""),
            "run_id": str(detail.get("run_id") or ""),
            "errors": errors,
        }, provider_order),
    )
    return {}


def _report_chart_assets(report_type: str, run_id: str, lang: str, context: dict[str, Any]) -> dict[str, Path]:
    assets_dir = _report_dir(run_id, lang) / f"assets_{report_type}"
    assets_dir.mkdir(parents=True, exist_ok=True)
    return build_report_charts(report_type, run_id, lang, context, _artifact_path_map(run_id), assets_dir)


def _available_or_message(required: bool, lang: str) -> dict[str, Any] | None:
    if required:
        return None
    return {"ok": False, "message": t("admin_eic_report_unavailable_no_validated_run", lang=lang)}


def _append_source_note(doc: Document, lang: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(10)
    run = para.add_run(_copy(lang, "source_note"))
    _set_font(run, 8.7, italic=True, color=MUTED)


def _bool_text(value: Any, lang: str) -> str:
    normalized = _clean_text(value).lower()
    return _copy(lang, "yes") if normalized in {"1", "true", "yes"} else _copy(lang, "no")


def _model_name(value: str, lang: str) -> str:
    return str(get_model_comparison_value_display("model_name", value, lang=lang))


def _narrative_robust_winner(detail: dict[str, Any], run_summary: dict[str, Any], lang: str) -> str:
    evaluation = run_summary.get("evaluation") or {}
    leader = _model_name(_clean_text(evaluation.get("primary_metric_leader") or detail.get("primary_metric_leader")), lang)
    if _clean_text(detail.get("run_status")) == "VALIDATED_EXPLORATORY_RUN":
        if lang == "es":
            return f"{leader} lideró la comparación validada con suficiente solidez como para justificar un shadow testing controlado."
        if lang == "tr":
            return f"{leader}, kontrollü shadow testini haklı çıkaracak kadar güçlü şekilde doğrulanmış karşılaştırmaya liderlik etti."
        return f"{leader} led the validated comparison strongly enough to justify controlled shadow testing."
    return _copy(lang, "academic_conclusion")


def _required_artifacts_present(report_type: str, context: dict[str, Any], artifacts: dict[str, Path]) -> bool:
    if report_type == "executive_docx":
        return bool(context.get("detail") and context.get("portfolio") and context.get("decisions"))
    if report_type == "academic_docx":
        return bool(context.get("run_summary") and context.get("dataset_summary") and ((artifacts.get("model_comparison_csv") and artifacts["model_comparison_csv"].exists()) or (context.get("detail", {}).get("model_results") or {}).get("models_compared")))
    if report_type == "technical_docx":
        artifact_complete = all(
            artifacts.get(key) and artifacts[key].exists()
            for key in ["run_summary_json", "dataset_summary_json", "model_comparison_csv", "holdout_predictions_csv", "feature_audit_csv"]
        )
        context_complete = bool(
            context.get("run_summary")
            and context.get("dataset_summary")
            and ((context.get("detail", {}).get("model_results") or {}).get("models_compared"))
        )
        return artifact_complete or context_complete
    return False


def _metadata_rows(detail: dict[str, Any], run_summary: dict[str, Any], dataset_summary: dict[str, Any], lang: str) -> list[list[str]]:
    dataset = run_summary.get("dataset") or dataset_summary
    evaluation = run_summary.get("evaluation") or {}
    return [
        [_copy(lang, "meta_field"), _copy(lang, "meta_value")],
        [t("admin_eic_run_status", lang=lang), get_run_status_display(str(detail.get("run_status") or ""), lang=lang)],
        [t("admin_eic_integrity_status", lang=lang), get_integrity_status_display(str(detail.get("integrity_status") or ""), lang=lang)],
        [t("admin_eic_registry_maturity", lang=lang), get_maturity_display(str(detail.get("maturity_verdict") or ""), lang=lang)],
        [t("admin_eic_report_target_version", lang=lang), _humanize_identifier(dataset_summary.get("target_version") or "opened_within_7d_v1")],
        [t("admin_eic_report_feature_schema", lang=lang), _humanize_identifier(dataset_summary.get("feature_schema_version") or dataset.get("feature_schema_version") or "n/a")],
        [t("admin_eic_report_run_id", lang=lang, value="").split(":")[0], str(detail.get("run_id") or "")],
        [t("admin_eic_report_dataset_fingerprint", lang=lang, value="").split(":")[0], str(detail.get("dataset_fingerprint") or "")],
        [_phrase(lang, "chronological_cutoff"), str(detail.get("chronological_cutoff") or evaluation.get("cutoff_timestamp") or "n/a")],
    ]


def _dataset_accounting_rows(detail: dict[str, Any], run_summary: dict[str, Any], dataset_summary: dict[str, Any], lang: str) -> list[list[str]]:
    dataset = run_summary.get("dataset") or dataset_summary
    evaluation = run_summary.get("evaluation") or {}
    rows = [
        [_phrase(lang, "source_rows"), str(dataset.get("source_row_count") or "n/a")],
        [_phrase(lang, "included_rows"), str(dataset.get("included_row_count") or detail.get("included_row_count") or "n/a")],
        [_phrase(lang, "positive_labels"), str(dataset.get("positive_count") or detail.get("positive_label_count") or "n/a")],
        [_phrase(lang, "negative_labels"), str(dataset.get("negative_count") or detail.get("negative_label_count") or "n/a")],
        [_phrase(lang, "right_censored"), str(dataset.get("excluded_row_count") or "0")],
        [_phrase(lang, "teachers_represented"), str(dataset.get("teacher_count") or detail.get("teachers_represented") or "0")],
        [_phrase(lang, "students_represented"), str(dataset.get("student_count") or detail.get("students_represented") or "0")],
        [_phrase(lang, "resources_represented"), str(dataset.get("resource_count") or detail.get("resources_represented") or "0")],
        [_phrase(lang, "date_range"), f"{(dataset.get('date_range') or {}).get('assigned_at_min') or detail.get('source_start_at') or 'n/a'} {_date_connector(lang)} {(dataset.get('date_range') or {}).get('assigned_at_max') or detail.get('source_end_at') or 'n/a'}"],
        [_phrase(lang, "chronological_cutoff"), str(evaluation.get("cutoff_timestamp") or detail.get("chronological_cutoff") or "n/a")],
        [_phrase(lang, "development_rows"), str(evaluation.get("development_count") or "n/a")],
        [_phrase(lang, "holdout_rows"), str(evaluation.get("holdout_count") or "n/a")],
    ]
    for reason, count in (dataset.get("exclusion_reasons") or {}).items():
        if reason == "included":
            continue
        rows.append([f"{_humanize_identifier(reason)}", str(count)])
    return rows


def _portfolio_rows(portfolio: list[dict[str, Any]], lang: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in portfolio:
        rows.append(
            [
                get_component_display_name(str(row.get("component_id") or ""), lang=lang),
                get_component_type_display(str(row.get("component_type") or ""), lang=lang),
                str(row.get("production_use") or row.get("decision_supported") or ""),
                get_evidence_display(str(row.get("evidence_maturity") or row.get("validated_evidence_status") or ""), lang=lang),
                get_business_action_display(str(row.get("recommended_next_action") or ""), lang=lang),
            ]
        )
    return rows


def _decision_rows(decisions: list[dict[str, Any]], lang: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in decisions:
        issue = str(row.get("issue") or "")
        evidence = str(row.get("evidence") or "")
        impact = str(row.get("business_impact") or "")
        rows.append(
            [
                str(row.get("urgency") or "").title() or _copy(lang, "not_available"),
                get_component_display_name(str(row.get("component_id") or ""), lang=lang),
                evidence,
                impact,
                get_business_action_display(str(row.get("recommended_action") or ""), lang=lang),
                str(row.get("responsible_area") or ""),
                issue,
            ]
        )
    return rows


def _feature_health_rows(run_summary: dict[str, Any], lang: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in (run_summary.get("review") or {}).get("feature_health") or []:
        rows.append(
            [
                _humanize_identifier(row.get("feature") or ""),
                _phrase(lang, "stored_review"),
                _phrase(lang, "past_only"),
                f"{round(_float(row.get('overall_missing_percentage')), 2)}%",
                f"{round(_float(row.get('development_missing_percentage')), 2)}%",
                f"{round(_float(row.get('holdout_missing_percentage')), 2)}%",
                _bool_text(not row.get("excluded_from_logistic_regression_reduced"), lang),
                _translate_known_evidence_text(row.get("missingness_explanation") or "", lang),
            ]
        )
    return rows


def _model_metric_groups(model_rows: list[dict[str, Any]], lang: str) -> list[tuple[str, list[str], list[list[str]], list[float], set[int]]]:
    group_specs = [
        (
            _phrase(lang, "discrimination_ranking"),
            ["model_name", "roc_auc", "average_precision", "balanced_accuracy", "f1", "delta_vs_dummy_roc_auc"],
            [1, 2, 3, 4, 5],
            [2.2, 1.2, 1.2, 1.2, 1.0, 1.5],
        ),
        (
            _phrase(lang, "calibration_thresholding"),
            ["model_name", "log_loss", "brier_score", "predicted_positive_rate", "best_f1", "best_balanced_accuracy"],
            [1, 2, 3, 4, 5],
            [2.1, 1.2, 1.2, 1.4, 1.0, 1.5],
        ),
        (
            _phrase(lang, "runtime_status"),
            ["model_name", "status", "cv_primary_metric_mean", "cv_primary_metric_variance", "train_duration_seconds", "inference_duration_seconds"],
            [2, 3, 4, 5],
            [2.1, 1.2, 1.5, 1.5, 1.5, 1.6],
        ),
    ]
    groups = []
    for title, columns, numeric_cols, widths in group_specs:
        headers = [get_model_comparison_column_display(column, lang=lang) for column in columns]
        rows: list[list[str]] = []
        for row in model_rows:
            rendered = [str(get_model_comparison_value_display(column, row.get(column), lang=lang) or "—") for column in columns]
            rows.append(rendered)
        groups.append((title, headers, rows, widths, set(numeric_cols)))
    return groups


def _artifact_manifest_rows(artifacts: dict[str, Path], lang: str) -> tuple[list[list[str]], list[str]]:
    purpose_map = {
        "run_summary_json": _phrase(lang, "validated_run_summary"),
        "dataset_summary_json": _phrase(lang, "dataset_accounting_purpose"),
        "model_comparison_csv": _phrase(lang, "stored_model_metrics"),
        "holdout_predictions_csv": _phrase(lang, "stored_holdout_probabilities"),
        "anchor_resource_candidates_csv": "Program-topic to candidate-resource alignment candidates" if lang == "en" else ("Candidatos de alineación entre tema del programa y recurso candidato" if lang == "es" else "Program konusu ile aday kaynak hizalama adayları"),
        "feature_audit_csv": _phrase(lang, "feature_availability_audit"),
        "cluster_assignments_csv": "Winning clustering assignments" if lang == "en" else ("Asignaciones de clúster del modelo ganador" if lang == "es" else "Kazanan küme atamaları"),
        "label_reconciliation_csv": _phrase(lang, "label_reconciliation_review"),
        "integrity_report_md": _phrase(lang, "integrity_narrative"),
        "academic_report_md": _phrase(lang, "findings_markdown_baseline"),
        "findings_interpretation_report_md": _phrase(lang, "findings_markdown_baseline"),
        "technical_report_md": _phrase(lang, "technical_markdown_baseline"),
    }
    rows: list[list[str]] = []
    filenames: list[str] = []
    for artifact_type, path in sorted(artifacts.items()):
        filenames.append(path.name)
        rows.append(
            [
                _artifact_display_name(artifact_type, path),
                purpose_map.get(artifact_type, _phrase(lang, "stored_supporting_artifact")),
                path.suffix.lstrip(".") or "file",
                _hash_file(path),
                _copy(lang, "yes") if path.exists() else _copy(lang, "no"),
            ]
        )
    return rows, filenames


def _editable_placeholder(lang: str) -> str:
    if lang == "es":
        return "[Por completar por la persona responsable de revisión en Classio.]"
    if lang == "tr":
        return "[Classio'da sorumlu inceleyici tarafından doldurulacak.]"
    return "[To be completed by the responsible Classio reviewer.]"


def _context_label(key: str, lang: str) -> str:
    return t(f"report_context_{key}", lang=lang)


def _context_text(context: dict[str, Any], key: str, lang: str, ai_narrative: dict[str, Any] | None = None) -> str:
    polished = _clean_text((((ai_narrative or {}).get("context_rewrites") or {}).get(key)))
    if polished:
        return polished
    value = _clean_text((context or {}).get(key))
    return value or _editable_placeholder(lang)


def _ai_used_text(ai_narrative: dict[str, Any], lang: str) -> str:
    values: list[str] = []
    for value in (ai_narrative or {}).values():
        if isinstance(value, list):
            continue
        if isinstance(value, dict):
            values.extend(_clean_text(item) for item in value.values())
        else:
            values.append(_clean_text(value))
    return _copy(lang, "yes") if any(values) else _copy(lang, "no")


def _report_context_missing_labels(context: dict[str, Any], lang: str) -> list[str]:
    completion = context.get("report_context_completion") or report_context_completion(context.get("report_context") or context)
    return [_context_label(field, lang) for field in list(completion.get("missing_fields") or [])]


def _scorecard_status(value: str, lang: str) -> str:
    normalized = _clean_text(value).lower()
    labels = {
        "strong": {"en": "Strong", "es": "Fuerte", "tr": "Güçlü"},
        "moderate": {"en": "Moderate", "es": "Moderada", "tr": "Orta"},
        "limited": {"en": "Limited", "es": "Limitada", "tr": "Sınırlı"},
        "insufficient": {"en": "Insufficient", "es": "Insuficiente", "tr": "Yetersiz"},
        "incomplete": {"en": "Incomplete", "es": "Incompleta", "tr": "Incompleta"},
        "not assessed": {"en": "Not assessed", "es": "No evaluada", "tr": "Değerlendirilmedi"},
    }
    return labels.get(normalized, {}).get(lang, value)


def _decision_summary_rows(context: dict[str, Any], ai_narrative: dict[str, Any], lang: str) -> list[list[str]]:
    detail = context.get("detail") or {}
    run_summary = context.get("run_summary") or {}
    report_context = context.get("report_context") or {}
    evaluation = run_summary.get("evaluation") or {}
    limitations = _unique_nonempty(
        list((context.get("academic") or {}).get("limitations") or [])
        + list(detail.get("limitations") or [])
        + list(((run_summary.get("review") or {}).get("label_reconciliation") or {}).get("limitations") or [])
    )
    localized_limitations = _localized_bullets(limitations, lang)
    return [
        [_copy(lang, "meta_field"), _copy(lang, "meta_value")],
        ["Experiment" if lang == "en" else ("Experimento" if lang == "es" else "Deney"), str(detail.get("experiment_id") or "")],
        ["Business problem" if lang == "en" else ("Problema de negocio" if lang == "es" else "İş problemi"), _context_text(report_context, "business_problem", lang, ai_narrative)],
        ["Decision supported" if lang == "en" else ("Decisión apoyada" if lang == "es" else "Desteklenen karar"), _context_text(report_context, "decision_supported", lang, ai_narrative)],
        ["Current evidence status" if lang == "en" else ("Estado actual de la evidencia" if lang == "es" else "Mevcut kanıt durumu"), get_evidence_display(str(detail.get("evidence_level") or detail.get("evidence_verdict") or "not_available"), lang=lang)],
        ["Robust winner" if lang == "en" else ("Ganador robusto" if lang == "es" else "Güçlü kazanan"), _bool_text((detail.get("model_results") or {}).get("robust_winner"), lang)],
        ["Model leader" if lang == "en" else ("Modelo líder" if lang == "es" else "Lider model"), _model_name(str(evaluation.get("primary_metric_leader") or detail.get("primary_metric_leader") or ""), lang)],
        ["Next review trigger" if lang == "en" else ("Disparador de la próxima revisión" if lang == "es" else "Sonraki inceleme tetikleyicisi"), _context_text(report_context, "next_review_trigger", lang, ai_narrative)],
        ["Responsible person or team" if lang == "en" else ("Persona o equipo responsable" if lang == "es" else "Sorumlu kişi veya ekip"), _context_text(report_context, "responsible_person_or_team", lang, ai_narrative)],
        ["Main limitation" if lang == "en" else ("Limitación principal" if lang == "es" else "Ana sınırlama"), _clean_text(ai_narrative.get("main_limitation_text")) or _context_text(report_context, "main_limitation", lang, ai_narrative) or (localized_limitations[0] if localized_limitations else _editable_placeholder(lang))],
    ]


def _readiness_scorecard_rows(context: dict[str, Any], lang: str) -> list[list[str]]:
    detail = context.get("detail") or {}
    run_summary = context.get("run_summary") or {}
    dataset_summary = context.get("dataset_summary") or {}
    completion = context.get("report_context_completion") or {}
    telemetry_summary = ((context.get("telemetry") or {}).get("summary") or {})
    included = int(detail.get("included_row_count") or dataset_summary.get("included_row_count") or 0)
    teachers = int(detail.get("teachers_represented") or dataset_summary.get("teacher_count") or 0)
    feature_health = (run_summary.get("review") or {}).get("feature_health") or []
    has_feature_health = bool(feature_health)
    evidence = str(detail.get("evidence_level") or detail.get("evidence_verdict") or "")
    rows = [
        ["Dimension" if lang == "en" else ("Dimensión" if lang == "es" else "Boyut"), "Status" if lang == "en" else ("Estado" if lang == "es" else "Durum"), _copy(lang, "interpretation")],
        ["Data coverage" if lang == "en" else ("Cobertura de datos" if lang == "es" else "Veri kapsamı"), _scorecard_status("moderate" if included >= 100 else "limited", lang), (f"{included} labels" if lang == "en" else (f"{included} etiquetas" if lang == "es" else f"{included} etiket"))],
        ["Teacher diversity" if lang == "en" else ("Diversidad docente" if lang == "es" else "Öğretmen çeşitliliği"), _scorecard_status("limited" if teachers <= 1 else "moderate", lang), (f"{teachers} teacher(s)" if lang == "en" else (f"{teachers} docente(s)" if lang == "es" else f"{teachers} öğretmen"))],
        ["Feature readiness" if lang == "en" else ("Preparación de variables" if lang == "es" else "Özellik hazırlığı"), _scorecard_status("moderate" if has_feature_health else "not assessed", lang), ("Stored feature audit available." if lang == "en" else ("La auditoría almacenada de variables está disponible." if lang == "es" else "Saklanan özellik denetimi mevcut.")) if has_feature_health else _editable_placeholder(lang)],
        ["Model discrimination" if lang == "en" else ("Discriminación del modelo" if lang == "es" else "Model ayrıştırması"), _scorecard_status("moderate" if evidence in {"exploratory", "validated"} else "insufficient", lang), _narrative_robust_winner(detail, run_summary, lang)],
        ["Calibration" if lang == "en" else ("Calibración" if lang == "es" else "Kalibrasyon"), _scorecard_status("not assessed" if not (run_summary.get("evaluation") or {}).get("calibration_leader") else "moderate", lang), _model_name(str((run_summary.get("evaluation") or {}).get("calibration_leader") or ""), lang) or _editable_placeholder(lang)],
        ["Production readiness" if lang == "en" else ("Preparación para producción" if lang == "es" else "Üretim hazırlığı"), _scorecard_status("insufficient", lang), get_run_status_display(str(detail.get("run_status") or ""), lang=lang)],
        ["Business-context completeness" if lang == "en" else ("Completitud del contexto de negocio" if lang == "es" else "İş bağlamı bütünlüğü"), _scorecard_status("moderate" if bool(completion.get("complete")) else "incomplete", lang), f"{int(completion.get('completed_fields') or 0)}/{int(completion.get('total_fields') or 0)}"],
    ]
    if int(telemetry_summary.get("unmatched_opens") or 0) > 0:
        rows.append(["Telemetry quality" if lang == "en" else ("Calidad de telemetría" if lang == "es" else "Telemetri kalitesi"), _scorecard_status("limited", lang), (f"{int(telemetry_summary.get('unmatched_opens') or 0)} unmatched opens" if lang == "en" else (f"{int(telemetry_summary.get('unmatched_opens') or 0)} aperturas sin emparejar" if lang == "es" else f"{int(telemetry_summary.get('unmatched_opens') or 0)} eşleşmeyen açılma"))])
    return rows


def _action_plan_rows(context: dict[str, Any], lang: str) -> list[list[str]]:
    detail = context.get("detail") or {}
    report_context = context.get("report_context") or {}
    run_summary = context.get("run_summary") or {}
    dataset_summary = context.get("dataset_summary") or {}
    teachers = int(detail.get("teachers_represented") or dataset_summary.get("teacher_count") or 0)
    included = int(detail.get("included_row_count") or dataset_summary.get("included_row_count") or 0)
    rows = [[
        _copy(lang, "priority"),
        "Action" if lang == "en" else ("Acción" if lang == "es" else "Eylem"),
        "Reason" if lang == "en" else ("Razón" if lang == "es" else "Gerekçe"),
        _copy(lang, "owner"),
        "Trigger or deadline" if lang == "en" else ("Disparador o fecha" if lang == "es" else "Tetikleyici veya tarih"),
        _phrase(lang, "status_col"),
    ]]
    owner = _context_text(report_context, "responsible_person_or_team", lang)
    if teachers <= 1:
        rows.append(["High", "Expand teacher coverage" if lang == "en" else ("Ampliar cobertura docente" if lang == "es" else "Öğretmen kapsamını genişlet"), "Current evidence comes from one teacher." if lang == "en" else ("La evidencia actual proviene de un solo docente." if lang == "es" else "Mevcut kanıt tek bir öğretmenden geliyor."), owner, "At least five teachers represented" if lang == "en" else ("Al menos cinco docentes representados" if lang == "es" else "En az beş öğretmen temsil edildiğinde"), "Open" if lang == "en" else ("Abierto" if lang == "es" else "Açık")])
    if included < 500:
        rows.append(["Medium", "Continue collecting labels" if lang == "en" else ("Continuar recolectando etiquetas" if lang == "es" else "Etiket toplamaya devam et"), "Current validated sample is still limited." if lang == "en" else ("La muestra validada actual todavía es limitada." if lang == "es" else "Mevcut doğrulanmış örneklem hâlâ sınırlı."), owner, "500 additional mature labels" if lang == "en" else ("500 etiquetas maduras adicionales" if lang == "es" else "500 ek olgun etiket"), "Open" if lang == "en" else ("Abierto" if lang == "es" else "Açık")])
    rows.append(["Medium", "Regenerate report after context review" if lang == "en" else ("Regenerar informe tras revisar el contexto" if lang == "es" else "Bağlam gözden geçirildikten sonra raporu yeniden üret"), "Business context improves meeting readiness." if lang == "en" else ("El contexto de negocio mejora la preparación para reuniones." if lang == "es" else "İş bağlamı toplantı hazırlığını güçlendirir."), owner, _context_text(report_context, "next_review_trigger", lang), "Open" if lang == "en" else ("Abierto" if lang == "es" else "Açık")])
    return rows


def _meeting_notes_rows(report_context: dict[str, Any], lang: str) -> list[list[str]]:
    return [
        [_copy(lang, "meta_field"), _copy(lang, "meta_value")],
        ["Proposed decision" if lang == "en" else ("Decisión propuesta" if lang == "es" else "Önerilen karar"), _context_text(report_context, "decision_supported", lang)],
        ["Decision owner" if lang == "en" else ("Responsable de la decisión" if lang == "es" else "Karar sahibi"), _context_text(report_context, "responsible_person_or_team", lang)],
        ["Review date" if lang == "en" else ("Fecha de revisión" if lang == "es" else "İnceleme tarihi"), _context_text(report_context, "next_review_date", lang)],
        ["Additional comments" if lang == "en" else ("Comentarios adicionales" if lang == "es" else "Ek yorumlar"), _context_text(report_context, "meeting_notes", lang)],
    ]


def build_executive_report_docx(run_id: str, language: str) -> dict[str, Any]:
    lang = _lang(language)
    context = _report_base_context(run_id, lang)
    if not context:
        return {"ok": False, "message": t("admin_eic_report_unavailable_no_validated_run", lang=lang)}
    artifacts = _artifact_path_map(run_id)
    if not _required_artifacts_present("executive_docx", context, artifacts):
        return {"ok": False, "message": t("admin_eic_report_unavailable_no_validated_run", lang=lang)}
    detail = context["detail"]
    summary = context["latest_summary"]
    telemetry = context["telemetry"]
    portfolio = context["portfolio"]
    decisions = context["decisions"]
    run_summary = context["run_summary"]
    dataset = run_summary.get("dataset") or context["dataset_summary"]
    charts = _report_chart_assets("executive_docx", run_id, lang, context)

    doc = Document()
    _set_doc_defaults(doc)
    _cover(
        doc,
        _copy(lang, "executive_report_name"),
        t("admin_eic_report_exec_subtitle", lang=lang),
        get_run_status_display(str(detail.get("run_status") or ""), lang=lang),
        [
            _copy(lang, "cover_generated", value=_now_text(lang)),
            _copy(lang, "cover_run", value=run_id),
            _copy(lang, "cover_fingerprint", value=_short_fingerprint(str(detail.get("dataset_fingerprint") or "n/a"))),
        ],
        lang,
    )
    _apply_page_furniture(doc, _copy(lang, "executive_report_name"), run_id, lang)

    _add_heading(doc, t("admin_eic_report_exec_summary_heading", lang=lang), 1)
    _add_paragraph(doc, _copy(lang, "exec_intro"))
    top_decision = summary.get("top_decision") or {}
    summary_rows = [
        [_copy(lang, "what_was_evaluated"), _translated_business_question(detail, lang)],
        [_copy(lang, "what_evidence_showed"), _narrative_robust_winner(detail, run_summary, lang)],
        [_copy(lang, "replacement_question"), _copy(lang, "exec_replace_hold")],
        [_copy(lang, "leadership_action"), get_business_action_display(str(top_decision.get("recommended_action") or detail.get("recommended_business_action") or ""), lang=lang)],
        [_copy(lang, "missing_evidence"), _copy(lang, "exec_missing")],
    ]
    _add_table(doc, [_copy(lang, "meta_field"), _copy(lang, "meta_value")], summary_rows, [2.1, 4.6], caption=_copy(lang, "summary_snapshot"))

    _add_heading(doc, _copy(lang, "exec_findings"), 1)
    kpi_rows = [
        ["Mature labelled observations", str(dataset.get("included_row_count") or detail.get("included_row_count") or 0)],
        ["Teachers represented", str(detail.get("teachers_represented") or 0)],
        ["Students represented", str(detail.get("students_represented") or 0)],
        ["Resources represented", str(detail.get("resources_represented") or 0)],
        ["Validation status", get_run_status_display(str(detail.get("run_status") or ""), lang=lang)],
        ["Evidence strength", get_evidence_display(str(detail.get("evidence_level") or detail.get("evidence_verdict") or ""), lang=lang)],
        ["Robust winner", _bool_text(detail.get("robust_winner"), lang)],
        ["Current recommendation", get_business_action_display(str(detail.get("recommended_business_action") or top_decision.get("recommended_action") or ""), lang=lang)],
    ]
    _add_table(doc, [_copy(lang, "meta_field"), _copy(lang, "meta_value")], kpi_rows, [2.8, 3.9])

    _add_picture_with_caption(doc, charts.get("component_maturity", Path("__missing__")), "Figure 1. Intelligence component maturity distribution.", width=6.3)
    _add_picture_with_caption(doc, charts.get("label_balance", Path("__missing__")), "Figure 2. Label balance for the validated supervised run.", width=5.7)
    _add_picture_with_caption(doc, charts.get("telemetry_surface", Path("__missing__")), "Figure 3. Matched-open coverage by telemetry surface in the recent admin telemetry window.", width=6.4)
    _add_picture_with_caption(doc, charts.get("decision_urgency", Path("__missing__")), "Figure 4. Prioritized decisions by urgency.", width=5.8)
    _append_source_note(doc, lang)

    _add_heading(doc, _copy(lang, "portfolio_summary"), 1)
    _add_table(
        doc,
        [_copy(lang, "system"), _copy(lang, "current_approach"), _copy(lang, "business_use"), _copy(lang, "evidence"), _copy(lang, "recommended_action")],
        _portfolio_rows(portfolio, lang),
        [1.6, 1.2, 2.2, 0.9, 1.0],
        caption=_copy(lang, "caption_portfolio"),
    )

    _add_heading(doc, _copy(lang, "data_feedback_health"), 1)
    telemetry_summary = telemetry.get("summary") or {}
    zero_note = (
        "Zero recent canonical exposures were recorded in the selected telemetry window; matched-open coverage cannot yet be interpreted as model failure."
        if _int(telemetry_summary.get("total_canonical_exposures")) == 0
        else ""
    )
    _add_bullets(
        doc,
        [
            t("admin_eic_report_data_health_line_exposures", lang=lang, value=str(telemetry_summary.get("total_canonical_exposures") or 0)),
            t("admin_eic_report_data_health_line_match", lang=lang, value=_pct(telemetry_summary.get("matched_open_coverage") or 0.0)),
            t("admin_eic_report_data_health_line_unmatched", lang=lang, value=str(telemetry_summary.get("unmatched_opens") or 0)),
            t("admin_eic_report_data_health_line_freshness", lang=lang, value=str(telemetry_summary.get("telemetry_freshness_hours") or "n/a")),
            zero_note,
        ],
    )

    _add_heading(doc, _copy(lang, "business_risks"), 1)
    _add_bullets(doc, _localized_bullets(detail.get("limitations"), lang) or [t("admin_eic_report_no_additional_risks", lang=lang)])

    _add_heading(doc, _copy(lang, "prioritized_actions"), 1)
    _add_table(
        doc,
        [_copy(lang, "priority"), _copy(lang, "component"), _copy(lang, "evidence"), _copy(lang, "impact"), _copy(lang, "recommended_action"), _copy(lang, "owner"), _copy(lang, "review_trigger")],
        _decision_rows(decisions, lang),
        [0.7, 1.3, 2.0, 1.6, 1.1, 0.9, 1.3],
        caption=_copy(lang, "caption_actions"),
    )

    _add_heading(doc, _copy(lang, "roadmap"), 1)
    _add_note_box(
        doc,
        "Next review milestone" if lang == "en" else ("Próximo hito de revisión" if lang == "es" else "Sonraki inceleme kilometre taşı"),
        "Review again after broader teacher coverage is available or when repeated validated runs show stable evidence with cleaner telemetry matching."
        if lang == "en"
        else (
            "Volver a revisar cuando exista una cobertura docente más amplia o cuando ejecuciones validadas repetidas muestren evidencia estable con una telemetría mejor conciliada."
            if lang == "es"
            else "Daha geniş öğretmen kapsamı mevcut olduğunda veya yinelenen doğrulanmış çalışmalar daha temiz telemetri eşleşmesiyle istikrarlı kanıt gösterdiğinde yeniden inceleyin."
        ),
    )

    _add_heading(doc, _copy(lang, "run_metadata"), 1)
    _add_table(doc, [_copy(lang, "meta_field"), _copy(lang, "meta_value")], _metadata_rows(detail, run_summary, context["dataset_summary"], lang)[1:], [2.2, 4.5])

    report_dir = _report_dir(run_id, lang)
    report_dir.mkdir(parents=True, exist_ok=True)
    path = report_dir / _report_filename("executive_docx", run_id)
    doc.save(path)
    return {"ok": True, "path": str(path), "bytes": _verified_docx_bytes(path), "report_type": "executive_docx"}


def build_academic_report_docx(run_id: str, language: str) -> dict[str, Any]:
    lang = _lang(language)
    context = _report_base_context(run_id, lang)
    if not context:
        return {"ok": False, "message": t("admin_eic_report_unavailable_no_validated_run", lang=lang)}
    artifacts = _artifact_path_map(run_id)
    if not _required_artifacts_present("academic_docx", context, artifacts):
        return {"ok": False, "message": t("admin_eic_report_unavailable_no_validated_run", lang=lang)}
    detail = context["detail"]
    academic = context["academic"]
    run_summary = context["run_summary"]
    dataset_summary = context["dataset_summary"]
    evaluation = run_summary.get("evaluation") or {}
    charts = _report_chart_assets("academic_docx", run_id, lang, context)
    model_rows = list((detail.get("model_results") or {}).get("models_compared") or [])
    ai_narrative = _generate_report_ai_narrative(context, report_kind="academic", lang=lang)

    doc = Document()
    _set_doc_defaults(doc)
    _cover(
        doc,
        _copy(lang, "academic_report_name"),
        t("admin_eic_report_academic_subtitle", lang=lang),
        get_run_status_display(str(detail.get("run_status") or ""), lang=lang),
        [
            _copy(lang, "cover_generated", value=_now_text(lang)),
            _copy(lang, "cover_run", value=run_id),
            _copy(lang, "cover_fingerprint", value=_short_fingerprint(str(detail.get("dataset_fingerprint") or "n/a"))),
        ],
        lang,
    )
    _apply_page_furniture(doc, _copy(lang, "academic_report_name"), run_id, lang)

    _add_heading(doc, _copy(lang, "academic_abstract"), 1)
    _add_paragraph(doc, _copy(lang, "academic_conclusion"))

    academic_sections = [
        (_copy(lang, "company_context"), str(academic.get("company_context") or "")),
        (_copy(lang, "problem_statement"), _copy(lang, "exec_intro")),
        (_copy(lang, "solution_statement"), _copy(lang, "exec_replace_hold")),
        (_copy(lang, "smart_objective"), _copy(lang, "exec_missing")),
        (_copy(lang, "supervised_formulation"), str(academic.get("target_definition") or "")),
    ]
    for heading, body in academic_sections:
        _add_heading(doc, heading, 1)
        _add_paragraph(doc, body)

    _add_heading(doc, _copy(lang, "dataset_sources"), 1)
    dataset_rows = [
        [_copy(lang, "cover_fingerprint", value="").split(":")[0], str(academic.get("dataset_fingerprint") or "")],
        [_phrase(lang, "dataset_size"), str(academic.get("dataset_size") or 0)],
        ["Positive class balance" if lang == "en" else ("Balance de clase positiva" if lang == "es" else "Pozitif sınıf dengesi"), _pct(academic.get("class_balance") or 0.0)],
        ["Date coverage" if lang == "en" else ("Cobertura temporal" if lang == "es" else "Tarih kapsamı"), f"{(academic.get('date_range') or {}).get('start') or 'n/a'} {_date_connector(lang)} {(academic.get('date_range') or {}).get('end') or 'n/a'}"],
        ["Data sources" if lang == "en" else ("Fuentes de datos" if lang == "es" else "Veri kaynakları"), ", ".join(academic.get("data_sources") or [])],
        ["Unit of analysis" if lang == "en" else ("Unidad de análisis" if lang == "es" else "Analiz birimi"), str(academic.get("unit_of_analysis") or "")],
    ]
    _add_table(doc, [_copy(lang, "meta_field"), _copy(lang, "meta_value")], dataset_rows, [2.0, 4.7])
    _add_picture_with_caption(doc, charts.get("label_balance", Path("__missing__")), "Figure 1. Stored class distribution for the validated run." if lang == "en" else ("Figura 1. Distribución de clases almacenada para la ejecución validada." if lang == "es" else "Şekil 1. Doğrulanmış çalışma için saklanan sınıf dağılımı."), width=5.6)
    _add_picture_with_caption(doc, charts.get("split_timeline", Path("__missing__")), "Figure 2. Event timeline with the stored chronological cutoff." if lang == "en" else ("Figura 2. Línea temporal de eventos con el corte cronológico almacenado." if lang == "es" else "Şekil 2. Saklanan kronolojik kesim ile olay zaman çizelgesi."), width=6.4)

    _add_heading(doc, _copy(lang, "data_preparation"), 1)
    _add_bullets(
        doc,
        [
            "Inclusion and exclusion rules are taken from the stored dataset summary and the run review." if lang == "en" else ("Las reglas de inclusión y exclusión se toman del resumen de datos y de la revisión almacenada de la ejecución." if lang == "es" else "Dahil etme ve hariç tutma kuralları, saklanan veri kümesi özeti ve çalışma incelemesinden alınır."),
            (f"Included rows: {dataset_summary.get('included_row_count') or detail.get('included_row_count') or 0}." if lang == "en" else (f"Filas incluidas: {dataset_summary.get('included_row_count') or detail.get('included_row_count') or 0}." if lang == "es" else f"Dahil edilen satırlar: {dataset_summary.get('included_row_count') or detail.get('included_row_count') or 0}.")),
            (f"Excluded rows: {dataset_summary.get('excluded_row_count') or 0}, primarily due to open observation windows." if lang == "en" else (f"Filas excluidas: {dataset_summary.get('excluded_row_count') or 0}, principalmente por ventanas de observación abiertas." if lang == "es" else f"Hariç tutulan satırlar: {dataset_summary.get('excluded_row_count') or 0}; bunun başlıca nedeni açık gözlem pencereleridir.")),
            "Right censoring is handled by excluding records whose seven-day observation window had not closed at extraction time." if lang == "en" else ("La censura por la derecha se maneja excluyendo registros cuya ventana de observación de siete días aún no había cerrado al momento de la extracción." if lang == "es" else "Sağ sansürleme, çıkarım anında yedi günlük gözlem penceresi kapanmamış kayıtların dışarıda bırakılmasıyla uygulanır."),
            "Feature engineering is restricted to information available at or before assignment time to avoid leakage." if lang == "en" else ("La ingeniería de variables se restringe a información disponible en o antes del momento de la asignación para evitar fuga de información." if lang == "es" else "Bilgi sızıntısını önlemek için özellik mühendisliği yalnızca atama anında veya öncesinde mevcut olan bilgilerle sınırlandırılır."),
        ],
    )

    _add_heading(doc, _copy(lang, "feature_selection"), 1)
    feature_health = _feature_health_rows(run_summary, lang)
    if feature_health:
        _add_table(
            doc,
            [_phrase(lang, "feature_col"), _phrase(lang, "source_col"), _phrase(lang, "prediction_time_availability_col"), _phrase(lang, "overall_missing_col"), _phrase(lang, "dev_missing_col"), _phrase(lang, "holdout_missing_col"), _phrase(lang, "included_col"), _phrase(lang, "explanation_col")],
            feature_health,
            [0.75, 0.75, 0.9, 0.65, 0.65, 0.75, 0.5, 1.75],
            caption=_copy(lang, "caption_feature_health"),
        )

    _add_heading(doc, _copy(lang, "target_construction"), 1)
    _add_paragraph(doc, "The target is opened_within_7d, derived from teacher_assignments.assigned_at and subsequent opened_at or viewed_at events inside a seven-day window. Negative cases are records whose observation window closed without a qualifying open event." if lang == "en" else ("El objetivo es opened_within_7d, derivado de teacher_assignments.assigned_at y de eventos posteriores opened_at o viewed_at dentro de una ventana de siete días. Los casos negativos son registros cuya ventana de observación cerró sin un evento de apertura válido." if lang == "es" else "Hedef opened_within_7d olup teacher_assignments.assigned_at ve sonraki yedi günlük penceredeki opened_at veya viewed_at olaylarından türetilir. Negatif vakalar, geçerli bir açılma olayı olmadan gözlem penceresi kapanan kayıtlardır."))

    _add_heading(doc, _copy(lang, "methodology"), 1)
    methodology_rows = [
        [_phrase(lang, "chronological_cutoff"), str((academic.get("train_holdout_split") or {}).get("chronological_cutoff") or evaluation.get("cutoff_timestamp") or "n/a")],
        ["Evaluation design" if lang == "en" else ("Diseño de evaluación" if lang == "es" else "Değerlendirme tasarımı"), str(academic.get("evaluation_design") or "")],
        ["Baseline comparator" if lang == "en" else ("Comparador base" if lang == "es" else "Temel karşılaştırıcı"), str(academic.get("baseline") or "DummyClassifier")],
        ["Primary metric leader" if lang == "en" else ("Líder de métrica principal" if lang == "es" else "Ana metrik lideri"), _model_name(str(academic.get("selected_metric_leader") or detail.get("primary_metric_leader") or ""), lang)],
    ]
    _add_table(doc, [_copy(lang, "meta_field"), _copy(lang, "meta_value")], methodology_rows, [2.0, 4.7])

    _add_heading(doc, _copy(lang, "models_evaluated"), 1)
    model_summary_rows = [
        [
            str(row.get("model_name") or ""),
            str(get_model_comparison_value_display("model_kind", row.get("model_kind"), lang=lang) or ""),
            str(get_model_comparison_value_display("status", row.get("status"), lang=lang) or ""),
            str(row.get("overall_interpretation") or ("Stored comparison available." if lang == "en" else ("Comparación almacenada disponible." if lang == "es" else "Saklanan karşılaştırma mevcut."))),
        ]
        for row in model_rows
    ]
    _add_table(doc, [_copy(lang, "model"), _phrase(lang, "kind_col"), _phrase(lang, "status_col"), _copy(lang, "interpretation")], model_summary_rows, [1.5, 0.8, 0.8, 2.8], caption=_copy(lang, "caption_models"))

    _add_heading(doc, _copy(lang, "evaluation_metrics"), 1)
    _add_bullets(
        doc,
        [
            "Primary ranking metrics: ROC AUC and average precision." if lang == "en" else ("Métricas principales de ranking: ROC AUC y average precision." if lang == "es" else "Ana sıralama metrikleri: ROC AUC ve average precision."),
            "Decision metrics: balanced accuracy, precision, recall, specificity, and F1-score." if lang == "en" else ("Métricas de decisión: balanced accuracy, precision, recall, specificity y F1-score." if lang == "es" else "Karar metrikleri: balanced accuracy, precision, recall, specificity ve F1-score."),
            "Calibration diagnostics: Brier score and log loss." if lang == "en" else ("Diagnósticos de calibración: Brier score y log loss." if lang == "es" else "Kalibrasyon tanıları: Brier score ve log loss."),
            "Confidence intervals are taken from the stored model comparison artifact when available." if lang == "en" else ("Los intervalos de confianza se toman del artefacto almacenado de comparación de modelos cuando está disponible." if lang == "es" else "Güven aralıkları mevcut olduğunda saklanan model karşılaştırma artefaktından alınır."),
        ],
    )

    _add_heading(doc, _copy(lang, "results"), 1)
    _add_picture_with_caption(doc, charts.get("model_metrics", Path("__missing__")), "Figure 3. Stored model comparison snapshot using ROC AUC for readability.", width=6.5)
    _add_picture_with_caption(doc, charts.get("roc_curves", Path("__missing__")), "Figure 4. ROC curves from stored holdout predictions.", width=6.2)
    _add_picture_with_caption(doc, charts.get("pr_curves", Path("__missing__")), "Figure 5. Precision-recall curves from stored holdout predictions.", width=6.2)
    _add_picture_with_caption(doc, charts.get("leader_confusion", Path("__missing__")), "Figure 6. Confusion matrix for the primary ROC leader.", width=3.8)
    _add_picture_with_caption(doc, charts.get("calibration_curves", Path("__missing__")), "Figure 7. Calibration curves from stored holdout probabilities.", width=6.2)
    _append_source_note(doc, lang)

    _add_heading(doc, _copy(lang, "comparative_analysis"), 1)
    _add_paragraph(doc, str(ai_narrative.get("analysis_paragraph") or _narrative_robust_winner(detail, run_summary, lang)))
    _add_bullets(
        doc,
        [
            f"Primary ROC AUC leader: {_model_name(str(evaluation.get('primary_metric_leader') or detail.get('primary_metric_leader') or ''), lang)}.",
            f"Best thresholded classifier: {_model_name(str(evaluation.get('best_thresholded_classifier') or ''), lang)}.",
            f"Best precision-recall ranking: {_model_name(str(evaluation.get('best_precision_recall_ranking') or ''), lang)}.",
            f"Calibration leader: {_model_name(str(evaluation.get('calibration_leader') or ''), lang)}.",
        ],
    )

    _add_heading(doc, _copy(lang, "conclusions"), 1)
    _add_note_box(doc, _copy(lang, "section_note"), str(ai_narrative.get("conclusion_paragraph") or _copy(lang, "academic_conclusion")))

    _add_heading(doc, _copy(lang, "implementation"), 1)
    _add_paragraph(doc, str(ai_narrative.get("implementation_paragraph") or ("The present evidence supports continued offline evaluation, telemetry improvement, and internal academic review rather than immediate live replacement of the existing heuristic workflow." if lang == "en" else ("La evidencia actual respalda continuar con evaluación offline, mejora de telemetría y revisión académica interna antes de cualquier sustitución inmediata de la lógica heurística en vivo." if lang == "es" else "Mevcut kanıt, canlı sezgisel iş akışının hemen değiştirilmesi yerine çevrimdışı değerlendirme, telemetri iyileştirmesi ve iç akademik incelemenin sürdürülmesini destekler."))))

    _add_heading(doc, _copy(lang, "limitations"), 1)
    _add_bullets(doc, _localized_bullets(ai_narrative.get("limitations") or academic.get("limitations") or detail.get("limitations") or [], lang))

    _add_heading(doc, _copy(lang, "future_work"), 1)
    _add_bullets(doc, list(ai_narrative.get("future_work") or academic.get("future_improvements") or []))

    _add_heading(doc, _copy(lang, "references"), 1)
    _add_bullets(
        doc,
        [
            "docs/classio_ml_blueprint.md",
            "reports/ml_architecture/assigned_resource_open_7d/.../assigned_resource_open_7d_run_summary.json",
            "reports/ml_architecture/assigned_resource_open_7d/.../assigned_resource_open_7d_integrity_review.md",
            "reports/ml_architecture/assigned_resource_open_7d/.../assigned_resource_open_7d_model_comparison.csv",
        ],
    )

    _add_heading(doc, _copy(lang, "appendix"), 1)
    _add_table(doc, [_copy(lang, "meta_field"), _copy(lang, "meta_value")], _dataset_accounting_rows(detail, run_summary, dataset_summary, lang), [2.2, 4.5])

    report_dir = _report_dir(run_id, lang)
    report_dir.mkdir(parents=True, exist_ok=True)
    path = report_dir / _report_filename("academic_docx", run_id)
    doc.save(path)
    return {"ok": True, "path": str(path), "bytes": _verified_docx_bytes(path), "report_type": "academic_docx"}


def build_technical_report_docx(run_id: str, language: str) -> dict[str, Any]:
    lang = _lang(language)
    context = _report_base_context(run_id, lang)
    if not context:
        return {"ok": False, "message": t("admin_eic_report_unavailable_no_validated_run", lang=lang)}
    missing_context_labels = _report_context_missing_labels(context, lang)
    if missing_context_labels:
        return {
            "ok": False,
            "message": t("report_context_generation_blocked", lang=lang, fields=", ".join(missing_context_labels)),
        }
    artifacts = _artifact_path_map(run_id)
    if not _required_artifacts_present("technical_docx", context, artifacts):
        return {"ok": False, "message": t("admin_eic_report_unavailable_no_validated_run", lang=lang)}
    detail = context["detail"]
    academic = context["academic"]
    run_summary = context["run_summary"]
    dataset_summary = context["dataset_summary"]
    evaluation = run_summary.get("evaluation") or {}
    review = run_summary.get("review") or {}
    report_context = context.get("report_context") or {}
    charts = _report_chart_assets("technical_docx", run_id, lang, context)
    model_rows = list((detail.get("model_results") or {}).get("models_compared") or [])
    manifest_rows, filenames = _artifact_manifest_rows(artifacts, lang)
    ai_narrative = _generate_report_ai_narrative(context, report_kind="technical", lang=lang)

    doc = Document()
    _set_doc_defaults(doc)
    _cover(
        doc,
        _copy(lang, "technical_report_name"),
        t("admin_eic_report_technical_subtitle", lang=lang),
        get_run_status_display(str(detail.get("run_status") or ""), lang=lang),
        [
            _copy(lang, "cover_generated", value=_now_text(lang)),
            _copy(lang, "cover_run", value=run_id),
            _copy(lang, "cover_fingerprint", value=_short_fingerprint(str(detail.get("dataset_fingerprint") or "n/a"))),
        ],
        lang,
    )
    _apply_page_furniture(doc, _copy(lang, "technical_report_name"), run_id, lang)

    _add_heading(doc, "Decision Summary" if lang == "en" else ("Resumen de decisión" if lang == "es" else "Karar özeti"), 1)
    _add_table(doc, [_copy(lang, "meta_field"), _copy(lang, "meta_value")], _decision_summary_rows(context, ai_narrative, lang)[1:], [2.4, 4.3])
    if _clean_text(ai_narrative.get("decision_summary_paragraph")):
        _add_note_box(
            doc,
            _copy(lang, "section_note"),
            str(ai_narrative.get("decision_summary_paragraph") or ""),
        )
    _add_note_box(
        doc,
        "What the evidence says" if lang == "en" else ("Qué dice la evidencia" if lang == "es" else "Kanıt ne söylüyor"),
        str(ai_narrative.get("analysis_paragraph") or _narrative_robust_winner(detail, run_summary, lang)),
    )
    _add_note_box(
        doc,
        "What it does not prove" if lang == "en" else ("Qué no demuestra" if lang == "es" else "Ne kanıtlamıyor"),
        str(
            ai_narrative.get("non_proof_paragraph")
            or _context_text(report_context, "evidence_non_proof", lang, ai_narrative)
            or (_localized_bullets(_unique_nonempty(list(academic.get("limitations") or []) + list(detail.get("limitations") or [])), lang) or [_editable_placeholder(lang)])[0]
        ),
    )
    _add_note_box(
        doc,
        "What Classio should do next" if lang == "en" else ("Qué debe hacer Classio después" if lang == "es" else "Classio'nun sıradaki adımı"),
        str(ai_narrative.get("implementation_paragraph") or ai_narrative.get("proposed_next_action_text") or _context_text(report_context, "recommended_next_action", lang, ai_narrative) or get_business_action_display(str(detail.get("recommended_business_action") or ""), lang=lang) or _editable_placeholder(lang)),
    )

    _add_heading(doc, "Business and educational context" if lang == "en" else ("Contexto de negocio y educativo" if lang == "es" else "İş ve eğitim bağlamı"), 1)
    context_rows = [
        [_copy(lang, "meta_field"), _copy(lang, "meta_value")],
        ["Why this experiment is being evaluated" if lang == "en" else ("Por qué se está evaluando este experimento" if lang == "es" else "Bu deney neden değerlendiriliyor"), _context_text(report_context, "business_problem", lang, ai_narrative)],
        ["Decision supported" if lang == "en" else ("Decisión apoyada" if lang == "es" else "Desteklenen karar"), _context_text(report_context, "decision_supported", lang, ai_narrative)],
        ["Expected value" if lang == "en" else ("Valor esperado" if lang == "es" else "Beklenen değer"), _context_text(report_context, "expected_value", lang, ai_narrative)],
        ["Product impact" if lang == "en" else ("Impacto en producto" if lang == "es" else "Ürün etkisi"), _context_text(report_context, "product_impact", lang, ai_narrative)],
        ["Success definition" if lang == "en" else ("Definición de éxito" if lang == "es" else "Başarı tanımı"), _context_text(report_context, "success_definition", lang, ai_narrative)],
        ["Minimum evidence required" if lang == "en" else ("Evidencia mínima requerida" if lang == "es" else "Gerekli asgari kanıt"), _context_text(report_context, "minimum_evidence_required", lang, ai_narrative)],
        ["Risks" if lang == "en" else ("Riesgos" if lang == "es" else "Riskler"), _context_text(report_context, "risks", lang, ai_narrative)],
    ]
    _add_table(doc, [_copy(lang, "meta_field"), _copy(lang, "meta_value")], context_rows[1:], [2.3, 4.4])

    _add_heading(doc, "Business readiness scorecard" if lang == "en" else ("Scorecard de preparación de negocio" if lang == "es" else "İş hazırlık puan kartı"), 1)
    scorecard_rows = _readiness_scorecard_rows(context, lang)
    _add_table(doc, scorecard_rows[0], scorecard_rows[1:], [1.6, 1.0, 4.1])

    _add_heading(doc, "Action plan" if lang == "en" else ("Plan de acción" if lang == "es" else "Eylem planı"), 1)
    action_rows = _action_plan_rows(context, lang)
    _add_table(doc, action_rows[0], action_rows[1:], [0.8, 1.7, 2.0, 1.2, 1.5, 0.8])

    _add_heading(doc, _copy(lang, "technical_metadata"), 1)
    metadata_rows = _metadata_rows(detail, run_summary, dataset_summary, lang)[1:]
    metadata_rows.extend(
        [
            ["Business context supplied by" if lang == "en" else ("Contexto de negocio aportado por" if lang == "es" else "İş bağlamını sağlayan"), _clean_text(report_context.get("created_by")) or _editable_placeholder(lang)],
            ["Context last updated" if lang == "en" else ("Última actualización del contexto" if lang == "es" else "Bağlamın son güncellemesi"), _clean_text(report_context.get("updated_at")) or _editable_placeholder(lang)],
            ["AI-assisted wording" if lang == "en" else ("Redacción asistida por IA" if lang == "es" else "YZ destekli ifade"), _ai_used_text(ai_narrative, lang)],
            ["Report language" if lang == "en" else ("Idioma del informe" if lang == "es" else "Rapor dili"), lang.upper()],
            ["Report version" if lang == "en" else ("Versión del informe" if lang == "es" else "Rapor sürümü"), str(EXPERIMENT_REPORT_TEMPLATE_VERSION)],
        ]
    )
    _add_table(doc, [_copy(lang, "meta_field"), _copy(lang, "meta_value")], metadata_rows, [2.2, 4.5])

    _add_heading(doc, _copy(lang, "experiment_definition"), 1)
    _add_paragraph(doc, _translated_business_question(detail, lang))
    _add_paragraph(doc, "Prediction target: opened_within_7d, based on assignment-open behavior within a seven-day window." if lang == "en" else ("Objetivo de predicción: opened_within_7d, basado en comportamiento de apertura de asignaciones dentro de una ventana de siete días." if lang == "es" else "Tahmin hedefi: opened_within_7d; yedi günlük pencere içindeki atama-açılma davranışına dayanır."))

    _add_heading(doc, _copy(lang, "source_logic"), 1)
    _add_bullets(
        doc,
        [
            "Source tables: teacher_assignments, teacher_assignment_attempts, practice_sessions, resource_exposures, and resource_exposure_events." if lang == "en" else ("Tablas fuente: teacher_assignments, teacher_assignment_attempts, practice_sessions, resource_exposures y resource_exposure_events." if lang == "es" else "Kaynak tablolar: teacher_assignments, teacher_assignment_attempts, practice_sessions, resource_exposures ve resource_exposure_events."),
            "Extraction logic is frozen in the validated run artifacts and summarized in the dataset summary JSON." if lang == "en" else ("La lógica de extracción queda congelada en los artefactos validados de la ejecución y se resume en el JSON de resumen del conjunto de datos." if lang == "es" else "Çıkarım mantığı doğrulanmış çalışma artefaktlarında dondurulmuştur ve veri kümesi özeti JSON dosyasında özetlenir."),
            "All feature construction is constrained to information known at or before assignment time." if lang == "en" else ("Toda la construcción de variables se limita a información conocida en o antes del momento de la asignación." if lang == "es" else "Tüm özellik oluşturma yalnızca atama anında veya öncesinde bilinen bilgilerle sınırlandırılmıştır."),
        ],
    )

    _add_heading(doc, _copy(lang, "dataset_accounting"), 1)
    _add_table(doc, [_copy(lang, "meta_field"), _copy(lang, "meta_value")], _dataset_accounting_rows(detail, run_summary, dataset_summary, lang), [2.1, 4.6])
    _add_picture_with_caption(doc, charts.get("label_balance", Path("__missing__")), "Figure 1. Stored target/class distribution." if lang == "en" else ("Figura 1. Distribución almacenada de objetivo y clases." if lang == "es" else "Şekil 1. Saklanan hedef/sınıf dağılımı."), width=5.6)
    _add_picture_with_caption(doc, charts.get("split_timeline", Path("__missing__")), "Figure 2. Assignment timeline and chronological split cutoff." if lang == "en" else ("Figura 2. Línea temporal de asignaciones y corte cronológico." if lang == "es" else "Şekil 2. Atama zaman çizelgesi ve kronolojik ayrım kesimi."), width=6.4)

    _add_heading(doc, _copy(lang, "label_reconciliation"), 1)
    reconciliation = (review.get("label_reconciliation") or {})
    recon_rows = [
        ["Final verdict" if lang == "en" else ("Veredicto final" if lang == "es" else "Nihai hüküm"), get_run_status_display(str(review.get("final_verdict") or detail.get("run_status") or ""), lang=lang)],
        ["Exact row-level reconciliation" if lang == "en" else ("Conciliación exacta a nivel de fila" if lang == "es" else "Satır düzeyinde tam uzlaştırma"), _bool_text(reconciliation.get("exact_row_level_reconciliation_available"), lang)],
        ["Legacy audit reconciliation applicable" if lang == "en" else ("Conciliación de auditoría legacy aplicable" if lang == "es" else "Eski denetim uzlaştırması uygulanabilir"), _bool_text(reconciliation.get("legacy_audit_reconciliation_applicable"), lang)],
        ["Audit counts" if lang == "en" else ("Conteos de auditoría" if lang == "es" else "Denetim sayıları"), json.dumps(reconciliation.get("audit_documented_counts") or {}, ensure_ascii=True)],
        ["Phase 3 counts" if lang == "en" else ("Conteos fase 3" if lang == "es" else "Faz 3 sayıları"), json.dumps(reconciliation.get("phase3_counts") or {}, ensure_ascii=True)],
    ]
    _add_table(doc, [_copy(lang, "meta_field"), _copy(lang, "meta_value")], recon_rows, [2.4, 4.3])
    _add_bullets(
        doc,
        [
            _translate_known_evidence_text(reconciliation.get("likely_difference_explanation") or "", lang),
            _translate_known_evidence_text(detail.get("validation_notes") or "", lang),
        ],
    )

    _add_heading(doc, _copy(lang, "leakage_controls"), 1)
    _add_bullets(
        doc,
        [
            "Chronological holdout keeps future assignments out of development data." if lang == "en" else ("El holdout cronológico mantiene asignaciones futuras fuera de los datos de desarrollo." if lang == "es" else "Kronolojik holdout, gelecekteki atamaları geliştirme verisinin dışında tutar."),
            "Feature timestamps are checked against assignment time to enforce past-only availability." if lang == "en" else ("Las marcas de tiempo de las variables se comparan con la fecha de asignación para garantizar disponibilidad solo del pasado." if lang == "es" else "Özellik zaman damgaları, yalnızca geçmişte mevcut olan bilgilerin kullanıldığını doğrulamak için atama zamanı ile karşılaştırılır."),
            "Feature missingness explanations in the stored review confirm that sparsity comes from history availability rather than identifier mismatch." if lang == "en" else ("Las explicaciones de ausencia de variables en la revisión almacenada indican que la dispersión proviene de la disponibilidad histórica y no de un desajuste de identificadores." if lang == "es" else "Saklanan incelemedeki özellik eksikliği açıklamaları, seyrekliğin kimlik uyuşmazlığından değil geçmiş veri kullanılabilirliğinden kaynaklandığını doğrular."),
        ],
    )

    _add_heading(doc, _copy(lang, "feature_health"), 1)
    _add_picture_with_caption(doc, charts.get("feature_missingness", Path("__missing__")), "Figure 3. Highest overall missingness rates from the stored feature audit." if lang == "en" else ("Figura 3. Tasas de ausencia más altas según la auditoría almacenada de variables." if lang == "es" else "Şekil 3. Saklanan özellik denetimine göre en yüksek genel eksiklik oranları."), width=6.4)
    feature_health_rows = _feature_health_rows(run_summary, lang)
    if feature_health_rows:
        _add_table(
            doc,
            [_phrase(lang, "feature_col"), _phrase(lang, "source_col"), _phrase(lang, "availability_col"), _phrase(lang, "overall_missing_col"), _phrase(lang, "dev_missing_col"), _phrase(lang, "holdout_missing_col"), _phrase(lang, "included_col"), _phrase(lang, "explanation_col")],
            feature_health_rows,
            [0.75, 0.75, 0.85, 0.65, 0.65, 0.75, 0.5, 2.2],
            caption=_copy(lang, "caption_feature_health"),
        )

    _add_heading(doc, _copy(lang, "preprocessing"), 1)
    _add_bullets(
        doc,
        [
            "Missing values are imputed only when the stored evaluation pipeline marked the feature as retained." if lang == "en" else ("Los valores faltantes se imputan solo cuando la canalización almacenada de evaluación marcó la variable como retenida." if lang == "es" else "Eksik değerler yalnızca saklanan değerlendirme hattı ilgili özelliği korunmuş olarak işaretlediğinde doldurulur."),
            "Categorical features are encoded within the stored pipeline configuration; no report-time transformation modifies model outcomes." if lang == "en" else ("Las variables categóricas se codifican dentro de la configuración almacenada de la canalización; ninguna transformación en tiempo de informe modifica los resultados del modelo." if lang == "es" else "Kategorik özellikler saklanan hat yapılandırması içinde kodlanır; rapor üretimi sırasında yapılan hiçbir dönüşüm model sonuçlarını değiştirmez."),
            "Fully missing development-slice features are automatically excluded before fitting, as documented in the integrity review." if lang == "en" else ("Las variables totalmente ausentes en el tramo de desarrollo se excluyen automáticamente antes del ajuste, tal como se documenta en la revisión de integridad." if lang == "es" else "Geliştirme bölümünde tamamen eksik olan özellikler, bütünlük incelemesinde belgelendiği gibi eğitimden önce otomatik olarak hariç tutulur."),
        ],
    )

    _add_heading(doc, _copy(lang, "cv_section"), 1)
    _add_bullets(
        doc,
        [
            (f"Development rows: {evaluation.get('development_count') or 'n/a'} with {evaluation.get('development_positive_count') or 'n/a'} positives." if lang == "en" else (f"Filas de desarrollo: {evaluation.get('development_count') or 'n/a'} con {evaluation.get('development_positive_count') or 'n/a'} positivas." if lang == "es" else f"Geliştirme satırları: {evaluation.get('development_count') or 'n/a'}; pozitif sayısı {evaluation.get('development_positive_count') or 'n/a'}.")),
            (f"Holdout rows: {evaluation.get('holdout_count') or 'n/a'} with {evaluation.get('holdout_positive_count') or 'n/a'} positives." if lang == "en" else (f"Filas holdout: {evaluation.get('holdout_count') or 'n/a'} con {evaluation.get('holdout_positive_count') or 'n/a'} positivas." if lang == "es" else f"Holdout satırları: {evaluation.get('holdout_count') or 'n/a'}; pozitif sayısı {evaluation.get('holdout_positive_count') or 'n/a'}.")),
            "Cross-validation fold counts and mean metrics are read directly from the stored model comparison artifact." if lang == "en" else ("Los conteos de folds y las métricas medias de validación cruzada se leen directamente del artefacto almacenado de comparación de modelos." if lang == "es" else "Çapraz doğrulama fold sayıları ve ortalama metrikler doğrudan saklanan model karşılaştırma artefaktından okunur."),
        ],
    )

    _add_heading(doc, _copy(lang, "model_configuration"), 1)
    config_rows = [
        [
            _model_name(str(row.get("model_name") or ""), lang),
            get_model_result_status_display(str(row.get("status") or ""), lang=lang),
            _display_jsonish(row.get("parameters_json") or {}, max_chars=200),
        ]
        for row in model_rows
    ]
    _add_table(doc, [_copy(lang, "model"), _phrase(lang, "status_col"), _phrase(lang, "selected_hyperparameters")], config_rows, [1.8, 0.9, 4.0])

    _add_heading(doc, _copy(lang, "baseline_results"), 1)
    _add_picture_with_caption(doc, charts.get("model_metrics", Path("__missing__")), "Figure 4. Model metric comparison snapshot using stored ROC AUC values." if lang == "en" else ("Figura 4. Comparación de métricas de modelos con valores ROC AUC almacenados." if lang == "es" else "Şekil 4. Saklanan ROC AUC değerleriyle model metrik karşılaştırması."), width=6.6)
    _add_picture_with_caption(doc, charts.get("runtime", Path("__missing__")), "Figure 5. Training-runtime comparison from stored model outputs." if lang == "en" else ("Figura 5. Comparación de tiempos de entrenamiento a partir de salidas almacenadas del modelo." if lang == "es" else "Şekil 5. Saklanan model çıktılarından eğitim süresi karşılaştırması."), width=6.6)

    _add_heading(doc, _copy(lang, "threshold_analysis"), 1)
    threshold_rows = [
        ["Primary ROC leader" if lang == "en" else ("Líder principal por ROC" if lang == "es" else "Birincil ROC lideri"), _model_name(str(evaluation.get("primary_metric_leader") or detail.get("primary_metric_leader") or ""), lang)],
        ["Best thresholded classifier" if lang == "en" else ("Mejor clasificador por umbral" if lang == "es" else "En iyi eşiklenmiş sınıflandırıcı"), _model_name(str(evaluation.get("best_thresholded_classifier") or ""), lang)],
        ["Calibration leader" if lang == "en" else ("Líder de calibración" if lang == "es" else "Kalibrasyon lideri"), _model_name(str(evaluation.get("calibration_leader") or ""), lang)],
        ["Best precision-recall ranking" if lang == "en" else ("Mejor ranking precision-recall" if lang == "es" else "En iyi precision-recall sıralaması"), _model_name(str(evaluation.get("best_precision_recall_ranking") or ""), lang)],
    ]
    _add_table(doc, [_copy(lang, "meta_field"), _copy(lang, "meta_value")], threshold_rows, [2.4, 4.3])

    confusion_targets = [
        ("baseline_confusion", "Figure 6. Confusion matrix for the stored baseline." if lang == "en" else ("Figura 6. Matriz de confusión para la línea base almacenada." if lang == "es" else "Şekil 6. Saklanan temel model için karmaşıklık matrisi."), _clean_text("DummyClassifier")),
        ("leader_confusion", "Figure 7. Confusion matrix for the primary ROC leader." if lang == "en" else ("Figura 7. Matriz de confusión para el líder principal por ROC." if lang == "es" else "Şekil 7. Birincil ROC lideri için karmaşıklık matrisi."), _clean_text(evaluation.get("primary_metric_leader") or detail.get("primary_metric_leader") or "")),
        ("threshold_confusion", "Figure 8. Confusion matrix for the best thresholded classifier." if lang == "en" else ("Figura 8. Matriz de confusión para el mejor clasificador por umbral." if lang == "es" else "Şekil 8. En iyi eşiklenmiş sınıflandırıcı için karmaşıklık matrisi."), _clean_text(evaluation.get("best_thresholded_classifier") or "")),
    ]
    seen_confusion_models: set[str] = set()
    for chart_key, caption, model_name in confusion_targets:
        if not model_name or model_name in seen_confusion_models:
            continue
        seen_confusion_models.add(model_name)
        _add_picture_with_caption(doc, charts.get(chart_key, Path("__missing__")), caption, width=3.7)

    _add_heading(doc, _copy(lang, "roc_analysis"), 1)
    _add_picture_with_caption(doc, charts.get("roc_curves", Path("__missing__")), "Figure 9. ROC curves for score-capable stored models." if lang == "en" else ("Figura 9. Curvas ROC para los modelos almacenados capaces de producir puntuaciones." if lang == "es" else "Şekil 9. Skor üretebilen saklanan modeller için ROC eğrileri."), width=6.4)
    _add_paragraph(doc, _narrative_robust_winner(detail, run_summary, lang))

    _add_heading(doc, _copy(lang, "pr_analysis"), 1)
    _add_picture_with_caption(doc, charts.get("pr_curves", Path("__missing__")), "Figure 10. Precision-recall curves for the same stored holdout predictions." if lang == "en" else ("Figura 10. Curvas precision-recall para las mismas predicciones holdout almacenadas." if lang == "es" else "Şekil 10. Aynı saklanan holdout tahminleri için precision-recall eğrileri."), width=6.4)

    _add_heading(doc, _copy(lang, "calibration_analysis"), 1)
    _add_picture_with_caption(doc, charts.get("calibration_curves", Path("__missing__")), "Figure 11. Calibration curves derived from stored probabilities." if lang == "en" else ("Figura 11. Curvas de calibración derivadas de probabilidades almacenadas." if lang == "es" else "Şekil 11. Saklanan olasılıklardan türetilen kalibrasyon eğrileri."), width=6.4)
    _add_picture_with_caption(doc, charts.get("probability_distribution", Path("__missing__")), "Figure 12. Predicted-probability distributions for the stored leader, threshold leader, and baseline." if lang == "en" else ("Figura 12. Distribuciones de probabilidad predicha para el líder almacenado, el líder por umbral y la línea base." if lang == "es" else "Şekil 12. Saklanan lider, eşik lideri ve temel model için tahmin edilen olasılık dağılımları."), width=6.4)

    _add_heading(doc, _copy(lang, "error_analysis"), 1)
    _add_bullets(
        doc,
        [
            "False positives and false negatives can be inspected only through the stored anonymized holdout predictions; no raw user identifiers are exposed." if lang == "en" else ("Los falsos positivos y falsos negativos solo pueden inspeccionarse mediante las predicciones holdout anonimizadas almacenadas; no se exponen identificadores crudos de usuario." if lang == "es" else "Yanlış pozitifler ve yanlış negatifler yalnızca saklanan anonimleştirilmiş holdout tahminleri üzerinden incelenebilir; ham kullanıcı kimlikleri açığa çıkarılmaz."),
            "Subgroup sample sizes are too small for strong technical claims by topic, student-history bucket, or resource type, so subgroup interpretation should remain descriptive only." if lang == "en" else ("Los tamaños muestrales por subgrupo son demasiado pequeños para afirmaciones técnicas sólidas por tema, historial del estudiante o tipo de recurso, por lo que la interpretación debe seguir siendo descriptiva." if lang == "es" else "Alt grup örneklem boyutları; konu, öğrenci geçmişi kovası veya kaynak türü bazında güçlü teknik iddialar için çok küçüktür; bu nedenle yorum yalnızca betimleyici kalmalıdır."),
            "The confusion matrices show that threshold choice changes the balance between misses and false alarms more than it changes the overall conclusion." if lang == "en" else ("Las matrices de confusión muestran que la elección del umbral cambia más el equilibrio entre omisiones y falsas alarmas que la conclusión general." if lang == "es" else "Karmaşıklık matrisleri, eşik seçiminin genel sonucu değiştirmekten çok kaçırmalar ile yanlış alarmlar arasındaki dengeyi değiştirdiğini gösterir."),
        ],
    )

    _add_heading(doc, _copy(lang, "uncertainty"), 1)
    _add_bullets(
        doc,
        [
            "Confidence intervals are read from the stored model comparison CSV." if lang == "en" else ("Los intervalos de confianza se leen del CSV almacenado de comparación de modelos." if lang == "es" else "Güven aralıkları saklanan model karşılaştırma CSV dosyasından okunur."),
            "Small holdout size and single-teacher coverage mean all uncertainty intervals should be interpreted conservatively." if lang == "en" else ("El tamaño pequeño del holdout y la cobertura de un solo docente implican que todos los intervalos de incertidumbre deben interpretarse con prudencia." if lang == "es" else "Küçük holdout boyutu ve tek öğretmen kapsamı, tüm belirsizlik aralıklarının temkinli yorumlanması gerektiği anlamına gelir."),
            (f"Stored overall verdict: {get_run_status_display(str(review.get('final_verdict') or detail.get('run_status') or ''), lang=lang)}." if lang == "en" else (f"Veredicto general almacenado: {get_run_status_display(str(review.get('final_verdict') or detail.get('run_status') or ''), lang=lang)}." if lang == "es" else f"Saklanan genel hüküm: {get_run_status_display(str(review.get('final_verdict') or detail.get('run_status') or ''), lang=lang)}.")),
        ],
    )

    _add_heading(doc, _copy(lang, "runtime"), 1)
    _add_paragraph(doc, "Training and inference durations are reported directly from the validated model comparison artifact and shown in Figure 5." if lang == "en" else ("Las duraciones de entrenamiento e inferencia se reportan directamente desde el artefacto validado de comparación de modelos y se muestran en la Figura 5." if lang == "es" else "Eğitim ve çıkarım süreleri doğrudan doğrulanmış model karşılaştırma artefaktından raporlanır ve Şekil 5'te gösterilir."))

    _add_heading(doc, _copy(lang, "integrity"), 1)
    _add_bullets(
        doc,
        [
            _copy(lang, "technical_conclusion"),
            _translate_known_evidence_text(detail.get("validation_notes") or "", lang),
            "All intended models executed successfully in the validated run." if lang == "en" else ("Todos los modelos previstos se ejecutaron con éxito en la ejecución validada." if lang == "es" else "Planlanan tüm modeller doğrulanmış çalışmada başarıyla çalıştı."),
            "Artifact consistency checks confirmed shared run ID and data fingerprint across stored outputs." if lang == "en" else ("Las comprobaciones de consistencia de artefactos confirmaron ID de ejecución y huella de datos compartidos entre las salidas almacenadas." if lang == "es" else "Artefakt tutarlılık kontrolleri, saklanan çıktılar arasında ortak çalışma kimliği ve veri parmak izi olduğunu doğruladı."),
        ]
        + [_translate_known_evidence_text(item, lang) for item in list((review.get("label_reconciliation") or {}).get("limitations") or [])],
    )

    _add_heading(doc, _copy(lang, "conclusions"), 1)
    _add_note_box(
        doc,
        _copy(lang, "section_note"),
        str(ai_narrative.get("conclusion_paragraph") or _narrative_robust_winner(detail, run_summary, lang)),
    )

    _add_heading(doc, _copy(lang, "reproducibility"), 1)
    _add_bullets(
        doc,
        [
            (f"Run ID: {run_id}" if lang == "en" else (f"ID de ejecución: {run_id}" if lang == "es" else f"Çalışma kimliği: {run_id}")),
            (f"Dataset fingerprint: {detail.get('dataset_fingerprint') or 'n/a'}" if lang == "en" else (f"Huella del conjunto de datos: {detail.get('dataset_fingerprint') or 'n/a'}" if lang == "es" else f"Veri kümesi parmak izi: {detail.get('dataset_fingerprint') or 'n/a'}")),
            (f"Feature schema version: {_humanize_identifier(dataset_summary.get('feature_schema_version') or run_summary.get('feature_schema_version') or 'n/a')}" if lang == "en" else (f"Versión del esquema de variables: {_humanize_identifier(dataset_summary.get('feature_schema_version') or run_summary.get('feature_schema_version') or 'n/a')}" if lang == "es" else f"Özellik şeması sürümü: {_humanize_identifier(dataset_summary.get('feature_schema_version') or run_summary.get('feature_schema_version') or 'n/a')}")),
            "This report was generated from stored artifacts only; it does not recompute training metrics independently." if lang == "en" else ("Este informe se generó solo a partir de artefactos almacenados; no recompone ni recalcula métricas de entrenamiento de forma independiente." if lang == "es" else "Bu rapor yalnızca saklanan artefaktlardan üretildi; eğitim metriklerini bağımsız olarak yeniden hesaplamaz."),
        ],
    )

    _add_heading(doc, _copy(lang, "implementation"), 1)
    _add_note_box(
        doc,
        _copy(lang, "section_note"),
        str(
            ai_narrative.get("implementation_paragraph")
            or (
                "Current evidence supports continued supervised experimentation, but not a production change. Any operational change should wait for broader teacher coverage, stronger telemetry consistency, and another validated run."
                if lang == "en"
                else (
                    "La evidencia actual respalda seguir experimentando con modelos supervisados, pero no un cambio en producción. Cualquier cambio operativo debería esperar a una cobertura docente más amplia, mayor consistencia de telemetría y otra ejecución validada."
                    if lang == "es"
                    else "Mevcut kanıtlar denetimli deneylerin sürdürülmesini destekliyor, ancak üretimde bir değişikliği desteklemiyor. Operasyonel bir değişiklik için daha geniş öğretmen kapsamı, daha tutarlı telemetri ve yeni bir doğrulanmış çalışma beklenmelidir."
                )
            )
        ),
    )

    _add_heading(doc, _copy(lang, "limitations"), 1)
    limitation_rows = _unique_nonempty(
        list(ai_narrative.get("limitations") or [])
        + list(academic.get("limitations") or [])
        + list(detail.get("limitations") or [])
        + list((review.get("label_reconciliation") or {}).get("limitations") or [])
    )
    _add_bullets(doc, _localized_bullets(limitation_rows, lang) or [t("admin_eic_report_no_additional_risks", lang=lang)])

    _add_heading(doc, _copy(lang, "appendix"), 1)
    _add_heading(doc, _copy(lang, "artifact_manifest"), 2)
    _add_table(
        doc,
        [_copy(lang, "artifact_name"), _copy(lang, "artifact_purpose"), _copy(lang, "artifact_format"), _copy(lang, "artifact_checksum"), _copy(lang, "artifact_availability")],
        manifest_rows,
        [1.5, 2.5, 0.7, 1.0, 1.0],
        caption=_copy(lang, "caption_manifest"),
    )
    for file_name in filenames:
        _add_paragraph(doc, f"{_copy(lang, 'artifact_filename')}: {file_name}", color=MUTED)

    _add_heading(doc, "Meeting Notes and Decision Record" if lang == "en" else ("Notas de reunión y registro de decisión" if lang == "es" else "Toplantı notları ve karar kaydı"), 2)
    _add_table(doc, [_copy(lang, "meta_field"), _copy(lang, "meta_value")], _meeting_notes_rows(report_context, lang)[1:], [2.2, 4.5])

    if model_rows:
        _add_landscape_section(doc)
        _apply_page_furniture(doc, _copy(lang, "technical_report_name"), run_id, lang)
        _add_heading(doc, _phrase(lang, "full_model_comparison"), 2)
        for title, headers, rows, widths, numeric_cols in _model_metric_groups(model_rows, lang):
            _add_heading(doc, title, 3)
            _add_table(doc, headers, rows, widths, numeric_cols=numeric_cols, caption=_copy(lang, "caption_models"))

    report_dir = _report_dir(run_id, lang)
    report_dir.mkdir(parents=True, exist_ok=True)
    path = report_dir / _report_filename("technical_docx", run_id)
    doc.save(path)
    return {
        "ok": True,
        "path": str(path),
        "bytes": _verified_docx_bytes(path),
        "report_type": "technical_docx",
        "generation_mode": "ai" if bool(ai_narrative.get("_ai_used")) else "template",
        "provider": str(ai_narrative.get("_provider") or ""),
    }


def _is_unsupervised_report_context(context: dict[str, Any]) -> bool:
    detail = context.get("detail") or {}
    return get_experiment_paradigm(str(detail.get("experiment_id") or "")) == EXPERIMENT_PARADIGM_UNSUPERVISED


def _read_csv_records(path: Path, *, limit: int = 12) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    try:
        frame = pd.read_csv(path)
    except Exception:
        return []
    if frame.empty:
        return []
    return frame.head(max(1, int(limit))).fillna("").to_dict("records")


def _unsupervised_metric_rows(summary: dict[str, Any], detail: dict[str, Any], lang: str) -> list[list[str]]:
    evaluation = summary.get("evaluation") or {}
    best = evaluation.get("best_model") or {}
    dataset = summary.get("dataset") or {}
    embedding = evaluation.get("embedding_manifest") or {}
    return [
        ["Best model" if lang == "en" else ("Mejor modelo" if lang == "es" else "En iyi model"), str(evaluation.get("winner") or best.get("model_name") or detail.get("primary_metric_leader") or "—")],
        ["Primary metric" if lang == "en" else ("Métrica principal" if lang == "es" else "Ana metrik"), _AFFINITY_METRIC_LABELS.get(lang, {}).get(str(evaluation.get("primary_metric") or detail.get("primary_metric") or ""), str(evaluation.get("primary_metric") or detail.get("primary_metric") or "—"))],
        ["Silhouette score" if lang == "en" else ("Silhouette score" if lang == "es" else "Silhouette skoru"), _display_scalar(best.get("silhouette_score"))],
        ["Calinski-Harabasz" if lang == "en" else ("Calinski-Harabasz" if lang == "es" else "Calinski-Harabasz"), _display_scalar(best.get("calinski_harabasz"))],
        ["Davies-Bouldin" if lang == "en" else ("Davies-Bouldin" if lang == "es" else "Davies-Bouldin"), _display_scalar(best.get("davies_bouldin"))],
        ["Clusters" if lang == "en" else ("Clústeres" if lang == "es" else "Kümeler"), _display_scalar(best.get("cluster_count"))],
        ["Noise ratio" if lang == "en" else ("Proporción de ruido" if lang == "es" else "Gürültü oranı"), _display_scalar(best.get("noise_ratio"))],
        ["Cross-subject contamination" if lang == "en" else ("Contaminación entre materias" if lang == "es" else "Dersler arası karışma"), _display_scalar(best.get("cross_subject_contamination_rate"))],
        ["Cross-language contamination" if lang == "en" else ("Contaminación entre idiomas" if lang == "es" else "Diller arası karışma"), _display_scalar(best.get("cross_language_contamination_rate"))],
        ["Selection score" if lang == "en" else ("Selection score" if lang == "es" else "Seçim skoru"), _display_scalar(best.get("selection_score"))],
        ["Resources included" if lang == "en" else ("Recursos incluidos" if lang == "es" else "Dahil edilen kaynaklar"), _display_scalar(dataset.get("resource_count") or detail.get("resources_represented") or detail.get("included_row_count"))],
        ["Curricular anchors" if lang == "en" else ("Anclas curriculares" if lang == "es" else "Müfredat çıpaları"), _display_scalar(dataset.get("curricular_anchor_count"))],
        ["Candidate resources" if lang == "en" else ("Recursos candidatos" if lang == "es" else "Aday kaynaklar"), _display_scalar(dataset.get("candidate_resource_count"))],
        ["Resource types" if lang == "en" else ("Tipos de recurso" if lang == "es" else "Kaynak türleri"), _display_scalar(dataset.get("resource_type_count"))],
        ["Vector method" if lang == "en" else ("Método vectorial" if lang == "es" else "Vektör yöntemi"), _display_scalar(embedding.get("embedding_method"))],
        ["Vector dimensions" if lang == "en" else ("Dimensiones vectoriales" if lang == "es" else "Vektör boyutları"), _display_scalar(embedding.get("vector_dimensions"))],
    ]


def _unsupervised_normalization_rows(summary: dict[str, Any], lang: str) -> list[list[str]]:
    evaluation = summary.get("evaluation") or {}
    dataset = summary.get("dataset") or {}
    normalization = evaluation.get("normalization_methodology") or {}
    completeness = dataset.get("completeness") or {}
    aliases = normalization.get("known_aliases") or {}
    if lang == "es":
        return [
            ["Unidad de análisis", _localize_known_text(evaluation.get("target_definition"), lang)],
            ["Roles", "program_topic = ancla curricular; worksheet, exam, video, lesson_plan y program = recursos candidatos."],
            ["Limpieza textual", "Los campos se convierten a texto, se compactan espacios múltiples y se eliminan espacios al inicio/final."],
            ["Normalización categórica", "Se aplica Unicode NFKC, casefold, normalización de separadores / y |, y alias conocidos para materia e idioma."],
            ["Alias usados", json.dumps(aliases, ensure_ascii=False, sort_keys=True)],
            ["Valores faltantes", "Los valores vacíos se auditan y no cuentan en denominadores de contaminación entre materias o idiomas."],
            ["Completitud", json.dumps(completeness, ensure_ascii=False, sort_keys=True)],
            ["Vectorización", "TF-IDF con unigramas y bigramas; SVD truncado cuando el tamaño lo permite; normalización L2 antes de similaridad coseno."],
            ["Artefactos auditables", ", ".join(normalization.get("audit_artifacts") or [])],
        ]
    if lang == "tr":
        return [
            ["Analiz birimi", _localize_known_text(evaluation.get("target_definition"), lang)],
            ["Roller", "program_topic = müfredat çıpası; worksheet, exam, video, lesson_plan ve program = aday kaynaklar."],
            ["Metin temizliği", "Alanlar metne çevrilir, fazla boşluklar tek boşluğa indirilir ve baştaki/sondaki boşluklar silinir."],
            ["Kategori normalizasyonu", "Unicode NFKC, casefold, / ve | ayırıcı normalizasyonu ve bilinen ders/dil alias eşleşmeleri uygulanır."],
            ["Kullanılan aliaslar", json.dumps(aliases, ensure_ascii=False, sort_keys=True)],
            ["Eksik değerler", "Boş değerler denetlenir ve ders/dil karışma paydalarına dahil edilmez."],
            ["Tamlık", json.dumps(completeness, ensure_ascii=False, sort_keys=True)],
            ["Vektörleştirme", "Unigram/bigram TF-IDF; uygun olduğunda Truncated SVD; kosinüs benzerliğinden önce L2 normalizasyonu."],
            ["Denetlenebilir artefaktlar", ", ".join(normalization.get("audit_artifacts") or [])],
        ]
    return [
        ["Analysis unit", _localize_known_text(evaluation.get("target_definition"), lang)],
        ["Roles", "program_topic = curricular anchor; worksheet, exam, video, lesson_plan, and program = candidate resources."],
        ["Text cleaning", "Fields are converted to text, repeated whitespace is collapsed, and leading/trailing whitespace is stripped."],
        ["Category normalization", "Unicode NFKC, casefolding, / and | separator normalization, and known subject/language aliases are applied."],
        ["Aliases used", json.dumps(aliases, ensure_ascii=False, sort_keys=True)],
        ["Missing values", "Empty values are audited and excluded from subject/language contamination denominators."],
        ["Completeness", json.dumps(completeness, ensure_ascii=False, sort_keys=True)],
        ["Vectorization", "TF-IDF with unigrams and bigrams; Truncated SVD when shape allows; L2 normalization before cosine similarity."],
        ["Auditable artifacts", ", ".join(normalization.get("audit_artifacts") or [])],
    ]


def _unsupervised_scoring_rows(summary: dict[str, Any], lang: str) -> list[list[str]]:
    evaluation = summary.get("evaluation") or {}
    scoring = evaluation.get("selection_scoring_methodology") or {}
    winner = evaluation.get("winner_explanation") or {}
    components = scoring.get("component_definitions") or {}
    rationale = scoring.get("weight_rationale") or []
    if lang == "es":
        return [
            ["Fórmula", _display_scalar(scoring.get("formula"))],
            ["Orden de selección", " -> ".join(_localize_known_text(item, lang) for item in scoring.get("sort_order") or [])],
            ["Ganador", f"{winner.get('winner') or 'n/a'}; selection_score={_display_scalar(winner.get('winner_selection_score'))}"],
            ["Líder por Silhouette", f"{winner.get('metric_leader') or 'n/a'}; selection_score={_display_scalar(winner.get('metric_leader_selection_score'))}"],
            ["Definición de componentes", _localized_json_map(components, lang, _AFFINITY_COMPONENT_LABELS)],
            ["Justificación de pesos", _localized_join(rationale, lang)],
            ["Por qué no Silhouette solamente", _localize_known_text(scoring.get("why_not_silhouette_only"), lang)],
        ]
    if lang == "tr":
        return [
            ["Formül", _display_scalar(scoring.get("formula"))],
            ["Seçim sırası", " -> ".join(_localize_known_text(item, lang) for item in scoring.get("sort_order") or [])],
            ["Kazanan", f"{winner.get('winner') or 'n/a'}; selection_score={_display_scalar(winner.get('winner_selection_score'))}"],
            ["Silhouette lideri", f"{winner.get('metric_leader') or 'n/a'}; selection_score={_display_scalar(winner.get('metric_leader_selection_score'))}"],
            ["Bileşen tanımları", _localized_json_map(components, lang, _AFFINITY_COMPONENT_LABELS)],
            ["Ağırlık gerekçesi", _localized_join(rationale, lang)],
            ["Neden yalnız Silhouette değil", _localize_known_text(scoring.get("why_not_silhouette_only"), lang)],
        ]
    return [
        ["Formula", _display_scalar(scoring.get("formula"))],
        ["Selection order", " -> ".join(scoring.get("sort_order") or [])],
        ["Winner", f"{winner.get('winner') or 'n/a'}; selection_score={_display_scalar(winner.get('winner_selection_score'))}"],
        ["Silhouette leader", f"{winner.get('metric_leader') or 'n/a'}; selection_score={_display_scalar(winner.get('metric_leader_selection_score'))}"],
        ["Component definitions", json.dumps(components, ensure_ascii=False, sort_keys=True)],
        ["Weight rationale", " ".join(str(item) for item in rationale)],
        ["Why not Silhouette only", _display_scalar(scoring.get("why_not_silhouette_only"))],
    ]


def _unsupervised_python_methodology_rows(summary: dict[str, Any], lang: str) -> list[list[str]]:
    evaluation = summary.get("evaluation") or {}
    methodology = evaluation.get("python_model_development_methodology") or {}
    data_sources = methodology.get("data_sources") or {}
    if lang == "es":
        return [
            ["Lenguaje", _display_scalar(methodology.get("language"))],
            ["Librerías principales", ", ".join(methodology.get("core_libraries") or [])],
            ["Datos utilizados", "; ".join(f"{_AFFINITY_DATA_SOURCE_LABELS.get(lang, {}).get(str(key), key)}: {_localize_known_text(value, lang)}" for key, value in data_sources.items())],
            ["Reglas de inclusión/exclusión", _localized_join(methodology.get("inclusion_rules"), lang)],
            ["Construcción de perfiles", _localized_join(methodology.get("profile_construction"), lang)],
            ["Pipeline del modelo", _localized_join(methodology.get("model_pipeline"), lang)],
            ["Reproducibilidad", _localized_join(methodology.get("reproducibility_controls"), lang)],
        ]
    if lang == "tr":
        return [
            ["Dil", _display_scalar(methodology.get("language"))],
            ["Ana kütüphaneler", ", ".join(methodology.get("core_libraries") or [])],
            ["Kullanılan veriler", "; ".join(f"{_AFFINITY_DATA_SOURCE_LABELS.get(lang, {}).get(str(key), key)}: {_localize_known_text(value, lang)}" for key, value in data_sources.items())],
            ["Dahil etme/dışlama kuralları", _localized_join(methodology.get("inclusion_rules"), lang)],
            ["Profil oluşturma", _localized_join(methodology.get("profile_construction"), lang)],
            ["Model hattı", _localized_join(methodology.get("model_pipeline"), lang)],
            ["Yeniden üretilebilirlik", _localized_join(methodology.get("reproducibility_controls"), lang)],
        ]
    return [
        ["Language", _display_scalar(methodology.get("language"))],
        ["Core libraries", ", ".join(methodology.get("core_libraries") or [])],
        ["Data used", "; ".join(f"{key}: {value}" for key, value in data_sources.items())],
        ["Inclusion/exclusion rules", " ".join(str(item) for item in methodology.get("inclusion_rules") or [])],
        ["Profile construction", " ".join(str(item) for item in methodology.get("profile_construction") or [])],
        ["Model pipeline", " ".join(str(item) for item in methodology.get("model_pipeline") or [])],
        ["Reproducibility", " ".join(str(item) for item in methodology.get("reproducibility_controls") or [])],
    ]


def build_unsupervised_experiment_report_docx(run_id: str, language: str, *, report_type: str = "experiment_docx") -> dict[str, Any]:
    lang = _lang(language)
    context = _report_base_context(run_id, lang)
    if not context:
        return {"ok": False, "message": t("admin_eic_report_unavailable_no_validated_run", lang=lang)}
    if not _is_unsupervised_report_context(context):
        return build_experiment_report_docx(run_id, lang)
    missing_context_labels = _report_context_missing_labels(context, lang)
    if missing_context_labels:
        return {
            "ok": False,
            "message": t("report_context_generation_blocked", lang=lang, fields=", ".join(missing_context_labels)),
        }
    artifacts = _artifact_path_map(run_id)
    if not artifacts.get("run_summary_json") or not artifacts["run_summary_json"].exists():
        return {"ok": False, "message": t("admin_eic_report_unavailable_no_validated_run", lang=lang)}

    detail = context["detail"]
    run_summary = context["run_summary"]
    dataset_summary = context["dataset_summary"]
    report_context = context.get("report_context") or {}
    academic = context.get("academic") or {}
    manifest_rows, filenames = _artifact_manifest_rows(artifacts, lang)
    model_rows = _read_csv_records(artifacts.get("model_comparison_csv", Path("__missing__")), limit=12)
    neighbor_rows = _read_csv_records(artifacts.get("holdout_predictions_csv", Path("__missing__")), limit=10)
    charts = build_report_charts("unsupervised_docx", run_id, lang, context, artifacts, _report_dir(run_id, lang) / "assets_unsupervised_docx")
    ai_narrative = _generate_report_ai_narrative(context, report_kind="unsupervised", lang=lang)

    doc = Document()
    _set_doc_defaults(doc)
    title = "Unsupervised Experiment Report" if lang == "en" else ("Informe de experimento no supervisado" if lang == "es" else "Denetimsiz deney raporu")
    _cover(
        doc,
        title,
        _translated_business_question(detail, lang),
        get_run_status_display(str(detail.get("run_status") or ""), lang=lang),
        [
            _copy(lang, "cover_generated", value=_now_text(lang)),
            _copy(lang, "cover_run", value=run_id),
            _copy(lang, "cover_fingerprint", value=_short_fingerprint(str(detail.get("dataset_fingerprint") or "n/a"))),
        ],
        lang,
    )
    _apply_page_furniture(doc, title, run_id, lang)

    _add_heading(doc, "Planteamiento de la solución" if lang == "es" else ("Solution framing" if lang == "en" else "Çözüm çerçevesi"), 1)
    _add_paragraph(doc, _translated_business_question(detail, lang))
    _add_table(
        doc,
        [_copy(lang, "meta_field"), _copy(lang, "meta_value")],
        [
            ["Business problem" if lang == "en" else ("Problema de negocio" if lang == "es" else "İş problemi"), _context_text(report_context, "business_problem", lang, ai_narrative)],
            ["Decision supported" if lang == "en" else ("Decisión apoyada" if lang == "es" else "Desteklenen karar"), _context_text(report_context, "decision_supported", lang, ai_narrative)],
            ["SMART objective" if lang == "en" else ("Objetivo SMART" if lang == "es" else "SMART hedef"), _context_text(report_context, "success_definition", lang, ai_narrative)],
            ["Expected value" if lang == "en" else ("Valor esperado" if lang == "es" else "Beklenen değer"), _context_text(report_context, "expected_value", lang, ai_narrative)],
        ],
        [2.1, 4.6],
    )

    _add_heading(doc, "Desarrollo del modelo" if lang == "es" else ("Model development" if lang == "en" else "Model geliştirme"), 1)
    _add_paragraph(
        doc,
        "Este experimento usa perfiles canónicos de recursos y algoritmos no supervisados para descubrir clústeres y vecinos semánticos sin construir etiquetas positivas/negativas."
        if lang == "es"
        else ("This experiment uses canonical resource profiles and unsupervised algorithms to discover clusters and semantic neighbors without positive/negative labels." if lang == "en" else "Bu deney, pozitif/negatif etiketler oluşturmadan kümeleri ve anlamsal komşuları keşfetmek için kanonik kaynak profilleri ve denetimsiz algoritmalar kullanır."),
    )
    _add_heading(doc, "Datos utilizados y desarrollo en Python" if lang == "es" else ("Data used and Python model development" if lang == "en" else "Kullanılan veri ve Python model geliştirme"), 2)
    _add_table(
        doc,
        [_copy(lang, "meta_field"), _copy(lang, "meta_value")],
        _unsupervised_python_methodology_rows(run_summary, lang),
        [2.1, 4.6],
    )
    _add_heading(doc, "Normalización y preparación de datos" if lang == "es" else ("Data normalization and preparation" if lang == "en" else "Veri normalizasyonu ve hazırlığı"), 2)
    _add_table(
        doc,
        [_copy(lang, "meta_field"), _copy(lang, "meta_value")],
        _unsupervised_normalization_rows(run_summary, lang),
        [2.1, 4.6],
    )
    _add_heading(doc, "Fórmula de scoring del ganador" if lang == "es" else ("Winner scoring formula" if lang == "en" else "Kazanan skor formülü"), 2)
    _add_table(
        doc,
        [_copy(lang, "meta_field"), _copy(lang, "meta_value")],
        _unsupervised_scoring_rows(run_summary, lang),
        [2.1, 4.6],
    )
    _add_picture_with_caption(
        doc,
        charts.get("unsupervised_model_quality", Path("__missing__")),
        "Figura 1. Ranking de modelos por Silhouette Score; barras más altas indican mayor coherencia interna de los clústeres." if lang == "es" else ("Figure 1. Model ranking by Silhouette Score; higher bars indicate stronger internal cluster coherence." if lang == "en" else "Şekil 1. Silhouette skoruna göre model sıralaması."),
        width=6.5,
    )
    _add_picture_with_caption(
        doc,
        charts.get("unsupervised_quality_tradeoff", Path("__missing__")),
        "Figura 2. Trade-off entre calidad y contaminación: el ganador debe combinar alto Silhouette con baja mezcla entre materias." if lang == "es" else ("Figure 2. Quality/contamination trade-off: the winner should combine high Silhouette with low cross-subject mixing." if lang == "en" else "Şekil 2. Kalite/karışma dengesi."),
        width=6.5,
    )
    _add_table(doc, [_copy(lang, "meta_field"), _copy(lang, "meta_value")], _unsupervised_metric_rows(run_summary, detail, lang), [2.4, 4.3])

    if model_rows:
        _add_heading(doc, "Model comparison" if lang == "en" else ("Comparación de modelos" if lang == "es" else "Model karşılaştırması"), 2)
        comparison_rows = [
            [
                str(row.get("model_name") or ""),
                _display_scalar(row.get("silhouette_score")),
                _display_scalar(row.get("cluster_count")),
                _display_scalar(row.get("noise_ratio")),
                _display_scalar(row.get("cross_subject_contamination_rate")),
                _display_scalar(row.get("selection_score")),
            ]
            for row in model_rows
        ]
        _add_table(
            doc,
            ["Model" if lang == "en" else ("Modelo" if lang == "es" else "Model"), "Silhouette", "Clusters" if lang == "en" else ("Clústeres" if lang == "es" else "Kümeler"), "Noise" if lang == "en" else ("Ruido" if lang == "es" else "Gürültü"), "Subject contamination" if lang == "en" else ("Contaminación por materia" if lang == "es" else "Ders karışması"), "Selection score"],
            comparison_rows,
            [2.0, 1.0, 0.8, 0.8, 1.2, 0.9],
            numeric_cols={1, 2, 3, 4, 5},
        )

    _add_heading(doc, "Visualización de clústeres" if lang == "es" else ("Cluster visualization" if lang == "en" else "Küme görselleştirme"), 2)
    _add_picture_with_caption(
        doc,
        charts.get("unsupervised_cluster_projection", Path("__missing__")),
        "Figura 3. Proyección 2D reproducible de los perfiles de recursos; las X marcan centroides aproximados de los clústeres principales y los puntos grises son ruido/otros clústeres." if lang == "es" else ("Figure 3. Reproducible 2D projection of resource profiles; X markers show approximate centroids for major clusters and gray points represent noise/other clusters." if lang == "en" else "Şekil 3. Kaynak profillerinin 2B projeksiyonu."),
        width=6.5,
    )
    _add_picture_with_caption(
        doc,
        charts.get("unsupervised_cluster_sizes", Path("__missing__")),
        "Figura 4. Tamaños de los clústeres principales del modelo ganador; ayuda a detectar clústeres dominantes o demasiado fragmentados." if lang == "es" else ("Figure 4. Main cluster sizes for the winning model; this helps detect dominant or overly fragmented clusters." if lang == "en" else "Şekil 4. Kazanan modelin ana küme boyutları."),
        width=6.2,
    )

    if neighbor_rows:
        _add_heading(doc, "Semantic neighbor audit" if lang == "en" else ("Auditoría de vecinos semánticos" if lang == "es" else "Anlamsal komşu denetimi"), 2)
        _add_table(
            doc,
            ["Source" if lang == "en" else ("Origen" if lang == "es" else "Kaynak"), "Neighbor" if lang == "en" else ("Vecino" if lang == "es" else "Komşu"), "Similarity" if lang == "en" else ("Similitud" if lang == "es" else "Benzerlik"), "Same subject" if lang == "en" else ("Misma materia" if lang == "es" else "Aynı ders")],
            [
                [
                    str(row.get("source_title") or row.get("source_resource_key") or "")[:90],
                    str(row.get("target_title") or row.get("target_resource_key") or "")[:90],
                    _display_scalar(row.get("similarity_score")),
                    _bool_text(row.get("same_subject"), lang),
                ]
                for row in neighbor_rows
            ],
            [2.4, 2.4, 0.9, 0.9],
            numeric_cols={2},
        )

    _add_heading(doc, "Conclusiones" if lang == "es" else ("Conclusions" if lang == "en" else "Sonuçlar"), 1)
    _add_note_box(
        doc,
        _copy(lang, "section_note"),
        str(
            ai_narrative.get("conclusion_paragraph")
            or "El resultado es exploratorio y apto para análisis de shadow testing; no debe activar cambios automáticos en producción sin revisión humana de vecinos/clústeres."
            if lang == "es"
            else (ai_narrative.get("conclusion_paragraph") or "The result is exploratory and suitable for shadow-test analysis; it should not trigger automatic production changes without human review of neighbors/clusters.")
        ),
    )
    limitations = _unique_nonempty(
        list(ai_narrative.get("limitations") or [])
        + list(academic.get("limitations") or [])
        + list((run_summary.get("evaluation") or {}).get("limitations") or [])
        + list(((run_summary.get("review") or {}).get("label_reconciliation") or {}).get("limitations") or [])
    )
    _add_heading(doc, _copy(lang, "limitations"), 2)
    _add_bullets(doc, [_localize_known_text(item, lang) for item in limitations] or [t("admin_eic_report_no_additional_risks", lang=lang)])

    _add_heading(doc, _copy(lang, "appendix"), 1)
    _add_heading(doc, _copy(lang, "artifact_manifest"), 2)
    _add_table(
        doc,
        [_copy(lang, "artifact_name"), _copy(lang, "artifact_purpose"), _copy(lang, "artifact_format"), _copy(lang, "artifact_checksum"), _copy(lang, "artifact_availability")],
        manifest_rows,
        [1.5, 2.5, 0.7, 1.0, 1.0],
        caption=_copy(lang, "caption_manifest"),
    )
    for file_name in filenames:
        _add_paragraph(doc, f"{_copy(lang, 'artifact_filename')}: {file_name}", color=MUTED)

    report_dir = _report_dir(run_id, lang)
    report_dir.mkdir(parents=True, exist_ok=True)
    safe_report_type = report_type if report_type in REPORT_TYPES else "experiment_docx"
    path = report_dir / _report_filename(safe_report_type, run_id)
    doc.save(path)
    return {
        "ok": True,
        "path": str(path),
        "bytes": _verified_docx_bytes(path),
        "report_type": safe_report_type,
        "generation_mode": "ai" if bool(ai_narrative.get("_ai_used")) else "template",
        "provider": str(ai_narrative.get("_provider") or ""),
    }


def build_experiment_report_docx(run_id: str, language: str) -> dict[str, Any]:
    lang = _lang(language)
    context = _report_base_context(run_id, lang)
    if context and _is_unsupervised_report_context(context):
        return build_unsupervised_experiment_report_docx(run_id, lang, report_type="experiment_docx")
    technical = build_technical_report_docx(run_id, lang)
    if not technical.get("ok"):
        return technical
    report_dir = _report_dir(run_id, lang)
    report_dir.mkdir(parents=True, exist_ok=True)
    source_path = Path(str(technical.get("path") or ""))
    target_path = report_dir / _report_filename("experiment_docx", run_id)
    target_path.write_bytes(_verified_docx_bytes(source_path))
    return {
        "ok": True,
        "path": str(target_path),
        "bytes": _verified_docx_bytes(target_path),
        "report_type": "experiment_docx",
        "generation_mode": str(technical.get("generation_mode") or "template"),
        "provider": str(technical.get("provider") or ""),
    }


def get_or_create_validated_report(run_id: str, report_type: str, language: str, *, force_regenerate: bool = False) -> dict[str, Any]:
    safe_type = _clean_text(report_type)
    safe_run_id = _clean_text(run_id)
    lang = _lang(language)
    if safe_type not in REPORT_TYPES:
        return {"ok": False, "message": t("admin_eic_report_generation_failed", lang=lang)}
    path = _report_dir(safe_run_id, lang) / _report_filename(safe_type, safe_run_id)
    if path.exists() and not force_regenerate:
        return {"ok": True, "path": str(path), "bytes": _verified_docx_bytes(path), "report_type": safe_type}
    context = _report_base_context(safe_run_id, lang)
    if context and _is_unsupervised_report_context(context):
        return build_unsupervised_experiment_report_docx(safe_run_id, lang, report_type=safe_type)
    builders = {
        "experiment_docx": build_experiment_report_docx,
        "executive_docx": build_executive_report_docx,
        "academic_docx": build_academic_report_docx,
        "technical_docx": build_technical_report_docx,
    }
    return builders[safe_type](safe_run_id, lang)


def list_available_eic_reports(run_id: str, user_capabilities: set[str] | None = None, *, language: str = "en") -> list[dict[str, Any]]:
    safe_run_id = _clean_text(run_id)
    lang = _lang(language)
    capabilities = user_capabilities or set()
    detail = _validated_run_detail(safe_run_id)
    has_validated_run = bool(detail)
    paradigm = get_experiment_paradigm(str(detail.get("experiment_id") or "")) if detail else "unknown"
    description = (
        "Informe Word del experimento no supervisado con métricas de clustering, vecinos semánticos y conclusiones."
        if lang == "es" and paradigm == EXPERIMENT_PARADIGM_UNSUPERVISED
        else (
            "Word report for the unsupervised experiment with clustering metrics, semantic neighbors, and conclusions."
            if paradigm == EXPERIMENT_PARADIGM_UNSUPERVISED
            else t("admin_eic_report_experiment_subtitle", lang=lang)
        )
    )
    base_rows = [
        {
            "report_type": "experiment_docx",
            "title": get_report_type_display("experiment_docx", lang=lang),
            "description": description,
            "status": "available" if has_validated_run else "no_validated_run",
            "restricted": False,
        }
    ]
    for row in base_rows:
        if row["status"] == "available":
            path = _report_dir(safe_run_id, lang) / _report_filename(str(row["report_type"]), safe_run_id)
            row["path"] = str(path) if path.exists() else ""
            row["download_ready"] = path.exists()
            row["modified_at"] = _format_file_timestamp(path, lang) if path.exists() else ""
            try:
                row["modified_epoch"] = int(path.stat().st_mtime) if path.exists() else 0
            except Exception:
                row["modified_epoch"] = 0
        else:
            row["path"] = ""
            row["download_ready"] = False
            row["modified_at"] = ""
            row["modified_epoch"] = 0
    return base_rows
