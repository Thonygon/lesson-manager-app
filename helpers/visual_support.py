import base64
import html
import json
import os
import re
from io import BytesIO
from urllib.request import Request, urlopen

import streamlit as st
from PIL import Image

from core.i18n import t
from helpers.ai_retry import run_with_ai_retries


# ── Hard rule appended to every image-generation prompt ──────────────
_NO_TEXT_RULE = (
    "\n\nCRITICAL — absolutely no text in the image: "
    "Do NOT render any letters, words, numbers, labels, captions, titles, "
    "headings, speech bubbles, or typographic elements of any kind. "
    "The image must contain ONLY illustrations/drawings. "
    "Any form of written language inside the image is strictly forbidden. "
    "Exception: text is allowed ONLY when the prompt explicitly requests "
    "a graph, chart, diagram axis labels, or mathematical notation."
)

_LOWER_PRIMARY_STAGES = {
    "pre_primary",
    "early_primary",
    "lower_primary",
    "primary_lower",
}

_STAGE_TIER_MAP = {
    "pre_primary": "lower_primary",
    "early_primary": "lower_primary",
    "lower_primary": "lower_primary",
    "primary_lower": "lower_primary",
    "upper_primary": "upper_primary",
    "primary_upper": "upper_primary",
    "lower_secondary": "secondary",
    "secondary_lower": "secondary",
    "upper_secondary": "secondary",
    "secondary_upper": "secondary",
    "adult": "adult",
    "professional": "adult",
}


def _stage_tier(stage: str) -> str:
    return _STAGE_TIER_MAP.get(_stage_key(stage), "secondary")


def _stage_label(stage: str) -> str:
    return {
        "lower_primary": "lower-primary",
        "upper_primary": "upper-primary",
        "secondary": "secondary-level",
        "adult": "adult-level",
    }.get(_stage_tier(stage), "educational")


def _visual_style(stage: str) -> str:
    return {
        "lower_primary": "simple educational illustration, warm and easy to recognize",
        "upper_primary": "clear educational illustration, age-appropriate and engaging",
        "secondary": "clean, realistic educational illustration",
        "adult": "professional, clear educational illustration",
    }.get(_stage_tier(stage), "clean educational illustration")


def _visual_mood(stage: str) -> str:
    return {
        "lower_primary": "bright, friendly, classroom-safe",
        "upper_primary": "clear, engaging, classroom-safe",
        "secondary": "clear, informative, print-friendly",
        "adult": "clear, professional, print-friendly",
    }.get(_stage_tier(stage), "clear, informative")


_ABSTRACT_VISUAL_TERMS = {
    "adjective",
    "adjectives",
    "autobiography",
    "biography",
    "conclusion",
    "definition",
    "definitions",
    "description",
    "descriptions",
    "fiction",
    "frequency",
    "genre",
    "genres",
    "grammar",
    "hypothesis",
    "meaning",
    "meanings",
    "mystery",
    "noun",
    "nouns",
    "poetry",
    "process",
    "science fiction",
    "sentence",
    "sentences",
    "story genre",
    "story genres",
    "synonym",
    "synonyms",
    "terminology",
    "verb",
    "verbs",
    "vocabulary",
    "writing genre",
    "writing genres",
}

_FUNCTION_WORDS = {
    "a", "an", "and", "as", "at", "by", "for", "from", "in", "is", "it", "its", "of", "on", "or", "the", "to", "with",
}

_PUNCTUATION_RE = re.compile(r"[.!?;:]")
_WORD_RE = re.compile(r"[A-Za-zÀ-ÿ0-9'/-]+")


def _clean_text(value) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _subject_key(subject: str) -> str:
    return _clean_text(subject).lower().replace("&", "and").replace(" ", "_")


def _stage_key(stage: str) -> str:
    return _clean_text(stage).lower().replace(" ", "_")


def _words(text: str) -> list[str]:
    return _WORD_RE.findall(_clean_text(text).lower())


def _is_lower_primary(stage: str) -> bool:
    return _stage_key(stage) in _LOWER_PRIMARY_STAGES


def _get_openrouter_api_key() -> str:
    try:
        key = str(st.secrets.get("OPENROUTER_API_KEY", "")).strip()
        if key:
            return key
    except Exception:
        pass
    return str(os.getenv("OPENROUTER_API_KEY", "")).strip()


def _get_gemini_api_key() -> str:
    try:
        key = str(st.secrets.get("GEMINI_API_KEY", "")).strip()
        if key:
            return key
    except Exception:
        pass
    return str(os.getenv("GEMINI_API_KEY", "")).strip()


def _get_openai_api_key() -> str:
    try:
        key = str(st.secrets.get("OPENAI_API_KEY", "")).strip()
        if key:
            return key
    except Exception:
        pass
    return str(os.getenv("OPENAI_API_KEY", "")).strip()


def _get_ai_provider() -> str:
    import helpers.lesson_planner as lp
    return lp.get_ai_provider()


def _get_image_provider_order() -> list[str]:
    import helpers.lesson_planner as lp
    order = lp.get_ai_provider_order()
    available = []
    for item in order:
        if item == "openrouter" and _get_openrouter_api_key():
            available.append(item)
        elif item == "gemini" and _get_gemini_api_key():
            available.append(item)
        elif item == "openai" and _get_openai_api_key():
            available.append(item)
    return available


def _provider_chain_label() -> str:
    order = _get_image_provider_order()
    return " -> ".join(order) if order else "none"


def _get_openrouter_image_model() -> str:
    custom = ""
    try:
        custom = str(st.secrets.get("OPENROUTER_IMAGE_MODEL", "")).strip()
    except Exception:
        custom = ""
    if not custom:
        custom = str(os.getenv("OPENROUTER_IMAGE_MODEL", "")).strip()
    return custom or "google/gemini-3.1-flash-image-preview"


def _get_gemini_image_model() -> str:
    custom = ""
    try:
        custom = str(st.secrets.get("GEMINI_IMAGE_MODEL", "")).strip()
    except Exception:
        custom = ""
    if not custom:
        custom = str(os.getenv("GEMINI_IMAGE_MODEL", "")).strip()
    return custom or "gemini-3.1-flash-image-preview"


@st.cache_resource(show_spinner=False)
def _get_gemini_client(api_key: str):
    from google import genai
    return genai.Client(api_key=api_key)


@st.cache_resource(show_spinner=False)
def _get_openai_client(api_key: str):
    from openai import OpenAI
    return OpenAI(api_key=api_key)


def _image_generation_enabled() -> bool:
    return bool(_get_image_provider_order())


def _is_short_label(text: str) -> bool:
    text = _clean_text(text)
    if not text or _PUNCTUATION_RE.search(text):
        return False
    return 1 <= len(_words(text)) <= 3


def _looks_sentence_like(text: str) -> bool:
    text = _clean_text(text)
    return bool(_PUNCTUATION_RE.search(text)) or len(_words(text)) >= 5


def _is_concrete_term(text: str) -> bool:
    text = _clean_text(text)
    if not _is_short_label(text):
        return False
    lower = text.lower()
    if lower in _ABSTRACT_VISUAL_TERMS:
        return False
    if any(token in _FUNCTION_WORDS for token in _words(text)) and len(_words(text)) > 1:
        return False
    return True


def _unique_terms(values) -> list[str]:
    out = []
    seen = set()
    for value in values or []:
        cleaned = _clean_text(value)
        if not cleaned:
            continue
        key = cleaned.casefold()
        if key in seen:
            continue
        seen.add(key)
        out.append(cleaned)
    return out


def _grid_layout_hint(n: int) -> str:
    """Return a composition/framing instruction that matches the actual item count."""
    if n <= 3:
        return f"{n} pictures in a single horizontal row"
    if n == 4:
        return "2x2 grid of separate picture cards"
    if n <= 6:
        return f"2-row grid: {(n + 1) // 2} pictures on top, {n // 2} on the bottom row"
    if n <= 9:
        return f"3-row grid with {(n + 2) // 3} columns, {n} pictures total"
    return f"grid layout with {n} separate picture cards arranged in rows of 4"


def _extract_matching_terms(pairs) -> list[str]:
    left_terms = _unique_terms([item.get("left", "") for item in pairs or [] if isinstance(item, dict)])
    right_terms = _unique_terms([item.get("right", "") for item in pairs or [] if isinstance(item, dict)])

    left_visual = [term for term in left_terms if _is_concrete_term(term)]
    right_visual = [term for term in right_terms if _is_concrete_term(term)]

    left_sentence_side = any(_looks_sentence_like(term) for term in left_terms)
    right_sentence_side = any(_looks_sentence_like(term) for term in right_terms)

    if len(left_visual) >= 3 and right_sentence_side:
        return left_visual[:12]
    if len(right_visual) >= 3 and left_sentence_side:
        return right_visual[:12]
    if len(left_visual) >= 3 and len(left_visual) == len(left_terms):
        return left_visual[:12]
    if len(right_visual) >= 3 and len(right_visual) == len(right_terms):
        return right_visual[:12]
    return []


def _extract_vocab_terms(payload: dict) -> list[str]:
    bank = _unique_terms(payload.get("vocabulary_bank", []))
    if bank:
        return [term for term in bank if _is_concrete_term(term)][:12]

    questions = payload.get("questions", []) or []
    terms = []
    for item in questions:
        if isinstance(item, dict):
            for key in ("word", "left", "right", "text", "stem"):
                if _is_concrete_term(item.get(key, "")):
                    terms.append(item.get(key, ""))
        elif _is_concrete_term(item):
            terms.append(item)
    return _unique_terms(terms)[:12]


def _observable_statements(statements) -> list[str]:
    out = []
    for statement in statements or []:
        text = _clean_text(statement)
        if not text:
            continue
        if len(_words(text)) > 10:
            return []
        out.append(text)
    return out[:5]


def _passage_excerpt(text: str, limit: int = 420) -> str:
    cleaned = _clean_text(text)
    if len(cleaned) <= limit:
        return cleaned
    clipped = cleaned[:limit].rsplit(" ", 1)[0].strip()
    return f"{clipped}..."


def _worksheet_prompt(ws: dict, subject: str, learner_stage: str, topic: str) -> tuple[str, str, str] | None:
    ws_type = _clean_text(ws.get("worksheet_type"))
    subject_label = _clean_text(subject or ws.get("subject", "")).replace("_", " ")
    topic_label = _clean_text(topic or ws.get("topic", "") or ws.get("title", ""))
    tier = _stage_tier(learner_stage)
    stage_lbl = _stage_label(learner_stage)
    style = _visual_style(learner_stage)
    mood = _visual_mood(learner_stage)

    if ws_type == "matching":
        # Picture matching boards help lower/upper primary identify concrete items;
        # secondary only benefits when there are enough highly-concrete terms; adults skip.
        if tier == "adult":
            return None
        min_terms = 4 if tier == "secondary" else 3
        terms = _extract_matching_terms(ws.get("matching_pairs", []))
        if len(terms) < min_terms:
            return None
        grid_layout = _grid_layout_hint(len(terms))
        prompt = (
            "Use case: illustration-story\n"
            "Asset type: worksheet support image\n"
            f"Primary request: Create a picture board with exactly {len(terms)} pictures for a {stage_lbl} {subject_label} matching activity.\n"
            f"Subject: {subject_label or 'general studies'}\n"
            f"Scene/backdrop: clean white worksheet background with separate picture tiles\n"
            f"Subject details: show one clear illustration for each of these items: {', '.join(terms)}.\n"
            f"Style/medium: {style}\n"
            f"Composition/framing: {grid_layout}\n"
            f"Lighting/mood: {mood}\n"
            "Text (verbatim): none\n"
            f"Constraints: exactly {len(terms)} pictures, no words, no letters, no labels, no decorative filler, each picture must be distinct and usable for matching\n"
            "Avoid: abstract icons, posters, banners, title cards, watermark"
        )
        return ("image_based_matching", topic_label or "Matching", prompt)

    if ws_type in {"word_search_vocab"}:
        # Picture vocabulary boards are most useful for primary stages only.
        if tier not in ("lower_primary", "upper_primary"):
            return None
        min_terms = 4 if tier == "lower_primary" else 5
        terms = _extract_vocab_terms(ws)
        if len(terms) < min_terms:
            return None
        grid_layout = _grid_layout_hint(len(terms))
        prompt = (
            "Use case: illustration-story\n"
            "Asset type: worksheet vocabulary support image\n"
            f"Primary request: Create a {stage_lbl} picture vocabulary board with exactly {len(terms)} pictures for these words: {', '.join(terms)}.\n"
            f"Subject: {subject_label or 'general studies'}\n"
            "Scene/backdrop: clean worksheet background with evenly spaced picture cards\n"
            f"Style/medium: {style}\n"
            f"Composition/framing: {grid_layout}\n"
            f"Lighting/mood: {mood}\n"
            "Text (verbatim): none\n"
            f"Constraints: exactly {len(terms)} pictures, no text, no labels, no decorative filler, every item must be easy to identify\n"
            "Avoid: collage chaos, title banners, watermark"
        )
        return ("picture_vocabulary", topic_label or "Vocabulary", prompt)

    if ws_type == "true_false":
        # Observation-scene support helps primary students; older students can reason without it.
        if tier not in ("lower_primary", "upper_primary"):
            return None
        if _clean_text(ws.get("source_text")):
            return None
        statements = _observable_statements(ws.get("true_false_statements", []))
        if len(statements) < 3:
            return None
        prompt = (
            "Use case: illustration-story\n"
            "Asset type: worksheet true/false support image\n"
            f"Primary request: Create one classroom-safe {stage_lbl} scene that supports observation-based true/false statements.\n"
            f"Subject: {subject_label or 'general studies'}\n"
            f"Scene/backdrop: The picture should plausibly support statements like: {' | '.join(statements)}\n"
            f"Style/medium: {style}\n"
            "Composition/framing: one coherent scene with visible details students can observe\n"
            f"Lighting/mood: {mood}\n"
            "Text (verbatim): none\n"
            "Constraints: no text, no labels, no unnecessary extra objects, the visible details must be easy to verify\n"
            "Avoid: decorative posters, abstract symbols, watermark"
        )
        return ("scene_truth_support", topic_label or "True or False", prompt)

    if ws_type == "reading_comprehension":
        # Scene illustrations aid comprehension at all stages; require richer passages for older learners.
        min_passage = {"lower_primary": 80, "upper_primary": 120, "secondary": 200, "adult": 300}.get(tier, 200)
        passage = _clean_text(ws.get("reading_passage"))
        if len(passage) < min_passage:
            return None
        reading_style = "warm educational story illustration" if tier in ("lower_primary", "upper_primary") else "clear educational illustration suited to the passage content"
        prompt = (
            "Use case: illustration-story\n"
            "Asset type: worksheet reading support image\n"
            f"Primary request: Create a {stage_lbl} reading support illustration for this passage.\n"
            f"Subject: {subject_label or 'language learning'}\n"
            f"Scene/backdrop: {_passage_excerpt(passage)}\n"
            f"Style/medium: {reading_style}\n"
            "Composition/framing: one coherent scene that reflects the passage without giving away answers\n"
            f"Lighting/mood: {mood}\n"
            "Text (verbatim): none\n"
            "Constraints: match the passage details, no text, no labels, no answer spoilers beyond the scene itself\n"
            "Avoid: decorative filler, watermark"
        )
        return ("reading_scene_support", topic_label or "Reading", prompt)

    if ws_type == "vocabulary":
        # Picture vocabulary boards are most useful for primary stages only.
        if tier not in ("lower_primary", "upper_primary"):
            return None
        min_terms = 4 if tier == "lower_primary" else 5
        terms = _extract_vocab_terms(ws)
        if len(terms) < min_terms:
            return None
        grid_layout = _grid_layout_hint(len(terms))
        prompt = (
            "Use case: illustration-story\n"
            "Asset type: worksheet picture vocabulary image\n"
            f"Primary request: Create a {stage_lbl} picture vocabulary board with exactly {len(terms)} pictures for these items: {', '.join(terms)}.\n"
            f"Subject: {subject_label or 'general studies'}\n"
            f"Style/medium: {style}\n"
            f"Composition/framing: {grid_layout}\n"
            f"Lighting/mood: {mood}\n"
            "Text (verbatim): none\n"
            f"Constraints: exactly {len(terms)} pictures, no labels, no decorative filler, every item must be easy to identify visually\n"
            "Avoid: abstract icons, posters, watermark"
        )
        return ("picture_vocabulary", topic_label or "Vocabulary", prompt)

    return None


def _exam_prompt(exam_data: dict, section: dict, subject: str, learner_stage: str, topic: str) -> tuple[str, str, str] | None:
    sec_type = _clean_text(section.get("type"))
    subject_label = _clean_text(subject or exam_data.get("subject", "")).replace("_", " ")
    topic_label = _clean_text(topic or section.get("title", "") or exam_data.get("title", ""))
    tier = _stage_tier(learner_stage)
    stage_lbl = _stage_label(learner_stage)
    style = _visual_style(learner_stage)
    mood = _visual_mood(learner_stage)

    if sec_type == "matching":
        if tier == "adult":
            return None
        min_terms = 4 if tier == "secondary" else 3
        terms = _extract_matching_terms(section.get("questions", []))
        if len(terms) < min_terms:
            return None
        grid_layout = _grid_layout_hint(len(terms))
        prompt = (
            "Use case: illustration-story\n"
            "Asset type: exam support image\n"
            f"Primary request: Create a picture board with exactly {len(terms)} pictures for a {stage_lbl} matching section using these items: {', '.join(terms)}.\n"
            f"Subject: {subject_label or 'general studies'}\n"
            "Scene/backdrop: clean worksheet-style background with separate picture cards\n"
            f"Style/medium: {style}\n"
            f"Composition/framing: {grid_layout}\n"
            f"Lighting/mood: {mood}\n"
            "Text (verbatim): none\n"
            f"Constraints: exactly {len(terms)} pictures, no labels, no decorative filler, each item must be visually distinct and usable for matching\n"
            "Avoid: posters, title cards, watermark"
        )
        return ("image_based_matching", topic_label or "Matching", prompt)

    if sec_type == "vocabulary":
        if tier not in ("lower_primary", "upper_primary"):
            return None
        min_terms = 4 if tier == "lower_primary" else 5
        terms = _extract_vocab_terms({"questions": section.get("questions", [])})
        if len(terms) < min_terms:
            return None
        grid_layout = _grid_layout_hint(len(terms))
        prompt = (
            "Use case: illustration-story\n"
            "Asset type: exam picture vocabulary image\n"
            f"Primary request: Create a {stage_lbl} picture vocabulary board with exactly {len(terms)} pictures for these items: {', '.join(terms)}.\n"
            f"Subject: {subject_label or 'general studies'}\n"
            f"Style/medium: {style}\n"
            f"Composition/framing: {grid_layout}\n"
            f"Lighting/mood: {mood}\n"
            "Text (verbatim): none\n"
            f"Constraints: exactly {len(terms)} pictures, no labels, no decorative filler, every item must be easy to recognize\n"
            "Avoid: abstract icons, posters, watermark"
        )
        return ("picture_vocabulary", topic_label or "Vocabulary", prompt)

    if sec_type == "reading_comprehension":
        min_passage = {"lower_primary": 80, "upper_primary": 120, "secondary": 200, "adult": 300}.get(tier, 200)
        passage = _clean_text(section.get("source_text"))
        if len(passage) < min_passage:
            return None
        reading_style = "warm educational story illustration" if tier in ("lower_primary", "upper_primary") else "clear educational illustration suited to the passage content"
        prompt = (
            "Use case: illustration-story\n"
            "Asset type: exam reading support image\n"
            f"Primary request: Create a {stage_lbl} reading support illustration for this passage.\n"
            f"Subject: {subject_label or 'language learning'}\n"
            f"Scene/backdrop: {_passage_excerpt(passage)}\n"
            f"Style/medium: {reading_style}\n"
            "Composition/framing: one coherent scene aligned to the passage\n"
            f"Lighting/mood: {mood}\n"
            "Text (verbatim): none\n"
            "Constraints: no labels, no decorative filler, do not introduce contradictory details\n"
            "Avoid: watermark"
        )
        return ("reading_scene_support", topic_label or "Reading", prompt)

    if sec_type == "true_false":
        if tier not in ("lower_primary", "upper_primary"):
            return None
        if _clean_text(section.get("source_text")):
            return None
        statements = _observable_statements(section.get("questions", []))
        if len(statements) < 3:
            return None
        prompt = (
            "Use case: illustration-story\n"
            "Asset type: exam true/false support image\n"
            f"Primary request: Create one {stage_lbl} scene that supports observation-based true/false statements.\n"
            f"Subject: {subject_label or 'general studies'}\n"
            f"Scene/backdrop: The image should plausibly support statements like: {' | '.join(statements)}\n"
            f"Style/medium: {style}\n"
            "Composition/framing: one coherent scene with visible details students can check\n"
            f"Lighting/mood: {mood}\n"
            "Text (verbatim): none\n"
            "Constraints: no labels, no decorative filler, no irrelevant objects\n"
            "Avoid: watermark"
        )
        return ("scene_truth_support", topic_label or "True or False", prompt)

    if sec_type in {"diagram_questions", "classification", "symbol_identification", "theory_questions", "terminology"}:
        # Diagrams and instructional visuals are pedagogically valuable at all stages.
        cue_text = _passage_excerpt(" ".join(_clean_text(q if isinstance(q, str) else q.get("text", q.get("stem", q.get("word", "")))) for q in (section.get("questions", [])[:4])))
        if not cue_text:
            cue_text = topic_label
        prompt = (
            "Use case: infographic-diagram\n"
            "Asset type: instructional support image\n"
            f"Primary request: Create a {stage_lbl} instructional visual for {subject_label or 'general studies'}.\n"
            f"Subject: {subject_label or 'general studies'}\n"
            f"Scene/backdrop: {cue_text}\n"
            f"Style/medium: {'clean educational diagram or reference illustration' if tier in ('lower_primary', 'upper_primary') else 'precise, informative educational diagram'}\n"
            "Composition/framing: one centered, easy-to-read instructional visual\n"
            f"Lighting/mood: {mood}\n"
            "Text (verbatim): none\n"
            "Constraints: no decorative filler, use only the fewest elements needed to support the task\n"
            "Avoid: posters, banners, watermark"
        )
        return ("instructional_diagram_support", f"{topic_label or 'Visual'} support", prompt)

    return None


def _status_payload(state: str, message: str, *, provider_chain: str = "", task_type: str = "", visual_role: str = "") -> dict:
    return {
        "state": _clean_text(state),
        "message": _clean_text(message),
        "provider_chain": _clean_text(provider_chain),
        "task_type": _clean_text(task_type),
        "visual_role": _clean_text(visual_role),
        "debug_detail": "",
    }


def _extract_image_data_url(response) -> str:
    data = getattr(response, "data", None) or []
    for item in data:
        if isinstance(item, dict):
            b64 = item.get("b64_json")
            url = item.get("url")
        else:
            b64 = getattr(item, "b64_json", None)
            url = getattr(item, "url", None)
        if b64:
            return "data:image/png;base64," + b64
        if url:
            try:
                with urlopen(url) as resp:
                    raw = resp.read()
                return "data:image/png;base64," + base64.b64encode(raw).decode("ascii")
            except Exception:
                continue
    return ""


def _extract_openrouter_image_data_url(payload: dict) -> str:
    try:
        choices = payload.get("choices") or []
        if not choices:
            return ""
        message = choices[0].get("message") or {}
        images = message.get("images") or []
        if not images:
            return ""
        first = images[0] or {}
        image_url = first.get("image_url") or first.get("imageUrl") or {}
        url = image_url.get("url") if isinstance(image_url, dict) else ""
        return _clean_text(url)
    except Exception:
        return ""


def _generate_with_openrouter(prompt: str) -> str:
    api_key = _get_openrouter_api_key()
    if not api_key:
        raise RuntimeError("missing_openrouter_api_key")

    def _request() -> str:
        body = {
            "model": _get_openrouter_image_model(),
            "messages": [{"role": "user", "content": prompt}],
            "modalities": ["image", "text"],
            "stream": False,
            "image_config": {
                "aspect_ratio": "4:3",
                "image_size": "1K",
            },
        }
        req = Request(
            "https://openrouter.ai/api/v1/chat/completions",
            data=json.dumps(body).encode("utf-8"),
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
                "HTTP-Referer": "https://classio.app",
                "X-Title": "Classio",
            },
            method="POST",
        )
        with urlopen(req, timeout=90) as resp:
            payload = json.loads(resp.read().decode("utf-8"))
        return _extract_openrouter_image_data_url(payload)

    return run_with_ai_retries(_request)


def _generate_with_gemini(prompt: str) -> str:
    api_key = _get_gemini_api_key()
    if not api_key:
        raise RuntimeError("missing_gemini_api_key")

    def _request() -> str:
        client = _get_gemini_client(api_key)
        response = client.models.generate_content(
            model=_get_gemini_image_model(),
            contents=[prompt],
        )

        candidate_parts = []
        candidate_parts.extend(getattr(response, "parts", None) or [])
        for candidate in getattr(response, "candidates", None) or []:
            content = getattr(candidate, "content", None)
            candidate_parts.extend(getattr(content, "parts", None) or [])

        for part in candidate_parts:
            inline = getattr(part, "inline_data", None)
            if inline is None and isinstance(part, dict):
                inline = part.get("inline_data") or part.get("inlineData")
            if inline is None:
                continue
            if isinstance(inline, dict):
                mime_type = _clean_text(inline.get("mime_type") or inline.get("mimeType") or "")
                data = inline.get("data")
            else:
                mime_type = _clean_text(getattr(inline, "mime_type", "") or getattr(inline, "mimeType", ""))
                data = getattr(inline, "data", None)
            if not data:
                continue
            if isinstance(data, str):
                encoded = data
            else:
                encoded = base64.b64encode(data).decode("ascii")
            return f"data:{mime_type or 'image/png'};base64,{encoded}"
        return ""

    return run_with_ai_retries(_request)


def _extract_openai_image_data_url(response) -> str:
    data = getattr(response, "data", None) or []
    for item in data:
        b64 = getattr(item, "b64_json", None) if not isinstance(item, dict) else item.get("b64_json")
        url = getattr(item, "url", None) if not isinstance(item, dict) else item.get("url")
        if b64:
            return f"data:image/png;base64,{b64}"
        if url:
            try:
                with urlopen(url) as resp:
                    raw = resp.read()
                return "data:image/png;base64," + base64.b64encode(raw).decode("ascii")
            except Exception:
                continue
    return ""


def _generate_with_openai(prompt: str) -> str:
    api_key = _get_openai_api_key()
    if not api_key:
        raise RuntimeError("missing_openai_api_key")

    def _request() -> str:
        client = _get_openai_client(api_key)
        response = client.images.generate(
            model="gpt-image-1",
            prompt=prompt,
            size="1024x1024",
            quality="low",
        )
        return _extract_openai_image_data_url(response)

    return run_with_ai_retries(_request)


@st.cache_data(show_spinner=False, ttl=60 * 60 * 24 * 7, max_entries=128)
def _generate_ai_image_data_url(prompt: str) -> str:
    last_errors = []
    for provider in _get_image_provider_order():
        try:
            if provider == "openrouter":
                result = _generate_with_openrouter(prompt)
            elif provider == "gemini":
                result = _generate_with_gemini(prompt)
            else:
                result = _generate_with_openai(prompt)
            if result:
                return result
        except Exception as exc:
            last_errors.append(f"{provider}: {exc}")
            continue
    return ""


def _normalize_existing_support(support: dict | None) -> dict | None:
    if not isinstance(support, dict):
        return None
    if _clean_text(support.get("generator")) == "classio_local_visual_v1":
        return None
    return dict(support)


def _build_ready_support(role: str, caption: str, prompt: str, subject: str, learner_stage: str, placement: str) -> dict | None:
    prompt = prompt + _NO_TEXT_RULE
    data_url = _generate_ai_image_data_url(prompt)
    if not data_url:
        return {
            "placement": placement,
            "generator": "classio_ai_visual_v1",
            "status": "planned",
            "purpose": "instructional_support",
            "visual_role": role,
            "caption": caption,
            "alt_text": caption,
            "image_prompt": prompt,
            "subject": _clean_text(subject),
            "learner_stage": _clean_text(learner_stage),
        }
    return {
        "placement": placement,
        "generator": "classio_ai_visual_v1",
        "status": "ready",
        "purpose": "instructional_support",
        "visual_role": role,
        "caption": caption,
        "alt_text": caption,
        "image_prompt": prompt,
        "image_data_url": data_url,
        "subject": _clean_text(subject),
        "learner_stage": _clean_text(learner_stage),
    }


def _should_render_support(support: dict | None) -> bool:
    if not isinstance(support, dict):
        return False
    if _clean_text(support.get("status")) != "ready":
        return False
    data_url = _clean_text(support.get("image_data_url"))
    return data_url.startswith("data:image/")


def worksheet_has_ready_visuals(ws: dict | None) -> bool:
    payload = dict(ws or {})
    for item in payload.get("visual_supports", []) or []:
        if _should_render_support(_normalize_existing_support(item)):
            return True
    return False


def exam_has_ready_visuals(exam_data: dict | None) -> bool:
    payload = dict(exam_data or {})
    for section in payload.get("sections", []) or []:
        if not isinstance(section, dict):
            continue
        if _should_render_support(_normalize_existing_support(section.get("visual_support"))):
            return True
    return False


def enrich_worksheet_with_visuals(ws: dict, *, subject: str = "", learner_stage: str = "", topic: str = "") -> dict:
    payload = dict(ws or {})
    stage = learner_stage or payload.get("learner_stage", "")

    existing_supports = [
        item for item in (_normalize_existing_support(s) for s in payload.get("visual_supports", []))
        if item
    ]
    ready_supports = [item for item in existing_supports if _should_render_support(item)]
    if ready_supports:
        payload["visual_supports"] = ready_supports
        payload["_visual_support_status"] = _status_payload(
            "ready",
            "image_support_ready",
            provider_chain=_provider_chain_label(),
            task_type=payload.get("worksheet_type", ""),
            visual_role=ready_supports[0].get("visual_role", ""),
        )
        return payload

    prompt_spec = _worksheet_prompt(
        payload,
        subject or payload.get("subject", ""),
        stage,
        topic or payload.get("topic", ""),
    )
    if not prompt_spec:
        payload["_visual_support_status"] = _status_payload(
            "skipped",
            "image_support_not_needed",
            provider_chain=_provider_chain_label(),
            task_type=payload.get("worksheet_type", ""),
        )
        payload["visual_supports"] = []
        return payload

    role, caption, prompt = prompt_spec

    if not _image_generation_enabled():
        payload["_visual_support_status"] = _status_payload(
            "provider_unavailable",
            "image_support_provider_unavailable",
            provider_chain=_provider_chain_label(),
            task_type=payload.get("worksheet_type", ""),
            visual_role=role,
        )
        payload["visual_supports"] = []
        return payload

    support = _build_ready_support(
        role,
        caption,
        prompt,
        subject or payload.get("subject", ""),
        stage,
        "worksheet_intro",
    )
    if _should_render_support(support):
        payload["visual_supports"] = [support]
        payload["_visual_support_status"] = _status_payload(
            "ready",
            "image_support_generated",
            provider_chain=_provider_chain_label(),
            task_type=payload.get("worksheet_type", ""),
            visual_role=role,
        )
    else:
        payload["visual_supports"] = []
        payload["_visual_support_status"] = _status_payload(
            "generation_failed",
            "image_support_generation_failed",
            provider_chain=_provider_chain_label(),
            task_type=payload.get("worksheet_type", ""),
            visual_role=role,
        )
    return payload


def regenerate_worksheet_visuals(ws: dict, *, subject: str = "", learner_stage: str = "", topic: str = "") -> dict:
    """Force-regenerate visual supports for an existing worksheet by clearing them first.

    Uses deep copies so the caller's original data is never mutated.
    """
    import copy
    _generate_ai_image_data_url.clear()          # bust cache so the API is called fresh
    payload = copy.deepcopy(ws or {})
    payload["visual_supports"] = []
    return enrich_worksheet_with_visuals(payload, subject=subject, learner_stage=learner_stage, topic=topic)


def regenerate_exam_visuals(exam_data: dict, *, subject: str = "", learner_stage: str = "", topic: str = "") -> dict:
    """Force-regenerate visual supports for an existing exam by clearing them first.

    Uses deep copies so the caller's original data is never mutated.
    """
    import copy
    _generate_ai_image_data_url.clear()          # bust cache so the API is called fresh
    payload = copy.deepcopy(exam_data or {})
    for sec in payload.get("sections", []) or []:
        if isinstance(sec, dict):
            sec.pop("visual_support", None)
    return enrich_exam_with_visuals(payload, subject=subject, learner_stage=learner_stage, topic=topic)


def worksheet_eligible_for_visuals(ws: dict, *, subject: str = "", learner_stage: str = "", topic: str = "") -> bool:
    """Return True if the worksheet type/stage qualifies for image generation."""
    if worksheet_has_ready_visuals(ws):
        return True
    stage = learner_stage or ws.get("learner_stage", "")
    return _worksheet_prompt(ws, subject or ws.get("subject", ""), stage, topic or ws.get("topic", "")) is not None


def exam_eligible_for_visuals(exam_data: dict, *, subject: str = "", learner_stage: str = "", topic: str = "") -> bool:
    """Return True if any exam section qualifies for image generation."""
    if exam_has_ready_visuals(exam_data):
        return True
    stage = learner_stage or exam_data.get("learner_stage", "")
    for sec in exam_data.get("sections", []) or []:
        if isinstance(sec, dict) and _exam_prompt(exam_data, sec, subject or exam_data.get("subject", ""), stage, topic or exam_data.get("topic", "")) is not None:
            return True
    return False


def enrich_exam_with_visuals(exam_data: dict, *, subject: str = "", learner_stage: str = "", topic: str = "") -> dict:
    payload = dict(exam_data or {})
    stage = learner_stage or payload.get("learner_stage", "")
    sections = []
    for section in payload.get("sections", []) or []:
        sec = dict(section or {})
        existing = _normalize_existing_support(sec.get("visual_support"))

        if _should_render_support(existing):
            sec["visual_support"] = existing
            sec["_visual_support_status"] = _status_payload(
                "ready",
                "image_support_ready",
                provider_chain=_provider_chain_label(),
                task_type=sec.get("type", ""),
                visual_role=existing.get("visual_role", ""),
            )
            sections.append(sec)
            continue

        prompt_spec = _exam_prompt(
            payload,
            sec,
            subject or payload.get("subject", ""),
            stage,
            topic or payload.get("topic", ""),
        )
        if not prompt_spec:
            sec["_visual_support_status"] = _status_payload(
                "skipped",
                "image_support_not_needed",
                provider_chain=_provider_chain_label(),
                task_type=sec.get("type", ""),
            )
            sec.pop("visual_support", None)
            sections.append(sec)
            continue

        if not _image_generation_enabled():
            sec["_visual_support_status"] = _status_payload(
                "provider_unavailable",
                "image_support_provider_unavailable",
                provider_chain=_provider_chain_label(),
                task_type=sec.get("type", ""),
                visual_role=prompt_spec[0],
            )
            sec.pop("visual_support", None)
            sections.append(sec)
            continue

        role, caption, prompt = prompt_spec
        support = _build_ready_support(
            role,
            caption,
            prompt,
            subject or payload.get("subject", ""),
            stage,
            "section_intro",
        )
        if _should_render_support(support):
            sec["visual_support"] = support
            sec["_visual_support_status"] = _status_payload(
                "ready",
                "image_support_generated",
                provider_chain=_provider_chain_label(),
                task_type=sec.get("type", ""),
                visual_role=role,
            )
        else:
            sec.pop("visual_support", None)
            sec["_visual_support_status"] = _status_payload(
                "generation_failed",
                "image_support_generation_failed",
                provider_chain=_provider_chain_label(),
                task_type=sec.get("type", ""),
                visual_role=role,
            )
        sections.append(sec)
    payload["sections"] = sections
    return payload


def data_uri_to_bytes(value: str) -> bytes | None:
    text = _clean_text(value)
    if not text.startswith("data:image/"):
        return None
    try:
        _, encoded = text.split(",", 1)
        return base64.b64decode(encoded)
    except Exception:
        return None


def render_streamlit_visual_support(support: dict | None) -> None:
    if not _should_render_support(support):
        return
    data_url = _clean_text(support.get("image_data_url"))
    if not data_url.startswith("data:image/"):
        return
    caption = _clean_text(support.get("caption"))
    caption_html = f'<p style="text-align:center;color:var(--muted);font-size:0.85rem;margin:6px 0 0;">{caption}</p>' if caption else ""
    st.markdown(
        f'<div style="text-align:center;margin:12px auto 16px;max-width:520px;">'
        f'<img src="{data_url}" style="max-width:100%;border-radius:10px;box-shadow:0 2px 8px rgba(0,0,0,0.10);" />'
        f'{caption_html}'
        f'</div>',
        unsafe_allow_html=True,
    )


def render_streamlit_visual_supports(items) -> None:
    for item in items or []:
        render_streamlit_visual_support(item)


def render_visual_support_status(status: dict | None, *, compact: bool = False) -> None:
    if not isinstance(status, dict):
        return
    state = _clean_text(status.get("state"))
    message_key = _clean_text(status.get("message"))
    provider_chain = _clean_text(status.get("provider_chain"))
    visual_role = _clean_text(status.get("visual_role"))
    debug_detail = _clean_text(status.get("debug_detail"))

    body = t(message_key) if message_key and t(message_key) != message_key else message_key
    if not body:
        return

    if compact:
        if state == "ready":
            st.caption(body)
        elif state in {"provider_unavailable", "generation_failed"}:
            st.warning(body)
        else:
            st.info(body)
        return

    with st.expander(t("image_support_status"), expanded=False):
        if state == "ready":
            st.success(body)
        elif state in {"provider_unavailable", "generation_failed"}:
            st.warning(body)
        else:
            st.info(body)
        if provider_chain:
            st.caption(f"{t('image_support_provider_chain')}: {provider_chain}")
        if visual_role:
            st.caption(f"{t('image_support_role')}: {visual_role.replace('_', ' ')}")
        if debug_detail:
            st.caption(f"{t('image_support_debug_details')}: {debug_detail}")


def render_visual_support_status_group(items, *, title: str | None = None) -> None:
    normalized = []
    for item in items or []:
        if isinstance(item, dict):
            label = _clean_text(item.get("label"))
            status = item.get("status")
        else:
            label = ""
            status = item
        if isinstance(status, dict):
            normalized.append({"label": label, "status": status})

    if not normalized:
        return

    with st.expander(title or t("image_support_status"), expanded=False):
        for idx, item in enumerate(normalized):
            label = item.get("label", "")
            status = item.get("status") or {}
            state = _clean_text(status.get("state"))
            message_key = _clean_text(status.get("message"))
            provider_chain = _clean_text(status.get("provider_chain"))
            visual_role = _clean_text(status.get("visual_role"))
            debug_detail = _clean_text(status.get("debug_detail"))

            body = t(message_key) if message_key and t(message_key) != message_key else message_key
            if not body:
                continue

            if idx > 0:
                st.markdown("---")

            if label:
                st.markdown(f"**{label}**")

            if state == "ready":
                st.success(body)
            elif state in {"provider_unavailable", "generation_failed"}:
                st.warning(body)
            else:
                st.info(body)

            if provider_chain:
                st.caption(f"{t('image_support_provider_chain')}: {provider_chain}")
            if visual_role:
                st.caption(f"{t('image_support_role')}: {visual_role.replace('_', ' ')}")
            if debug_detail:
                st.caption(f"{t('image_support_debug_details')}: {debug_detail}")


def build_pdf_visual_flowables(support: dict | None, *, max_width_cm: float, paragraph_style, max_height_cm: float = 0):
    if not _should_render_support(support):
        return []
    from reportlab.lib.units import cm
    from reportlab.platypus import Image as RLImage, Paragraph, Spacer

    image_bytes = data_uri_to_bytes(support.get("image_data_url", ""))
    if not image_bytes:
        return []
    img = Image.open(BytesIO(image_bytes))
    width_px, height_px = img.size
    target_width = max_width_cm * cm
    ratio = height_px / max(width_px, 1)
    target_height = target_width * ratio
    # Cap height if requested (e.g. word search needs compact images)
    if max_height_cm > 0 and target_height > max_height_cm * cm:
        target_height = max_height_cm * cm
        target_width = target_height / max(ratio, 0.01)
    flowables = [RLImage(BytesIO(image_bytes), width=target_width, height=target_height), Spacer(1, 6)]
    return flowables


def add_docx_visual_support(doc, support: dict | None, *, width_cm: float, font_name: str, font_size_pt: float, max_height_cm: float = 0):
    if not _should_render_support(support):
        return
    from docx.shared import Cm, Pt

    image_bytes = data_uri_to_bytes(support.get("image_data_url", ""))
    if not image_bytes:
        return
    para = doc.add_paragraph()
    run = para.add_run()
    # Cap height if requested to keep content on one page
    if max_height_cm > 0:
        img = Image.open(BytesIO(image_bytes))
        w_px, h_px = img.size
        ratio = h_px / max(w_px, 1)
        effective_height = width_cm * ratio
        if effective_height > max_height_cm:
            run.add_picture(BytesIO(image_bytes), height=Cm(max_height_cm))
        else:
            run.add_picture(BytesIO(image_bytes), width=Cm(width_cm))
    else:
        run.add_picture(BytesIO(image_bytes), width=Cm(width_cm))
    # Caption removed – image stands alone
    if False:  # caption block disabled
        cap = doc.add_paragraph()
        cap.alignment = 1
        cap_run = cap.add_run("")
        try:
            cap_run.font.name = font_name
            cap_run.font.size = Pt(font_size_pt)
        except Exception:
            pass


# ── Watermark helpers ────────────────────────────────────────────────

def _make_watermark_image(image_bytes: bytes, opacity: float = 0.10) -> bytes:
    """Blend *image_bytes* with white to produce a faded watermark PNG."""
    img = Image.open(BytesIO(image_bytes)).convert("RGBA")
    white = Image.new("RGBA", img.size, (255, 255, 255, 255))
    blended = Image.blend(white, img, opacity)
    buf = BytesIO()
    blended.convert("RGB").save(buf, format="PNG")
    return buf.getvalue()


def get_visual_watermark_bytes(resource_dict: dict, *, opacity: float = 0.10) -> bytes | None:
    """Return watermark PNG bytes from the first ready visual support, or *None*."""
    for support in resource_dict.get("visual_supports", []) or []:
        if _should_render_support(support):
            raw = data_uri_to_bytes(support.get("image_data_url", ""))
            if raw:
                return _make_watermark_image(raw, opacity=opacity)
    return None
