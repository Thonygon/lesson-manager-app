# CLASSIO — Word (DOCX) Export Generator
# ============================================================
"""
Generate Word (.docx) files for worksheets and lesson plans,
applying the user's selected font and size preferences.
"""
import os
import ast
import urllib.request
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from core.i18n import t


# ── Colour constants matching pdf_styles.py ──────────────────────────
_PRIMARY = RGBColor(0x1D, 0x4E, 0xD8)
_TEXT = RGBColor(0x0F, 0x17, 0x2A)
_TEXT_MUTED = RGBColor(0x33, 0x41, 0x55)
_BORDER = RGBColor(0xCB, 0xD5, 0xE1)
_BG_SUBTLE = "F8FAFC"

_DEFAULT_LOGO_PATH = os.path.abspath(
    os.path.join(os.path.dirname(__file__), os.pardir, "static", "logo_classio_light.png")
)


def _apply_font(run, font_name: str, size_pt: float, bold: bool = False, color: RGBColor = _TEXT):
    """Apply font settings to a run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = color


def _set_line_spacing(paragraph, ratio: float):
    """Set line spacing on a paragraph (e.g. 1.15 = 115%)."""
    from docx.shared import Pt as _Pt
    pf = paragraph.paragraph_format
    pf.line_spacing = ratio


def _add_heading(doc: Document, text: str, font_name: str, size: float,
                 color: RGBColor = _PRIMARY, align=WD_ALIGN_PARAGRAPH.LEFT,
                 leading_ratio: float = 1.15):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    _apply_font(run, font_name, size, bold=True, color=color)
    p.paragraph_format.space_after = Pt(4)
    _set_line_spacing(p, leading_ratio)
    return p


def _add_body(doc: Document, text: str, font_name: str, size: float,
              bold: bool = False, leading_ratio: float = 1.15):
    p = doc.add_paragraph()
    run = p.add_run(str(text or ""))
    _apply_font(run, font_name, size, bold=bold)
    p.paragraph_format.space_after = Pt(3)
    _set_line_spacing(p, leading_ratio)
    return p


def _add_separator(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = _BORDER


def _strip_leading_enum(text: str) -> str:
    """Remove leading '1.', 'a)', etc."""
    import re
    return re.sub(r"^\s*(\d+[\.\)]\s*|[a-z][\.\)]\s*)", "", str(text or "")).strip()


def _sentence_case_fragment(text: str) -> str:
    if isinstance(text, str):
        stripped = text.strip()
        if stripped.startswith("{") and stripped.endswith("}"):
            try:
                parsed = ast.literal_eval(stripped)
            except Exception:
                parsed = None
            if isinstance(parsed, dict) and len(parsed) == 1:
                text = next(iter(parsed.keys()))
    cleaned = _strip_leading_enum(text)
    if not cleaned:
        return ""
    if any(ch.isupper() for ch in cleaned):
        return cleaned
    chars = list(cleaned)
    for idx, ch in enumerate(chars):
        if ch.isalpha():
            chars[idx] = ch.upper()
            return "".join(chars)
    return cleaned


def _true_false_mark_label() -> str:
    true_mark = t("true_false_true_mark")
    false_mark = t("true_false_false_mark")
    if true_mark == "true_false_true_mark":
        true_mark = "T"
    if false_mark == "true_false_false_mark":
        false_mark = "F"
    return f"{true_mark} ☐   {false_mark} ☐"


def _coerce_legacy_mapping(value):
    if isinstance(value, str):
        stripped = value.strip()
        if stripped.startswith("{") and stripped.endswith("}"):
            try:
                parsed = ast.literal_eval(stripped)
            except Exception:
                parsed = None
            if isinstance(parsed, dict):
                return parsed
    return value


def _extract_mapping_side(value, *, prefer: str = "value") -> str:
    value = _coerce_legacy_mapping(value)
    if isinstance(value, dict):
        if prefer in value:
            return _sentence_case_fragment(value.get(prefer, ""))
        if len(value) == 1:
            only_key, only_val = next(iter(value.items()))
            return _sentence_case_fragment(only_val if prefer == "value" else only_key)
        fallback = value.get("text", value.get("answer", str(value)))
        return _sentence_case_fragment(fallback)
    return _sentence_case_fragment(value)


def _fetch_image_bytes(url: str) -> bytes | None:
    """Fetch image bytes from URL. Returns None on failure."""
    if not url:
        return None
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Classio/1.0"})
        with urllib.request.urlopen(req, timeout=8) as resp:
            return resp.read()
    except Exception:
        return None


def _add_logo_to_doc(doc: Document, branding: dict):
    """Add branding logo or default Classio logo to the document."""
    header_enabled = branding.get("header_enabled", False)
    logo_url = str(branding.get("header_logo_url") or "").strip()

    img_bytes = None
    if header_enabled and logo_url:
        img_bytes = _fetch_image_bytes(logo_url)

    if img_bytes:
        buf = BytesIO(img_bytes)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(buf, height=Cm(3.0))
        p.paragraph_format.space_after = Pt(4)
    elif os.path.isfile(_DEFAULT_LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (header_enabled and branding.get("header_style") == "school") else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        run.add_picture(_DEFAULT_LOGO_PATH, height=Cm(2.2))
        p.paragraph_format.space_after = Pt(4)


def _add_footer_with_image(doc: Document, branding: dict, font_name: str):
    """Add footer with optional image and brand text."""
    brand_name = str(branding.get("brand_name") or "").strip()
    footer_enabled = branding.get("footer_enabled", False)
    footer_url = str(branding.get("footer_image_url") or "").strip()

    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False

    # Footer image
    if footer_enabled and footer_url:
        img_bytes = _fetch_image_bytes(footer_url)
        if img_bytes:
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fp.add_run()
            run.add_picture(BytesIO(img_bytes), height=Cm(1.5))
            # Add text footer on a new line
            fp2 = footer.add_paragraph()
            fp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run2 = fp2.add_run(brand_name if (footer_enabled and brand_name) else "Classio")
            _apply_font(run2, font_name, 9, color=_TEXT_MUTED)
            return

    # Text-only footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_label = brand_name if (footer_enabled and brand_name) else "Classio"
    run = fp.add_run(footer_label)
    _apply_font(run, font_name, 9, color=_TEXT_MUTED)


def _set_cell_font(cell, text: str, font_name: str, size: float,
                   bold: bool = False, color: RGBColor = _TEXT):
    """Set text and font in a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text or ""))
    _apply_font(run, font_name, size, bold=bold, color=color)
    return p


def _shade_cell(cell, hex_color: str):
    """Apply background shading to a cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell, bottom: bool = False, right: bool = False):
    """Set thin borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        + (f'<w:bottom w:val="single" w:sz="4" w:color="CBD5E1"/>' if bottom else '')
        + (f'<w:right w:val="single" w:sz="4" w:color="CBD5E1"/>' if right else '')
        + '</w:tcBorders>'
    )
    tcPr.append(borders)


def _set_table_col_width(table, col_idx: int, width_cm: float):
    width = Cm(width_cm)
    for cell in table.columns[col_idx].cells:
        cell.width = width


# ── Worksheet DOCX generator ────────────────────────────────────────

def generate_docx_worksheet(ws: dict, student_only: bool = True) -> bytes:
    """
    Generate a Word document from a worksheet dict.
    Returns raw .docx bytes.
    """
    from helpers.branding import get_user_branding
    from helpers.font_manager import get_docx_font_name, get_font_sizes

    branding = get_user_branding()
    font_key = branding.get("branding_font", "dejavu")
    size_key = branding.get("branding_font_size", "standard")

    font_name = get_docx_font_name(font_key)
    sz = get_font_sizes(size_key)
    t_sz, sec_sz, b_sz = sz["title"], sz["section"], sz["body"]
    lr = sz.get("leading_ratio", 1.15)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(b_sz)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ── Logo ─────────────────────────────────────────────────
    _add_logo_to_doc(doc, branding)

    # ── Header (branding) ────────────────────────────────────
    brand_name = branding.get("brand_name", "").strip()
    department = branding.get("department", "").strip()
    header_enabled = branding.get("header_enabled", False)

    if header_enabled and brand_name:
        _add_heading(doc, brand_name, font_name, sec_sz + 1, color=_PRIMARY,
                     align=WD_ALIGN_PARAGRAPH.CENTER, leading_ratio=lr)
    if header_enabled and department:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(department)
        _apply_font(run, font_name, b_sz, color=_TEXT_MUTED)
        _set_line_spacing(p, lr)

    # Title
    title = str(ws.get("title") or t("untitled_worksheet")).strip()
    _add_heading(doc, title, font_name, t_sz, color=_PRIMARY, leading_ratio=lr)

    # School header fields (Name / Class / Date in a table)
    if branding.get("header_style") == "school" and header_enabled:
        table = doc.add_table(rows=1, cols=3)
        table.autofit = True
        cells = table.rows[0].cells
        _set_cell_font(cells[0], f"{t('student_name_label')}: _________________________________", font_name, b_sz, bold=True)
        _set_cell_font(cells[1], f"{t('class_label')}: __________", font_name, b_sz, bold=True)
        _set_cell_font(cells[2], f"{t('date_label')}: __________", font_name, b_sz, bold=True)
        # Remove table borders
        for cell in cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tcPr.append(borders)
        _add_separator(doc)

    # Instructions
    instructions = str(ws.get("instructions") or "").strip()
    if instructions:
        _add_heading(doc, t("ws_instructions") if t("ws_instructions") != "ws_instructions" else "Instructions",
                     font_name, sec_sz, color=_TEXT, leading_ratio=lr)
        _add_body(doc, instructions, font_name, b_sz, leading_ratio=lr)

    # Vocabulary bank (table layout)
    vocab = ws.get("vocabulary_bank", [])
    if vocab:
        _render_vocab_bank_docx(doc, vocab, font_name, sec_sz, b_sz, lr)

    # Reading passage
    ws_type = ws.get("worksheet_type", "")
    if ws_type == "reading_comprehension" and ws.get("reading_passage"):
        _add_heading(doc, t("ws_reading_passage") if t("ws_reading_passage") != "ws_reading_passage" else "Reading Passage",
                     font_name, sec_sz, color=_TEXT, leading_ratio=lr)
        _add_body(doc, ws["reading_passage"], font_name, b_sz, leading_ratio=lr)

    # ── Questions / exercises ────────────────────────────────
    questions = ws.get("questions", [])

    if ws_type == "matching":
        _render_matching_docx(doc, ws, font_name, sec_sz, b_sz, student_only, lr)
    elif ws_type == "true_false":
        _render_true_false_docx(doc, ws, font_name, sec_sz, b_sz, student_only, lr)
    elif ws_type == "word_search_vocab":
        _render_word_search_docx(doc, ws, font_name, sec_sz, b_sz, lr)
    elif ws_type == "multiple_choice":
        _render_multiple_choice_docx(doc, ws, font_name, sec_sz, b_sz, student_only, lr)
    elif ws_type == "short_answer" and questions:
        _add_heading(doc, t("ws_questions") if t("ws_questions") != "ws_questions" else "Questions",
                     font_name, sec_sz, color=_TEXT, leading_ratio=lr)
        for idx, q in enumerate(questions, 1):
            q_text = q if isinstance(q, str) else str(q.get("stem") or q.get("question") or q)
            _add_body(doc, f"{idx}. {q_text}", font_name, b_sz, leading_ratio=lr)
            _add_body(doc, "_" * 50, font_name, b_sz, leading_ratio=lr)
    elif ws_type in ("fill_in_the_blanks", "fill_in_blank") and questions:
        _add_heading(doc, t("ws_questions") if t("ws_questions") != "ws_questions" else "Questions",
                     font_name, sec_sz, color=_TEXT, leading_ratio=lr)
        for idx, q in enumerate(questions, 1):
            q_text = q if isinstance(q, str) else str(q.get("stem") or q.get("sentence") or q)
            _add_body(doc, f"{idx}. {q_text}", font_name, b_sz, leading_ratio=lr)
    elif ws_type == "error_correction" and questions:
        _add_heading(doc, t("ws_questions") if t("ws_questions") != "ws_questions" else "Questions",
                     font_name, sec_sz, color=_TEXT, leading_ratio=lr)
        for idx, q in enumerate(questions, 1):
            q_text = q if isinstance(q, str) else str(q.get("stem") or q.get("sentence") or q)
            _add_body(doc, f"{idx}. {q_text}", font_name, b_sz, leading_ratio=lr)
            _add_body(doc, f"    {t('correction') if t('correction') != 'correction' else 'Correction'}: _____________________", font_name, b_sz, leading_ratio=lr)
    elif ws_type == "reading_comprehension" and questions:
        _add_heading(doc, t("ws_questions") if t("ws_questions") != "ws_questions" else "Questions",
                     font_name, sec_sz, color=_TEXT, leading_ratio=lr)
        for idx, q in enumerate(questions, 1):
            q_text = q if isinstance(q, str) else str(q.get("stem") or q.get("question") or q)
            _add_body(doc, f"{idx}. {q_text}", font_name, b_sz, leading_ratio=lr)
            _add_body(doc, "_" * 50, font_name, b_sz, leading_ratio=lr)
    elif questions:
        _add_heading(doc, t("ws_questions") if t("ws_questions") != "ws_questions" else "Questions",
                     font_name, sec_sz, color=_TEXT, leading_ratio=lr)
        for idx, q in enumerate(questions, 1):
            if isinstance(q, dict):
                _render_question_item_docx(doc, q, idx, font_name, b_sz, student_only, lr)
            else:
                _add_body(doc, f"{idx}. {q}", font_name, b_sz, leading_ratio=lr)

    # Answer key section (teacher version)
    if not student_only:
        _add_separator(doc)
        _add_heading(doc, t("ws_answer_key") if t("ws_answer_key") != "ws_answer_key" else "Answer Key",
                     font_name, t_sz, color=_PRIMARY, leading_ratio=lr)
        _render_answer_key_docx(doc, ws, font_name, b_sz, lr)

    # Footer (with optional image)
    _add_footer_with_image(doc, branding, font_name)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_question_item_docx(doc, q: dict, idx: int, font_name: str, b_sz: float,
                               student_only: bool, lr: float = 1.15):
    """Render a single question item (MC, FIB, short_answer, etc.)."""
    q_type = q.get("type", "").lower()
    stem = str(q.get("stem") or q.get("question") or q.get("sentence") or "").strip()

    if q_type in ("mc", "multiple_choice"):
        _add_body(doc, f"{idx}. {stem}", font_name, b_sz, bold=True, leading_ratio=lr)
        for opt_idx, opt in enumerate(q.get("options", [])):
            letter = chr(65 + opt_idx)
            _add_body(doc, f"    {letter}) {opt}", font_name, b_sz, leading_ratio=lr)
    elif q_type in ("fib", "fill_in_blank", "fill_in_the_blank"):
        _add_body(doc, f"{idx}. {stem}", font_name, b_sz, leading_ratio=lr)
    elif q_type in ("tf", "true_false"):
        _add_body(doc, f"{idx}. {stem}", font_name, b_sz, leading_ratio=lr)
        _add_body(doc, "    True ☐   False ☐", font_name, b_sz, leading_ratio=lr)
    elif q_type in ("short_answer",):
        _add_body(doc, f"{idx}. {stem}", font_name, b_sz, leading_ratio=lr)
        _add_body(doc, "_" * 50, font_name, b_sz, leading_ratio=lr)
    elif q_type in ("error_correction",):
        _add_body(doc, f"{idx}. {stem}", font_name, b_sz, leading_ratio=lr)
        _add_body(doc, f"    {t('correction')}: _____________________", font_name, b_sz, leading_ratio=lr)
    else:
        _add_body(doc, f"{idx}. {stem}", font_name, b_sz, leading_ratio=lr)


def _render_vocab_bank_docx(doc, vocab_list: list, font_name: str, sec_sz: float,
                            b_sz: float, lr: float = 1.15):
    """Render vocabulary bank as a table (2-col or 5-col like PDF)."""
    _add_heading(doc, t("ws_vocabulary_bank") if t("ws_vocabulary_bank") != "ws_vocabulary_bank" else "Vocabulary",
                 font_name, sec_sz, color=_TEXT, leading_ratio=lr)

    cleaned = [str(x).strip() for x in vocab_list if str(x).strip()]
    has_definitions = any((" - " in item or ":" in item) for item in cleaned)

    if has_definitions:
        # 2-column table: term | definition
        table = doc.add_table(rows=0, cols=2)
        table.autofit = True
        for idx, raw in enumerate(cleaned):
            row = table.add_row()
            if " - " in raw:
                parts = raw.split(" - ", 1)
                _set_cell_font(row.cells[0], parts[0].strip().capitalize(), font_name, b_sz, bold=True)
                _set_cell_font(row.cells[1], parts[1].strip().capitalize(), font_name, b_sz)
            elif ":" in raw:
                parts = raw.split(":", 1)
                _set_cell_font(row.cells[0], parts[0].strip().capitalize(), font_name, b_sz, bold=True)
                _set_cell_font(row.cells[1], parts[1].strip().capitalize(), font_name, b_sz)
            else:
                _set_cell_font(row.cells[0], raw.capitalize(), font_name, b_sz, bold=True)
                _set_cell_font(row.cells[1], "", font_name, b_sz)
            # Alternate row shading
            if idx % 2 == 1:
                _shade_cell(row.cells[0], _BG_SUBTLE)
                _shade_cell(row.cells[1], _BG_SUBTLE)
            _set_cell_borders(row.cells[0], right=True)
    else:
        # 4-column table for single words
        col_count = min(4, len(cleaned))
        if col_count < 1:
            col_count = 1
        table = doc.add_table(rows=0, cols=col_count)
        table.autofit = True
        row_cells = None
        for idx, word in enumerate(cleaned):
            col_idx = idx % col_count
            if col_idx == 0:
                row_obj = table.add_row()
                row_cells = row_obj.cells
            _set_cell_font(row_cells[col_idx], word.capitalize(), font_name, b_sz, bold=True)
            if col_idx < col_count - 1:
                _set_cell_borders(row_cells[col_idx], right=True)
            if (idx // col_count) % 2 == 1:
                _shade_cell(row_cells[col_idx], _BG_SUBTLE)

    doc.add_paragraph()  # spacer


def _render_matching_docx(doc, ws: dict, font_name: str, sec_sz: float,
                          b_sz: float, student_only: bool, lr: float = 1.15):
    """Render matching exercise as a 3-column table like the PDF.
    Uses matching_pairs / left_items+right_items (same data path as PDF)."""
    _add_heading(doc, t("ws_matching_task") if t("ws_matching_task") != "ws_matching_task" else "Match the items",
                 font_name, sec_sz, color=_TEXT, leading_ratio=lr)

    # Extract pairs using the same logic as the PDF renderer
    pairs = ws.get("matching_pairs") or []
    left_items = []
    right_items = []

    if isinstance(pairs, list) and pairs:
        for p in pairs:
            if isinstance(p, dict):
                left = str(p.get("left") or p.get("term") or "").strip()
                right = str(p.get("right") or p.get("definition") or "").strip()
                if left and right:
                    left_items.append(left)
                    right_items.append(right)
    else:
        # Fallback: left_items / right_items lists
        li = ws.get("left_items") or []
        ri = ws.get("right_items") or []
        if isinstance(li, list) and isinstance(ri, list):
            for l_val, r_val in zip(li, ri):
                left_items.append(str(l_val).strip())
                right_items.append(str(r_val).strip())

    if not left_items:
        return

    import random
    rng = random.Random("|".join(left_items + right_items))
    shuffled_right = right_items[:]
    rng.shuffle(shuffled_right)

    max_len = max(len(left_items), len(shuffled_right))
    table = doc.add_table(rows=max_len, cols=3)
    table.autofit = True

    for i in range(max_len):
        cells = table.rows[i].cells
        left_text = f"{i+1}. {_extract_mapping_side(left_items[i], prefer='key')}" if i < len(left_items) else ""
        box_text = "[   ]" if i < len(left_items) else ""
        right_text = f"{chr(97+i)}) {_extract_mapping_side(shuffled_right[i], prefer='value')}" if i < len(shuffled_right) else ""

        _set_cell_font(cells[0], left_text, font_name, b_sz)
        p_box = _set_cell_font(cells[1], box_text, font_name, b_sz)
        p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_font(cells[2], right_text, font_name, b_sz)

        # Remove all borders for clean look
        for cell in cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tcPr.append(borders)

    doc.add_paragraph()  # spacer


def _render_true_false_docx(doc, ws: dict, font_name: str, sec_sz: float,
                            b_sz: float, student_only: bool, lr: float = 1.15):
    """Render true/false exercise as a 2-column table like the PDF.
    Uses true_false_statements / source_text (same data path as PDF)."""
    # Source text (same fallback chain as PDF)
    source_text = ""
    for key in ("source_text", "reading_passage", "text"):
        val = str(ws.get(key) or "").strip()
        if val:
            source_text = val
            break

    if source_text:
        _add_heading(doc, t("read_and_decide_true_false") if t("read_and_decide_true_false") != "read_and_decide_true_false" else "Read and decide: True or False",
                     font_name, sec_sz, color=_TEXT, leading_ratio=lr)
        _add_body(doc, source_text, font_name, b_sz, leading_ratio=lr)

    # Statements (same fallback chain as PDF: true_false_statements → questions)
    statements = ws.get("true_false_statements") or ws.get("questions") or []
    if isinstance(statements, list):
        statements = [str(x).strip() for x in statements if str(x).strip()]
    if statements:
        _add_heading(doc, t("ws_questions") if t("ws_questions") != "ws_questions" else "Questions",
                     font_name, sec_sz, color=_TEXT, leading_ratio=lr)

        table = doc.add_table(rows=len(statements), cols=2)
        table.autofit = False
        _set_table_col_width(table, 0, 12.8)
        _set_table_col_width(table, 1, 3.6)

        for idx, item in enumerate(statements):
            text = item if isinstance(item, str) else str(item.get("statement") or item.get("stem") or item)
            cells = table.rows[idx].cells
            _set_cell_font(cells[0], f"{idx+1}. {_strip_leading_enum(text)}", font_name, b_sz)
            p = _set_cell_font(cells[1], _true_false_mark_label(), font_name, b_sz)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Remove borders for clean look
            for cell in cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                borders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    '<w:top w:val="none" w:sz="0" w:color="auto"/>'
                    '<w:bottom w:val="none" w:sz="0" w:color="auto"/>'
                    '<w:left w:val="none" w:sz="0" w:color="auto"/>'
                    '<w:right w:val="none" w:sz="0" w:color="auto"/>'
                    '</w:tcBorders>'
                )
                tcPr.append(borders)

        doc.add_paragraph()  # spacer


def _render_word_search_docx(doc, ws: dict, font_name: str, sec_sz: float,
                             b_sz: float, lr: float = 1.15, show_heading: bool = True):
    """Render word search grid as a table."""
    grid, _ = _resolve_word_search_grid(ws)
    if not grid:
        return

    if show_heading:
        _add_heading(doc, t("word_search_grid") if t("word_search_grid") != "word_search_grid" else "Word Search",
                     font_name, sec_sz, color=_TEXT, leading_ratio=lr)

    grid_size = len(grid)
    table = doc.add_table(rows=grid_size, cols=grid_size)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r, row_data in enumerate(grid):
        for c, ch in enumerate(row_data):
            cell = table.rows[r].cells[c]
            p = _set_cell_font(cell, str(ch).upper(), font_name, max(b_sz, 11), bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Light grid borders
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="single" w:sz="4" w:color="94A3B8"/>'
                '<w:bottom w:val="single" w:sz="4" w:color="94A3B8"/>'
                '<w:left w:val="single" w:sz="4" w:color="94A3B8"/>'
                '<w:right w:val="single" w:sz="4" w:color="94A3B8"/>'
                '</w:tcBorders>'
            )
            tcPr.append(borders)

    doc.add_paragraph()  # spacer


def _extract_word_search_words(payload: dict) -> list[str]:
    raw_words = payload.get("vocabulary_bank") or payload.get("words") or []
    if not raw_words:
        raw_words = payload.get("questions") or []

    words = []
    for item in raw_words:
        if isinstance(item, dict):
            candidate = (
                item.get("word")
                or item.get("term")
                or item.get("left")
                or item.get("text")
                or item.get("stem")
                or ""
            )
        else:
            candidate = item
        text = str(candidate or "").strip()
        if text:
            words.append(text)
    return words


def _resolve_word_search_grid(payload: dict) -> tuple[list[list[str]], list[dict]]:
    grid = payload.get("wordsearch_grid") or payload.get("grid") or []
    placements = payload.get("wordsearch_placements") or payload.get("placements") or []
    if grid:
        return grid, placements

    words = _extract_word_search_words(payload)
    if not words:
        return [], []

    try:
        from helpers.worksheet_storage import _generate_wordsearch_grid, _normalize_wordsearch_words

        normalized_words = _normalize_wordsearch_words(words)
        if not normalized_words:
            return [], []
        seed = "|".join(normalized_words)
        size = payload.get("wordsearch_size") or payload.get("size") or 12
        grid, _, placements = _generate_wordsearch_grid(
            normalized_words,
            seed=seed,
            size=size,
        )
        return grid or [], placements or []
    except Exception:
        return [], []


def _render_multiple_choice_docx(doc, ws: dict, font_name: str, sec_sz: float,
                                 b_sz: float, student_only: bool, lr: float = 1.15):
    """Render multiple choice exercise using multiple_choice_items (same data as PDF)."""
    import re

    # Same extraction logic as PDF: multiple_choice_items first, then parse questions
    items = ws.get("multiple_choice_items") or []
    if not isinstance(items, list):
        items = []

    # Normalize: ensure each item has stem, options
    mc_items = []
    for item in items:
        if isinstance(item, dict):
            stem = str(item.get("stem") or item.get("question") or "").strip()
            options = item.get("options") or []
            answer = str(item.get("answer") or item.get("correct_answer") or "").strip()
            if stem and options:
                mc_items.append({"stem": stem, "options": options, "answer": answer})

    # Fallback: parse from questions
    if not mc_items:
        questions = ws.get("questions") or []
        for q in questions:
            text = str(q).strip()
            m = re.match(
                r"^(.*?)(?:\s+A\)|\s+A\.)(.*?)(?:\s+B\)|\s+B\.)(.*?)(?:\s+C\)|\s+C\.)(.*?)(?:(?:\s+D\)|\s+D\.)(.*))?$",
                text,
                flags=re.IGNORECASE,
            )
            if m:
                stem = m.group(1).strip()
                opts = [m.group(2).strip(), m.group(3).strip(), m.group(4).strip()]
                if m.group(5):
                    opts.append(m.group(5).strip())
                mc_items.append({"stem": stem, "options": opts, "answer": ""})

    if not mc_items:
        return

    _add_heading(doc, t("ws_questions") if t("ws_questions") != "ws_questions" else "Questions",
                 font_name, sec_sz, color=_TEXT, leading_ratio=lr)

    for idx, item in enumerate(mc_items, 1):
        stem = _strip_leading_enum(str(item.get("stem") or ""))
        _add_body(doc, f"{idx}. {stem}", font_name, b_sz, bold=True, leading_ratio=lr)
        for opt_idx, opt in enumerate(item.get("options", [])):
            letter = chr(65 + opt_idx)
            opt_text = _strip_leading_enum(str(opt))
            _add_body(doc, f"    {letter}) {opt_text}", font_name, b_sz, leading_ratio=lr)

    doc.add_paragraph()  # spacer


def _render_answer_key_docx(doc, ws: dict, font_name: str, b_sz: float, lr: float = 1.15):
    """Render the answer key section for teacher version."""
    ws_type = ws.get("worksheet_type", "")

    # Matching: show correct pairs
    if ws_type == "matching":
        pairs = ws.get("matching_pairs") or []
        if isinstance(pairs, list) and pairs:
            for i, p in enumerate(pairs):
                if isinstance(p, dict):
                    left = str(p.get("left") or p.get("term") or "").strip()
                    right = str(p.get("right") or p.get("definition") or "").strip()
                    _add_body(doc, f"{i+1}. {_strip_leading_enum(left)} → {_strip_leading_enum(right)}",
                              font_name, b_sz, leading_ratio=lr)
        else:
            li = ws.get("left_items") or []
            ri = ws.get("right_items") or []
            for i, (l_val, r_val) in enumerate(zip(li, ri)):
                _add_body(doc, f"{i+1}. {_strip_leading_enum(str(l_val))} → {_strip_leading_enum(str(r_val))}",
                          font_name, b_sz, leading_ratio=lr)
        return

    # True/False: show answers from answer_key
    if ws_type == "true_false":
        answer_key = ws.get("answer_key")
        if isinstance(answer_key, list):
            for i, line in enumerate(answer_key, 1):
                if str(line).strip():
                    _add_body(doc, f"{i}. {str(line).strip()}", font_name, b_sz, leading_ratio=lr)
        elif answer_key and str(answer_key).strip():
            for line in str(answer_key).strip().splitlines():
                if line.strip():
                    _add_body(doc, line.strip(), font_name, b_sz, leading_ratio=lr)
        return

    # Multiple choice: show answers
    if ws_type == "multiple_choice":
        items = ws.get("multiple_choice_items") or []
        for i, item in enumerate(items, 1):
            if isinstance(item, dict):
                answer = str(item.get("answer") or item.get("correct_answer") or "").strip()
                if answer:
                    _add_body(doc, f"{i}. {answer}", font_name, b_sz, leading_ratio=lr)
        if not any(isinstance(it, dict) and (it.get("answer") or it.get("correct_answer")) for it in items):
            # Fallback to freeform answer_key
            ak = ws.get("answer_key", "")
            if ak:
                for line in str(ak).strip().splitlines():
                    if line.strip():
                        _add_body(doc, line.strip(), font_name, b_sz, leading_ratio=lr)
        return

    # Generic: use answer_key string or questions with answers
    answer_key = ws.get("answer_key")
    if isinstance(answer_key, list):
        for i, line in enumerate(answer_key, 1):
            if str(line).strip():
                _add_body(doc, f"{i}. {str(line).strip()}", font_name, b_sz, leading_ratio=lr)
    elif answer_key and str(answer_key).strip():
        for line in str(answer_key).strip().splitlines():
            if line.strip():
                _add_body(doc, line.strip(), font_name, b_sz, leading_ratio=lr)
    else:
        questions = ws.get("questions", [])
        for i, q in enumerate(questions, 1):
            if isinstance(q, dict):
                answer = q.get("answer") or q.get("correct_answer") or q.get("correct") or ""
                if isinstance(answer, list):
                    answer = ", ".join(str(a) for a in answer)
                _add_body(doc, f"{i}. {answer}", font_name, b_sz, leading_ratio=lr)


# ── Exam DOCX generator ─────────────────────────────────────────────

def _exam_title_fallback(sec_type: str) -> str:
    label = t(sec_type)
    if label == sec_type:
        label = str(sec_type or "").replace("_", " ").title()
    part_title = t("quick_exam_part_title", section=label)
    return part_title if part_title != "quick_exam_part_title" else label


def _format_exam_question_text(sec_type: str, q) -> str:
    if sec_type == "multiple_choice" and isinstance(q, dict):
        return _strip_leading_enum(q.get("stem") or q.get("text") or "")
    if sec_type == "sentence_transformation" and isinstance(q, dict):
        original = _strip_leading_enum(q.get("original") or "")
        prompt = _strip_leading_enum(q.get("prompt") or "")
        return f"{original} ({prompt})" if prompt else original
    if sec_type == "vocabulary" and isinstance(q, dict):
        word = _sentence_case_fragment(q.get("word") or "")
        task = _strip_leading_enum(q.get("task") or "")
        return f"{word}: {task}" if task else word
    if sec_type == "matching" and isinstance(q, dict):
        left_text = _extract_mapping_side(q.get("left", ""), prefer="key")
        right_text = _extract_mapping_side(q.get("right", ""), prefer="value")
        return f"{left_text} ↔ {right_text}" if right_text else left_text
    if isinstance(q, dict):
        return _strip_leading_enum(q.get("text") or q.get("stem") or q.get("question") or q.get("sentence") or "")
    return _strip_leading_enum(q)


def _format_exam_answer_text(ans) -> str:
    if isinstance(ans, dict):
        if "left" in ans and "right" in ans:
            return f"{_extract_mapping_side(ans.get('left', ''), prefer='key')} -> {_extract_mapping_side(ans.get('right', ''), prefer='value')}"
        if "word" in ans and "answer" in ans:
            return f"{_sentence_case_fragment(ans.get('word', ''))}: {_strip_leading_enum(ans.get('answer', ''))}"
        return _strip_leading_enum(ans.get("answer") or ans.get("text") or str(ans))
    return _strip_leading_enum(ans)


def _add_response_lines(doc: Document, font_name: str, b_sz: float, line_count: int, lr: float = 1.15):
    for _ in range(max(1, line_count)):
        _add_body(doc, "_" * 50, font_name, b_sz, leading_ratio=lr)


def _render_exam_matching_docx(doc, questions: list, font_name: str, b_sz: float):
    left_items = []
    right_items = []
    for q in questions or []:
        if not isinstance(q, dict):
            continue
        left_items.append(_extract_mapping_side(q.get("left", ""), prefer="key"))
        right_items.append(_extract_mapping_side(q.get("right", ""), prefer="value"))

    if not left_items:
        return

    import random
    rng = random.Random("|".join(left_items + right_items))
    shuffled_right = right_items[:]
    rng.shuffle(shuffled_right)

    max_len = max(len(left_items), len(shuffled_right))
    table = doc.add_table(rows=max_len, cols=3)
    table.autofit = True

    for i in range(max_len):
        cells = table.rows[i].cells
        left_text = f"{i+1}. {left_items[i]}" if i < len(left_items) else ""
        box_text = "[   ]" if i < len(left_items) else ""
        right_text = f"{chr(97+i)}) {shuffled_right[i]}" if i < len(shuffled_right) else ""

        _set_cell_font(cells[0], left_text, font_name, b_sz)
        p_box = _set_cell_font(cells[1], box_text, font_name, b_sz)
        p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_font(cells[2], right_text, font_name, b_sz)

        for cell in cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tcPr.append(borders)

    doc.add_paragraph()


def _render_exam_true_false_docx(doc, questions: list, font_name: str, b_sz: float):
    if not questions:
        return

    table = doc.add_table(rows=len(questions), cols=2)
    table.autofit = False
    _set_table_col_width(table, 0, 12.8)
    _set_table_col_width(table, 1, 3.6)
    tf_label = _true_false_mark_label()

    for idx, q in enumerate(questions, 1):
        cells = table.rows[idx - 1].cells
        text = _format_exam_question_text("true_false", q)
        _set_cell_font(cells[0], f"{idx}. {text}", font_name, b_sz)
        p = _set_cell_font(cells[1], tf_label, font_name, b_sz)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tcPr.append(borders)

    doc.add_paragraph()


def _render_exam_section_docx(
    doc: Document,
    sec: dict,
    font_name: str,
    sec_sz: float,
    b_sz: float,
    lr: float = 1.15,
):
    sec_type = str(sec.get("type") or "").strip()
    sec_title = str(sec.get("title") or _exam_title_fallback(sec_type)).strip()

    _add_heading(doc, sec_title, font_name, sec_sz, color=_TEXT, leading_ratio=lr)

    instructions = str(sec.get("instructions") or "").strip()
    if instructions:
        _add_body(doc, instructions, font_name, b_sz, leading_ratio=lr)

    source_text = str(sec.get("source_text") or "").strip()
    if source_text:
        _add_body(doc, _strip_leading_enum(source_text), font_name, b_sz, leading_ratio=lr)

    questions = sec.get("questions") or []

    if sec_type in ("word_search_vocab", "word_search") and sec.get("vocabulary_bank"):
        _render_vocab_bank_docx(doc, sec.get("vocabulary_bank") or [], font_name, sec_sz, b_sz, lr)

    if sec_type == "multiple_choice":
        for idx, q in enumerate(questions, 1):
            if isinstance(q, dict):
                stem = _format_exam_question_text(sec_type, q)
                _add_body(doc, f"{idx}. {stem}", font_name, b_sz, bold=True, leading_ratio=lr)
                for opt_idx, opt in enumerate(q.get("options", [])):
                    letter = chr(65 + opt_idx)
                    _add_body(doc, f"    {letter}) {_strip_leading_enum(opt)}", font_name, b_sz, leading_ratio=lr)
            else:
                _add_body(doc, f"{idx}. {_format_exam_question_text(sec_type, q)}", font_name, b_sz, leading_ratio=lr)
    elif sec_type == "matching":
        _render_exam_matching_docx(doc, questions, font_name, b_sz)
    elif sec_type == "true_false":
        _render_exam_true_false_docx(doc, questions, font_name, b_sz)
    elif sec_type in ("word_search_vocab", "word_search"):
        _render_word_search_docx(doc, sec, font_name, sec_sz, b_sz, lr, show_heading=False)
    elif sec_type == "fill_in_blank":
        for idx, q in enumerate(questions, 1):
            text = _format_exam_question_text(sec_type, q).replace("___", "______________")
            _add_body(doc, f"{idx}. {text}", font_name, b_sz, leading_ratio=lr)
    elif sec_type in (
        "short_answer",
        "reading_comprehension",
        "writing_prompt",
        "problem_solving",
        "equation_solving",
        "table_interpretation",
        "word_problems",
        "data_analysis",
        "classification",
        "process_explanation",
        "hypothesis_and_conclusion",
        "diagram_questions",
        "theory_questions",
        "symbol_identification",
        "rhythm_counting",
        "terminology",
        "composer_period_matching",
        "show_your_work",
        "vocabulary",
        "sentence_transformation",
        "error_correction",
    ):
        line_count = 5 if sec_type in ("writing_prompt", "show_your_work") else 2
        for idx, q in enumerate(questions, 1):
            _add_body(doc, f"{idx}. {_format_exam_question_text(sec_type, q)}", font_name, b_sz, leading_ratio=lr)
            _add_response_lines(doc, font_name, b_sz, line_count, lr)
    else:
        for idx, q in enumerate(questions, 1):
            _add_body(doc, f"{idx}. {_format_exam_question_text(sec_type, q)}", font_name, b_sz, leading_ratio=lr)

    doc.add_paragraph()


def _render_exam_answer_key_docx(doc: Document, exam_data: dict, answer_key: dict, font_name: str, sec_sz: float, b_sz: float, lr: float = 1.15):
    exam_sections = exam_data.get("sections") or []
    ak_sections = answer_key.get("sections") or []

    for idx, ak_sec in enumerate(ak_sections):
        fallback_type = exam_sections[idx].get("type", "") if idx < len(exam_sections) else ""
        sec_title = str(ak_sec.get("title") or _exam_title_fallback(fallback_type)).strip()
        _add_heading(doc, sec_title, font_name, sec_sz, color=_TEXT, leading_ratio=lr)
        for answer_idx, ans in enumerate(ak_sec.get("answers", []), 1):
            _add_body(doc, f"{answer_idx}. {_format_exam_answer_text(ans)}", font_name, b_sz, leading_ratio=lr)
        doc.add_paragraph()


def generate_docx_exam(
    exam_data: dict,
    answer_key: dict | None = None,
    *,
    student_only: bool = True,
) -> bytes:
    """
    Generate a Word document from an exam dict.
    The teacher version appends an answer key section.
    """
    from helpers.branding import get_user_branding
    from helpers.font_manager import get_docx_font_name, get_font_sizes

    branding = get_user_branding()
    font_key = branding.get("branding_font", "dejavu")
    size_key = branding.get("branding_font_size", "standard")

    font_name = get_docx_font_name(font_key)
    sz = get_font_sizes(size_key)
    t_sz, sec_sz, b_sz = sz["title"], sz["section"], sz["body"]
    lr = sz.get("leading_ratio", 1.15)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(b_sz)

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    _add_logo_to_doc(doc, branding)

    brand_name = branding.get("brand_name", "").strip()
    department = branding.get("department", "").strip()
    header_enabled = branding.get("header_enabled", False)

    if header_enabled and brand_name:
        _add_heading(doc, brand_name, font_name, sec_sz + 1, color=_PRIMARY,
                     align=WD_ALIGN_PARAGRAPH.CENTER, leading_ratio=lr)
    if header_enabled and department:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(department)
        _apply_font(run, font_name, b_sz, color=_TEXT_MUTED)
        _set_line_spacing(p, lr)

    title = str(exam_data.get("title") or t("quick_exam_generic_exam_title")).strip()
    if not student_only:
        title = f"{title} - {t('ws_answer_key')}"
    _add_heading(doc, title, font_name, t_sz, color=_PRIMARY, leading_ratio=lr)

    if branding.get("header_style") == "school" and header_enabled:
        table = doc.add_table(rows=1, cols=3)
        table.autofit = True
        cells = table.rows[0].cells
        _set_cell_font(cells[0], f"{t('student_name_label')}: _________________________________", font_name, b_sz, bold=True)
        _set_cell_font(cells[1], f"{t('class_label')}: __________", font_name, b_sz, bold=True)
        _set_cell_font(cells[2], f"{t('date_label')}: __________", font_name, b_sz, bold=True)
        for cell in cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tcPr.append(borders)
        _add_separator(doc)

    instructions = str(exam_data.get("instructions") or "").strip()
    if instructions:
        _add_heading(doc, t("ws_instructions") if t("ws_instructions") != "ws_instructions" else "Instructions",
                     font_name, sec_sz, color=_TEXT, leading_ratio=lr)
        _add_body(doc, instructions, font_name, b_sz, leading_ratio=lr)

    for sec in exam_data.get("sections", []):
        _render_exam_section_docx(doc, sec, font_name, sec_sz, b_sz, lr)

    if not student_only and answer_key:
        _add_separator(doc)
        _add_heading(doc, t("ws_answer_key") if t("ws_answer_key") != "ws_answer_key" else "Answer Key",
                     font_name, t_sz, color=_PRIMARY, leading_ratio=lr)
        _render_exam_answer_key_docx(doc, exam_data, answer_key, font_name, sec_sz, b_sz, lr)

    _add_footer_with_image(doc, branding, font_name)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Lesson plan DOCX generator ──────────────────────────────────────

def generate_docx_lesson_plan(
    plan: dict,
    *,
    subject: str = "",
    topic: str = "",
    learner_stage: str = "",
    level_or_band: str = "",
    lesson_purpose: str = "",
) -> bytes:
    """
    Generate a Word document from a lesson plan dict.
    Mirrors the PDF layout: Overview → Lesson Flow → Teacher Notes → Materials.
    Returns raw .docx bytes.
    """
    from helpers.branding import get_user_branding
    from helpers.font_manager import get_docx_font_name, get_font_sizes

    branding = get_user_branding()
    font_key = branding.get("branding_font", "dejavu")
    size_key = branding.get("branding_font_size", "standard")

    font_name = get_docx_font_name(font_key)
    sz = get_font_sizes(size_key)
    t_sz, sec_sz, b_sz = sz["title"], sz["section"], sz["body"]
    lr = sz.get("leading_ratio", 1.15)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(b_sz)

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Logo
    _add_logo_to_doc(doc, branding)

    # Header branding
    brand_name = branding.get("brand_name", "").strip()
    if branding.get("header_enabled") and brand_name:
        _add_heading(doc, brand_name, font_name, sec_sz + 1, color=_PRIMARY,
                     align=WD_ALIGN_PARAGRAPH.CENTER, leading_ratio=lr)

    # Title
    title = str(plan.get("title") or t("untitled_plan")).strip()
    _add_heading(doc, title, font_name, t_sz, color=_PRIMARY,
                 align=WD_ALIGN_PARAGRAPH.CENTER, leading_ratio=lr)

    # Meta info line (labels bold, values normal — matches PDF)
    meta_pairs = []  # list of (label, value)
    if subject:
        meta_pairs.append((f"{t('subject_label')}:", subject))
    if topic:
        meta_pairs.append((f"{t('topic_label')}:", topic))
    if learner_stage:
        meta_pairs.append((f"{t('learner_stage')}:", learner_stage))
    if level_or_band:
        meta_pairs.append((f"{t('level_or_band')}:", level_or_band))
    if lesson_purpose:
        lp_label = t('lesson_purpose') if t('lesson_purpose') != 'lesson_purpose' else 'Purpose'
        meta_pairs.append((f"{lp_label}:", lesson_purpose))
    mat_lang = str(plan.get("student_material_language") or "").strip().upper()
    if mat_lang:
        ml_label = t('student_material_language')
        if ml_label == 'student_material_language':
            ml_label = t('material_language')
        meta_pairs.append((f"{ml_label}:", mat_lang))
    if meta_pairs:
        p = doc.add_paragraph()
        meta_sz = b_sz - 1
        for i, (label, value) in enumerate(meta_pairs):
            if i > 0:
                sep = p.add_run(" | ")
                _apply_font(sep, font_name, meta_sz)
            run_label = p.add_run(f"{label} ")
            _apply_font(run_label, font_name, meta_sz, bold=True)
            run_value = p.add_run(value)
            _apply_font(run_value, font_name, meta_sz)
        p.paragraph_format.space_after = Pt(3)
        _set_line_spacing(p, lr)

    # ── LESSON OVERVIEW ──────────────────────────────────────
    objective = str(plan.get("objective") or "").strip()
    success_criteria = plan.get("success_criteria") or []

    if objective or success_criteria:
        _add_heading(doc, (t("lesson_overview") if t("lesson_overview") != "lesson_overview" else "LESSON OVERVIEW").upper(),
                     font_name, sec_sz, color=_PRIMARY, leading_ratio=lr)

        if objective:
            _add_body(doc, f"{t('lesson_objective') if t('lesson_objective') != 'lesson_objective' else 'Objective'}:", font_name, b_sz, bold=True, leading_ratio=lr)
            _add_body(doc, objective, font_name, b_sz, leading_ratio=lr)

        if success_criteria:
            _add_body(doc, f"{t('success_criteria') if t('success_criteria') != 'success_criteria' else 'Success Criteria'}:", font_name, b_sz, bold=True, leading_ratio=lr)
            _render_list_or_text(doc, success_criteria, font_name, b_sz, lr)

    # ── LESSON FLOW ──────────────────────────────────────────
    flow_sections = [
        ("warm_up", plan.get("warm_up", [])),
        ("main_activity", plan.get("main_activity", [])),
        ("guided_practice", plan.get("guided_practice", [])),
        ("freer_task", plan.get("freer_task", [])),
        ("wrap_up", plan.get("wrap_up", [])),
    ]
    flow_has_content = any(_has_value(v) for _, v in flow_sections)

    if flow_has_content:
        _add_heading(doc, (t("lesson_flow") if t("lesson_flow") != "lesson_flow" else "LESSON FLOW").upper(),
                     font_name, sec_sz, color=_PRIMARY, leading_ratio=lr)

        for key, value in flow_sections:
            if not _has_value(value):
                continue
            label = t(key) if t(key) != key else _fallback_label(key)
            _add_body(doc, label, font_name, b_sz, bold=True, leading_ratio=lr)
            _render_list_or_text(doc, value, font_name, b_sz, lr)

    # ── TEACHER NOTES ────────────────────────────────────────
    teacher_sections = [
        ("core_examples", plan.get("core_examples", [])),
        ("practice_questions", plan.get("practice_questions", [])),
        ("teacher_moves", plan.get("teacher_moves", [])),
        ("extension_task", plan.get("extension_task", "")),
        ("optional_homework", plan.get("homework", "")),
    ]
    teacher_has_content = any(_has_value(v) for _, v in teacher_sections)

    if teacher_has_content:
        _add_heading(doc, (t("teacher_notes") if t("teacher_notes") != "teacher_notes" else "TEACHER NOTES").upper(),
                     font_name, sec_sz, color=_PRIMARY, leading_ratio=lr)

        for key, value in teacher_sections:
            if not _has_value(value):
                continue
            label = t(key) if t(key) != key else _fallback_label(key)
            _add_body(doc, label, font_name, b_sz, bold=True, leading_ratio=lr)
            _render_list_or_text(doc, value, font_name, b_sz, lr)

    # ── LESSON MATERIALS ─────────────────────────────────────
    cm = plan.get("core_material", {}) or {}
    material_items = [
        ("target_vocabulary", cm.get("target_vocabulary")),
        ("language_frames", cm.get("language_frames")),
        ("pre_task_questions", cm.get("pre_task_questions")),
        ("reading_passage", plan.get("reading_passage")),
        ("listening_script", plan.get("listening_script")),
        ("gist_questions", cm.get("gist_questions")),
        ("detail_questions", cm.get("detail_questions")),
        ("worked_example", cm.get("worked_example")),
        ("independent_practice", cm.get("independent_practice")),
        ("common_error_alert", cm.get("common_error_alert")),
        ("concept_explanation", cm.get("concept_explanation")),
        ("real_life_application", cm.get("real_life_application")),
        ("strategy_steps", cm.get("strategy_steps")),
        ("performance_goal", cm.get("performance_goal")),
        ("materials_needed", cm.get("materials_needed")),
        ("timing_guide", cm.get("timing_guide")),
        ("expected_output", cm.get("expected_output")),
        ("differentiation", cm.get("differentiation")),
        ("assessment_check", cm.get("assessment_check")),
        ("student_checklist", cm.get("student_checklist")),
        ("key_concept", cm.get("key_concept")),
        ("guided_problem_set", cm.get("guided_problem_set")),
        ("independent_problem_set", cm.get("independent_problem_set")),
        ("challenge_problem", cm.get("challenge_problem")),
        ("answer_key", cm.get("answer_key")),
        ("phenomenon_prompt", cm.get("phenomenon_prompt")),
        ("prediction_task", cm.get("prediction_task")),
        ("observation_task", cm.get("observation_task")),
        ("evidence_questions", cm.get("evidence_questions")),
        ("misconception_alert", cm.get("misconception_alert")),
        ("technical_focus", cm.get("technical_focus")),
        ("practice_pattern", cm.get("practice_pattern")),
        ("teacher_model", cm.get("teacher_model")),
        ("strategy_name", cm.get("strategy_name")),
        ("model_scenario", cm.get("model_scenario")),
        ("student_action_plan", cm.get("student_action_plan")),
        ("post_task", cm.get("post_task")),
    ]
    material_items = [(k, v) for k, v in material_items if _has_value(v)]

    if material_items:
        _add_heading(doc, (t("lesson_materials") if t("lesson_materials") != "lesson_materials" else "LESSON MATERIALS").upper(),
                     font_name, sec_sz, color=_PRIMARY, leading_ratio=lr)

        for key, value in material_items:
            label = t(key) if t(key) != key else _fallback_label(key)
            _add_body(doc, label, font_name, b_sz, bold=True, leading_ratio=lr)
            _render_list_or_text(doc, value, font_name, b_sz, lr)

    # Footer (with optional image)
    _add_footer_with_image(doc, branding, font_name)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _has_value(value) -> bool:
    """Check if a value has meaningful content."""
    if value is None:
        return False
    if isinstance(value, str):
        return bool(value.strip())
    if isinstance(value, list):
        return any(str(x).strip() for x in value)
    if isinstance(value, dict):
        return bool(value)
    return bool(value)


def _fallback_label(key: str) -> str:
    """Convert a snake_case key to a Title Case label."""
    return key.replace("_", " ").title()


def _render_list_or_text(doc, value, font_name: str, b_sz: float, lr: float = 1.15):
    """Render a value that could be a list of strings, a single string, or a list of dicts."""
    if isinstance(value, list):
        for item in value:
            if isinstance(item, dict):
                text = item.get("text") or item.get("description") or item.get("item") or str(item)
                _add_body(doc, f"• {text}", font_name, b_sz, leading_ratio=lr)
            elif isinstance(item, str) and item.strip():
                _add_body(doc, f"• {item}", font_name, b_sz, leading_ratio=lr)
    elif isinstance(value, str) and value.strip():
        for line in value.strip().splitlines():
            if line.strip():
                _add_body(doc, line.strip(), font_name, b_sz, leading_ratio=lr)
