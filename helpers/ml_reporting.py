from __future__ import annotations

from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from core.database import load_profile_row
from core.state import get_current_user_id


def _normalize_lang(lang: str | None) -> str:
    value = str(lang or "").strip().lower()
    return value if value in {"en", "es", "tr"} else "en"


def resolve_report_language(user_id: str | None = None, preferred: str | None = None) -> str:
    if preferred:
        return _normalize_lang(preferred)
    safe_user_id = str(user_id or get_current_user_id() or "").strip()
    if safe_user_id:
        profile = load_profile_row(safe_user_id)
        profile_lang = str((profile or {}).get("preferred_ui_language") or "").strip().lower()
        if profile_lang in {"en", "es", "tr"}:
            return profile_lang
    try:
        import streamlit as st

        return _normalize_lang(st.session_state.get("ui_lang", "en"))
    except Exception:
        return "en"


def _localized(value: Any, lang: str) -> str:
    if isinstance(value, dict):
        return str(value.get(lang) or value.get("en") or next(iter(value.values()), ""))
    return str(value or "")


_COPY = {
    "es": {
        "report_title": "Classio Admin ML Report: {model_name} ({scope_label})",
        "report_subtitle": "Reporte generado el {generated_at}. Fuente de datos: {source_label}.",
        "source_label": "datos operativos reales de Classio",
        "single_scope": "Single",
        "multi_scope": "Aggregate",
        "problem_heading": "1. Definición del Problema",
        "tool_heading": "2. Selección de la Herramienta",
        "data_heading": "3. Preparación de Datos",
        "features_heading": "4. Selección de Variables y Características",
        "config_heading": "5. Configuración del Modelo",
        "validation_heading": "6. Entrenamiento y Validación del Modelo",
        "analysis_heading": "7. Análisis de Resultados y Conclusiones",
        "summary_indicator": "Indicador",
        "summary_value": "Valor",
        "execution_mode": "Modo de ejecución",
        "live_mode": "Classio live",
        "total_samples": "Total de muestras",
        "train_count": "Entrenamiento",
        "test_count": "Prueba",
        "positive_rate": "Tasa positiva",
        "metrics_chart_caption": "Figura 3. Métricas de validación del modelo.",
        "distribution_chart_caption": "Figura 1. Distribución de muestras por tipo de recurso.",
        "features_chart_caption": "Figura 2. Variables con mayor influencia en el modelo.",
        "confusion_chart_caption": "Figura 4. Matriz de confusión en el conjunto de prueba.",
        "segments_chart_caption": "Figura 5. Peso sugerido por segmento.",
        "metric": "Métrica",
        "result": "Resultado",
        "ml_weight": "Peso ML sugerido",
        "maturity": "Madurez del modelo",
        "metric_interpretation_intro": "Interpretación automática de métricas:",
        "limitations_heading": "Limitaciones y mejoras recomendadas:",
        "segment_intro": "Seguimiento por segmento:",
        "segment": "Segmento",
        "samples": "Muestras",
        "f1": "F1",
        "recommended_weight": "Peso recomendado",
        "kind_distribution_low": "La distribución por tipo de recurso es todavía reducida, así que conviene interpretar los resultados con cautela.",
        "kind_distribution_balanced": "La distribución por tipo de recurso es suficientemente diversa para comparar comportamientos entre formatos.",
        "small_sample_warning": "El volumen histórico disponible en esta ejecución es reducido, por lo que las métricas deben interpretarse como señal inicial y no como validación definitiva.",
        "feature_empty": "No hubo suficientes muestras validadas para identificar variables influyentes con estabilidad estadística. En esta ejecución conviene ampliar la evidencia antes de interpretar pesos.",
        "feature_intro": "La descripción de variables se genera directamente desde los pesos aprendidos del modelo. En esta corrida, las señales con mayor influencia fueron: {feature_list}.",
        "imbalance_warning": "El dataset está altamente desbalanceado. La accuracy no debe interpretarse como métrica principal; en este contexto pesan más Precision, Recall, F1-score y ROC-AUC.",
        "accuracy_paradox": "La accuracy de {accuracy} es probablemente engañosa porque la tasa positiva es de solo {positive_rate}. Un clasificador que prediga casi todo como negativo puede seguir pareciendo preciso.",
        "accuracy_normal": "La accuracy observada fue de {accuracy} y su lectura es consistente con el balance actual del conjunto.",
        "precision_zero": "La Precision es 0.0%, por lo que el modelo no está generando aciertos positivos útiles en validación.",
        "precision_low": "La Precision sigue baja, lo que indica una proporción elevada de falsos positivos cuando el modelo intenta señalar casos positivos.",
        "precision_good": "La Precision ya muestra que una parte razonable de las predicciones positivas corresponde a casos realmente positivos.",
        "recall_zero": "El Recall es 0.0%, señal de que el modelo no está recuperando casos positivos en el conjunto de prueba.",
        "recall_low": "El Recall sigue limitado y deja escapar una parte importante de las oportunidades positivas.",
        "recall_good": "El Recall indica que el modelo sí está capturando una parte útil de los casos positivos.",
        "f1_zero": "El F1-score es 0.0%, por lo que el clasificador no tiene valor predictivo práctico en su estado actual.",
        "f1_low": "El F1-score todavía es bajo y confirma que el equilibrio entre Precision y Recall no es suficiente.",
        "f1_good": "El F1-score muestra un equilibrio razonable entre Precision y Recall para esta etapa.",
        "auc_excellent": "La discriminación del modelo es sobresaliente.",
        "auc_strong": "La discriminación del modelo es sólida y ya apunta a un uso operativo controlado.",
        "auc_moderate": "La discriminación del modelo es moderada: supera el azar, pero todavía requiere mejoras.",
        "auc_weak": "La discriminación del modelo es apenas superior al azar.",
        "auc_bad": "La discriminación del modelo es deficiente y está por debajo de un comportamiento útil.",
        "auc_with_value": "{text} ROC-AUC actual: {roc_auc}.",
        "degenerate_warning": "El modelo cayó en un comportamiento degenerado durante validación: no identificó ejemplos positivos. Antes de considerarlo efectivo, necesita más positivos o técnicas explícitas de manejo de desbalance.",
        "confusion_degenerate": "La matriz de confusión muestra {tn} verdaderos negativos, {fn} falsos negativos, {fp} falsos positivos y {tp} verdaderos positivos. El patrón dominante es la omisión de positivos, no una buena capacidad de clasificación.",
        "confusion_normal": "La matriz de confusión registra {tn} verdaderos negativos, {tp} verdaderos positivos, {fp} falsos positivos y {fn} falsos negativos, lo que ayuda a identificar si el error dominante proviene de sobre-predecir o de dejar pasar oportunidades relevantes.",
        "maturity_experimental": "Experimental",
        "maturity_prototype": "Prototype",
        "maturity_candidate": "Operational Candidate",
        "maturity_ready": "Production Ready",
        "maturity_experimental_text": "El modelo debe considerarse experimental. La evidencia actual no respalda un despliegue confiable como señal principal de decisión.",
        "maturity_prototype_text": "El modelo se encuentra en etapa de prototipo: ya muestra alguna señal útil, pero aún no debe presentarse como un componente maduro.",
        "maturity_candidate_text": "El modelo puede evaluarse como candidato operativo, pero todavía conviene monitorizar estabilidad y errores antes de ampliar su uso.",
        "maturity_ready_text": "El modelo puede considerarse listo para producción porque combina discriminación alta y equilibrio predictivo fuerte.",
        "conclusion_degenerate": "Conclusión: el modelo se clasifica actualmente como Experimental. Aunque puede mostrar una accuracy alta, la validación indica que no identifica casos positivos y por tanto no aporta valor predictivo práctico todavía.",
        "conclusion_normal": "Conclusión: el modelo se clasifica actualmente como {maturity}. Con un ROC-AUC de {roc_auc} y un F1-score de {f1}, su lectura debe centrarse en la capacidad real de discriminar y no en plantillas genéricas.",
        "analysis_sample_line": "El pipeline evaluó {sample_count} muestras y alcanzó un F1-score de {f1}.",
        "analysis_positive_line": "La tasa positiva observada fue de {positive_rate}.",
        "analysis_maturity_line": "El nivel de madurez asignado automáticamente al modelo es {maturity}.",
        "objective_selector_heading": "Lectura del selector pedagogico",
        "objective_selector_line": "La capa que decide entre siguiente tema, repaso o hueco pendiente evaluo {sample_count} muestras, con F1-score de {f1} y ROC-AUC de {roc_auc}.",
        "generic_limit_1": "El desempeño depende del volumen histórico disponible y del nivel de adopción actual del flujo evaluado.",
        "generic_limit_2": "Parte del target sigue aproximado a partir de eventos intermedios; todavía puede afinarse con señales más cercanas al resultado real de uso o impacto pedagógico.",
        "generic_limit_imbalance": "El desbalance de clases requiere acciones concretas como reponderación, más ejemplos positivos o umbrales calibrados antes de usar la accuracy como argumento de calidad.",
        "generic_limit_degenerate": "La validación muestra un clasificador degenerado para positivos. La prioridad inmediata es corregir ese fallo antes de afirmar utilidad operativa.",
        "generic_limit_next": "La siguiente iteración debería incorporar más señales de conversión y resultados posteriores para cerrar mejor el ciclo de feedback.",
    },
    "en": {
        "report_title": "Classio Admin ML Report: {model_name} ({scope_label})",
        "report_subtitle": "Generated on {generated_at}. Data source: {source_label}.",
        "source_label": "Classio live operational data",
        "single_scope": "Single",
        "multi_scope": "Aggregate",
        "problem_heading": "1. Problem Definition",
        "tool_heading": "2. Tool Selection",
        "data_heading": "3. Data Preparation",
        "features_heading": "4. Variable and Feature Selection",
        "config_heading": "5. Model Configuration",
        "validation_heading": "6. Model Training and Validation",
        "analysis_heading": "7. Results Analysis and Conclusions",
        "summary_indicator": "Indicator",
        "summary_value": "Value",
        "execution_mode": "Execution mode",
        "live_mode": "Classio live",
        "total_samples": "Total samples",
        "train_count": "Training",
        "test_count": "Test",
        "positive_rate": "Positive rate",
        "metrics_chart_caption": "Figure 3. Model validation metrics.",
        "distribution_chart_caption": "Figure 1. Sample distribution by resource type.",
        "features_chart_caption": "Figure 2. Most influential variables in the model.",
        "confusion_chart_caption": "Figure 4. Confusion matrix on the test set.",
        "segments_chart_caption": "Figure 5. Suggested weight by segment.",
        "metric": "Metric",
        "result": "Result",
        "ml_weight": "Suggested ML weight",
        "maturity": "Model maturity",
        "metric_interpretation_intro": "Automatic metric interpretation:",
        "limitations_heading": "Limitations and recommended improvements:",
        "segment_intro": "Segment follow-up:",
        "segment": "Segment",
        "samples": "Samples",
        "f1": "F1",
        "recommended_weight": "Recommended weight",
        "kind_distribution_low": "Distribution across resource types is still narrow, so results should be interpreted cautiously.",
        "kind_distribution_balanced": "Distribution across resource types is diverse enough to compare behaviour across formats.",
        "small_sample_warning": "The historical volume available in this run is limited, so metrics should be read as an initial signal rather than definitive validation.",
        "feature_empty": "There were not enough validated samples to identify influential variables with statistical stability. More evidence is needed before interpreting weights.",
        "feature_intro": "The variable description is generated directly from the model's learned weights. In this run, the most influential signals were: {feature_list}.",
        "imbalance_warning": "The dataset is highly imbalanced. Accuracy should not be treated as the main metric; Precision, Recall, F1-score and ROC-AUC are more informative here.",
        "accuracy_paradox": "The {accuracy} accuracy is likely misleading because the positive rate is only {positive_rate}. A classifier that predicts almost everything as negative can still appear accurate.",
        "accuracy_normal": "Observed accuracy was {accuracy}, and its reading is consistent with the current class balance.",
        "precision_zero": "Precision is 0.0%, so the model is not producing useful positive hits in validation.",
        "precision_low": "Precision remains low, indicating a high proportion of false positives when the model tries to flag positive cases.",
        "precision_good": "Precision already shows that a reasonable share of positive predictions corresponds to actual positive cases.",
        "recall_zero": "Recall is 0.0%, which means the model is not recovering positive cases in the test set.",
        "recall_low": "Recall remains limited and misses an important share of positive opportunities.",
        "recall_good": "Recall indicates that the model is capturing a useful share of positive cases.",
        "f1_zero": "F1-score is 0.0%, so the classifier currently has no practical predictive value.",
        "f1_low": "F1-score is still low and confirms that the Precision/Recall balance is not sufficient yet.",
        "f1_good": "F1-score shows a reasonable Precision/Recall balance for this stage.",
        "auc_excellent": "Model discrimination is outstanding.",
        "auc_strong": "Model discrimination is strong and already points toward controlled operational use.",
        "auc_moderate": "Model discrimination is moderate: it beats chance, but still needs improvement.",
        "auc_weak": "Model discrimination is only slightly better than random.",
        "auc_bad": "Model discrimination is poor and below a useful operating level.",
        "auc_with_value": "{text} Current ROC-AUC: {roc_auc}.",
        "degenerate_warning": "The model collapsed into degenerate behaviour during validation: it did not identify positive examples. It needs more positives or explicit imbalance handling before being considered effective.",
        "confusion_degenerate": "The confusion matrix shows {tn} true negatives, {fn} false negatives, {fp} false positives and {tp} true positives. The dominant pattern is missed positives, not good classification skill.",
        "confusion_normal": "The confusion matrix records {tn} true negatives, {tp} true positives, {fp} false positives and {fn} false negatives, which helps show whether the dominant error comes from over-predicting or from missing relevant opportunities.",
        "maturity_experimental": "Experimental",
        "maturity_prototype": "Prototype",
        "maturity_candidate": "Operational Candidate",
        "maturity_ready": "Production Ready",
        "maturity_experimental_text": "The model should be considered experimental. Current evidence does not support a reliable deployment as a primary decision signal.",
        "maturity_prototype_text": "The model is still at prototype stage: it shows some useful signal, but should not yet be presented as a mature component.",
        "maturity_candidate_text": "The model can be evaluated as an operational candidate, but stability and error patterns should still be monitored before wider use.",
        "maturity_ready_text": "The model can be considered production ready because it combines high discrimination with strong predictive balance.",
        "conclusion_degenerate": "Conclusion: the model is currently classified as Experimental. Even if accuracy looks high, validation shows that it is not identifying positive cases and therefore does not yet provide practical predictive value.",
        "conclusion_normal": "Conclusion: the model is currently classified as {maturity}. With a ROC-AUC of {roc_auc} and an F1-score of {f1}, interpretation should focus on real discrimination capacity rather than generic templates.",
        "analysis_sample_line": "The pipeline evaluated {sample_count} samples and reached an F1-score of {f1}.",
        "analysis_positive_line": "Observed positive rate was {positive_rate}.",
        "analysis_maturity_line": "The automatically assigned maturity level is {maturity}.",
        "objective_selector_heading": "Pedagogical objective selector",
        "objective_selector_line": "The layer that chooses between next topic, review or pending gap evaluated {sample_count} samples, with an F1-score of {f1} and ROC-AUC of {roc_auc}.",
        "generic_limit_1": "Performance depends on the available historical volume and on adoption of the evaluated workflow.",
        "generic_limit_2": "Part of the target is still approximated from intermediate events; it can be refined with signals closer to real usage or pedagogical impact.",
        "generic_limit_imbalance": "Class imbalance requires explicit actions such as reweighting, more positive examples or calibrated thresholds before accuracy can be treated as a quality argument.",
        "generic_limit_degenerate": "Validation shows a degenerate classifier for positives. Fixing that failure is the immediate priority before claiming operational usefulness.",
        "generic_limit_next": "The next iteration should incorporate more conversion signals and downstream outcomes to close the feedback loop better.",
    },
    "tr": {
        "report_title": "Classio Admin ML Report: {model_name} ({scope_label})",
        "report_subtitle": "{generated_at} tarihinde olusturuldu. Veri kaynagi: {source_label}.",
        "source_label": "Classio canli operasyon verisi",
        "single_scope": "Tekil",
        "multi_scope": "Toplu",
        "problem_heading": "1. Problem Tanimi",
        "tool_heading": "2. Arac Secimi",
        "data_heading": "3. Veri Hazirligi",
        "features_heading": "4. Degisken ve Ozellik Secimi",
        "config_heading": "5. Model Yapilandirmasi",
        "validation_heading": "6. Model Egitimi ve Dogrulama",
        "analysis_heading": "7. Sonuc Analizi ve Cikarimlar",
        "summary_indicator": "Gosterge",
        "summary_value": "Deger",
        "execution_mode": "Calisma modu",
        "live_mode": "Classio canli",
        "total_samples": "Toplam ornek",
        "train_count": "Egitim",
        "test_count": "Test",
        "positive_rate": "Pozitif oran",
        "metrics_chart_caption": "Sekil 3. Model dogrulama metrikleri.",
        "distribution_chart_caption": "Sekil 1. Kaynak turune gore ornek dagilimi.",
        "features_chart_caption": "Sekil 2. Modelde en etkili degiskenler.",
        "confusion_chart_caption": "Sekil 4. Test kumesindeki karisiklik matrisi.",
        "segments_chart_caption": "Sekil 5. Segmente gore onerilen agirlik.",
        "metric": "Metrik",
        "result": "Sonuc",
        "ml_weight": "Onerilen ML agirligi",
        "maturity": "Model olgunlugu",
        "metric_interpretation_intro": "Metriklerin otomatik yorumu:",
        "limitations_heading": "Sinirlar ve onerilen gelistirmeler:",
        "segment_intro": "Segment takibi:",
        "segment": "Segment",
        "samples": "Ornekler",
        "f1": "F1",
        "recommended_weight": "Onerilen agirlik",
        "kind_distribution_low": "Kaynak turleri arasindaki dagilim hala dar; bu nedenle sonuclar dikkatli yorumlanmalidir.",
        "kind_distribution_balanced": "Kaynak turleri arasindaki dagilim, formatlar arasinda davranis karsilastirmasi yapmak icin yeterince cesitlidir.",
        "small_sample_warning": "Bu calismadaki tarihsel hacim sinirli oldugu icin metrikler kesin dogrulama olarak degil, ilk sinyal olarak okunmalidir.",
        "feature_empty": "Istatistiksel olarak istikrarli etkili degiskenleri belirlemek icin yeterli dogrulanmis ornek yoktu. Agirliklari yorumlamadan once daha fazla kanit gereklidir.",
        "feature_intro": "Degisken aciklamasi dogrudan modelin ogrendigi agirliklardan uretilir. Bu kosuda en etkili sinyaller sunlardi: {feature_list}.",
        "imbalance_warning": "Veri kumesi ciddi sekilde dengesiz. Accuracy ana metrik olarak yorumlanmamalidir; Precision, Recall, F1-score ve ROC-AUC burada daha anlamlidir.",
        "accuracy_paradox": "{accuracy} accuracy degeri muhtemelen yanilticidir cunku pozitif oran yalnizca {positive_rate}. Neredeyse her seyi negatif tahmin eden bir siniflandirici yine de dogru gorunebilir.",
        "accuracy_normal": "Gozlenen accuracy {accuracy} oldu ve bu deger mevcut sinif dengesiyle uyumludur.",
        "precision_zero": "Precision %0.0, yani model dogrulamada yararli pozitif isabetler uretmiyor.",
        "precision_low": "Precision hala dusuk; model pozitif vakalari isaretlemeye calistiginda yanlis pozitif oraninin yuksek oldugunu gosteriyor.",
        "precision_good": "Precision, pozitif tahminlerin makul bir kisminin gercekten pozitif vakalara karsilik geldigini gosteriyor.",
        "recall_zero": "Recall %0.0; model test kumesindeki pozitif vakalari yakalayamiyor.",
        "recall_low": "Recall hala sinirli ve pozitif firsatlarin onemli bir kismini kaciriyor.",
        "recall_good": "Recall, modelin pozitif vakalarin faydali bir kismini yakaladigini gosteriyor.",
        "f1_zero": "F1-score %0.0; siniflandiricinin su an pratik tahmin degeri yok.",
        "f1_low": "F1-score hala dusuk ve Precision/Recall dengesinin yeterli olmadigini dogruluyor.",
        "f1_good": "F1-score bu asama icin makul bir Precision/Recall dengesi gosteriyor.",
        "auc_excellent": "Model ayristirma kapasitesi oldukca guclu.",
        "auc_strong": "Model ayristirma kapasitesi guclu ve kontrollu operasyonel kullanim ihtimaline isaret ediyor.",
        "auc_moderate": "Model ayristirma kapasitesi orta duzeyde: sansi asiyor ancak hala gelistirme gerekiyor.",
        "auc_weak": "Model ayristirma kapasitesi rastgele tahminden yalnizca biraz daha iyi.",
        "auc_bad": "Model ayristirma kapasitesi zayif ve kullanisli bir duzeyin altinda.",
        "auc_with_value": "{text} Mevcut ROC-AUC: {roc_auc}.",
        "degenerate_warning": "Model dogrulama sirasinda dejeneratif davrandi: pozitif ornekleri belirleyemedi. Etkili sayilmadan once daha fazla pozitif ornege veya acik dengesizlik yontemlerine ihtiyac var.",
        "confusion_degenerate": "Karisiklik matrisi {tn} gercek negatif, {fn} yanlis negatif, {fp} yanlis pozitif ve {tp} gercek pozitif gosteriyor. Baskin desen iyi siniflandirma degil, pozitifleri kacirmak.",
        "confusion_normal": "Karisiklik matrisi {tn} gercek negatif, {tp} gercek pozitif, {fp} yanlis pozitif ve {fn} yanlis negatif kaydediyor; bu da baskin hatanin fazla tahminden mi yoksa onemli firsatlari kacirmaktan mi geldigini gostermeye yardimci olur.",
        "maturity_experimental": "Experimental",
        "maturity_prototype": "Prototype",
        "maturity_candidate": "Operational Candidate",
        "maturity_ready": "Production Ready",
        "maturity_experimental_text": "Model deneysel kabul edilmelidir. Mevcut kanitlar onu birincil karar sinyali olarak guvenilir sekilde kullanmayi desteklemiyor.",
        "maturity_prototype_text": "Model prototip asamasinda: bazi yararli sinyaller gosteriyor ancak henuz olgun bir bilesen olarak sunulmamalidir.",
        "maturity_candidate_text": "Model operasyonel aday olarak degerlendirilebilir; ancak daha genis kullanimdan once istikrar ve hata kaliplari izlenmelidir.",
        "maturity_ready_text": "Model yuksek ayristirma ve guclu tahmin dengesi sundugu icin uretime hazir sayilabilir.",
        "conclusion_degenerate": "Sonuc: model su anda Experimental olarak siniflandiriliyor. Accuracy yuksek gorunse bile dogrulama, modelin pozitif vakalari tanimlamadigini ve bu nedenle henuz pratik tahmin degeri saglamadigini gosteriyor.",
        "conclusion_normal": "Sonuc: model su anda {maturity} olarak siniflandiriliyor. ROC-AUC degeri {roc_auc} ve F1-score degeri {f1}; yorum gercek ayristirma kapasitesine dayanmalidir, genel sablonlara degil.",
        "analysis_sample_line": "Pipeline {sample_count} ornek degerlendirdi ve {f1} F1-score sonucuna ulasti.",
        "analysis_positive_line": "Gozlenen pozitif oran {positive_rate} oldu.",
        "analysis_maturity_line": "Otomatik atanan olgunluk seviyesi {maturity}.",
        "objective_selector_heading": "Pedagojik hedef secici",
        "objective_selector_line": "Sonraki konu, tekrar veya bekleyen bosluk arasinda secim yapan katman {sample_count} ornegi degerlendirdi; F1-score {f1}, ROC-AUC ise {roc_auc}.",
        "generic_limit_1": "Performans mevcut tarihsel hacme ve degerlendirilen akisin benimsenme duzeyine baglidir.",
        "generic_limit_2": "Hedefin bir kismi hala ara olaylardan yaklasik olarak uretiliyor; gercek kullanim veya pedagojik etkiye daha yakin sinyallerle gelistirilebilir.",
        "generic_limit_imbalance": "Sinif dengesizligi, accuracy kalite argumani olarak kullanilmadan once yeniden agirliklandirma, daha fazla pozitif ornek veya ayarlanmis esikler gibi acik adimlar gerektirir.",
        "generic_limit_degenerate": "Dogrulama pozitifler icin dejeneratif bir siniflandirici gosteriyor. Operasyonel fayda iddiasindan once bu sorunun cozulmesi onceliklidir.",
        "generic_limit_next": "Bir sonraki iterasyon, geri bildirim dongusunu daha iyi kapatmak icin daha fazla donusum sinyali ve sonraki sonuclari dahil etmelidir.",
    },
}


def _text(lang: str, key: str, **kwargs) -> str:
    template = _COPY.get(lang, _COPY["en"]).get(key, key)
    return template.format(**kwargs)


def _pct(value: float) -> str:
    return f"{round(float(value or 0.0) * 100, 1)}%"


def _clamp(value: float, low: float = 0.0, high: float = 1.0) -> float:
    return max(low, min(high, float(value)))


def _ensure_heading_style(style, font_name: str, size: int, color: RGBColor, bold: bool = False) -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def _set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_after = Pt(8)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15
    _ensure_heading_style(doc.styles["Heading 1"], "Arial", 16, RGBColor(46, 116, 181))
    _ensure_heading_style(doc.styles["Heading 2"], "Arial", 13, RGBColor(46, 116, 181))
    _ensure_heading_style(doc.styles["Heading 3"], "Arial", 12, RGBColor(31, 77, 120))


def _add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    r.font.size = Pt(23)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(subtitle)
    r2.font.name = "Arial"
    r2._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r2._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(85, 85, 85)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for cell in hdr:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph("")


def _plot_bar(path: Path, title: str, labels: list[str], values: list[float], *, color: str = "#2E74B5", horizontal: bool = False) -> None:
    plt.figure(figsize=(7.0, 4.0))
    if horizontal:
        plt.barh(labels, values, color=color)
    else:
        plt.bar(labels, values, color=color)
    plt.title(title, fontsize=12)
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()


def _plot_confusion(path: Path, metrics: dict[str, float], lang: str) -> None:
    if lang == "es":
        x_labels = ["Pred. no", "Pred. si"]
        y_labels = ["Real no", "Real si"]
        title = "Matriz de confusion"
    elif lang == "tr":
        x_labels = ["Tah. hayir", "Tah. evet"]
        y_labels = ["Gercek hayir", "Gercek evet"]
        title = "Karisiklik matrisi"
    else:
        x_labels = ["Pred. no", "Pred. yes"]
        y_labels = ["Actual no", "Actual yes"]
        title = "Confusion matrix"
    matrix = [
        [metrics.get("tn", 0.0), metrics.get("fp", 0.0)],
        [metrics.get("fn", 0.0), metrics.get("tp", 0.0)],
    ]
    fig, ax = plt.subplots(figsize=(4.2, 4.0))
    ax.imshow(matrix, cmap="Blues")
    ax.set_xticks([0, 1], labels=x_labels)
    ax.set_yticks([0, 1], labels=y_labels)
    ax.set_title(title)
    for row_idx, row in enumerate(matrix):
        for col_idx, value in enumerate(row):
            ax.text(col_idx, row_idx, str(int(round(value))), ha="center", va="center", color="black")
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close(fig)


def build_report_charts(data: dict, out_dir: Path, lang: str) -> dict[str, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    diagnostics = data["diagnostics"]
    metrics = diagnostics["metrics"]
    charts = {
        "metrics": out_dir / "metrics.png",
        "kinds": out_dir / "kinds.png",
        "features": out_dir / "features.png",
        "confusion": out_dir / "confusion.png",
    }
    segment_rows = diagnostics.get("segment_rows") or []
    if segment_rows:
        charts["segments"] = out_dir / "segments.png"
    _plot_bar(
        charts["metrics"],
        "Validation Metrics" if lang == "en" else ("Dogrulama Metrikleri" if lang == "tr" else "Metricas de validacion"),
        ["Accuracy", "Precision", "Recall", "F1", "ROC-AUC"],
        [
            metrics.get("accuracy", 0.0),
            metrics.get("precision", 0.0),
            metrics.get("recall", 0.0),
            metrics.get("f1", 0.0),
            metrics.get("roc_auc", 0.0),
        ],
        color="#2563EB",
    )
    counts_by_kind = diagnostics.get("counts_by_kind") or {}
    _plot_bar(
        charts["kinds"],
        "Sample Distribution by Resource Type" if lang == "en" else ("Kaynak Turune Gore Ornek Dagilimi" if lang == "tr" else "Distribucion de muestras por tipo de recurso"),
        list(counts_by_kind.keys()) or ["no data"],
        list(counts_by_kind.values()) or [0],
        color="#10B981",
    )
    top_features = diagnostics.get("top_features") or [{"name": "bias", "weight": 0.0}]
    _plot_bar(
        charts["features"],
        "Most Influential Variables" if lang == "en" else ("En Etkili Degiskenler" if lang == "tr" else "Variables mas influyentes"),
        [str(item["name"]) for item in reversed(top_features)],
        [float(item["weight"]) for item in reversed(top_features)],
        color="#F59E0B",
        horizontal=True,
    )
    _plot_confusion(charts["confusion"], metrics, lang)
    if segment_rows:
        _plot_bar(
            charts["segments"],
            "Suggested Recalibration by Segment" if lang == "en" else ("Segmente Gore Onerilen Yeniden Kalibrasyon" if lang == "tr" else "Recalibracion recomendada por segmento"),
            [str(item["segment"]) for item in segment_rows],
            [float(item["recommended_blend_weight"]) for item in segment_rows],
            color="#8B5CF6",
        )
    return charts


def _add_picture_with_caption(doc: Document, image_path: Path, caption: str, width: float = 6.1) -> None:
    doc.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(10)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)


def _feature_label(name: str, feature_labels: dict[str, Any], lang: str) -> str:
    value = feature_labels.get(str(name or "").strip())
    if value is None:
        return str(name or "").replace("_", " ")
    return _localized(value, lang)


def assess_maturity_level(metrics: dict[str, float]) -> str:
    roc_auc = float(metrics.get("roc_auc") or 0.0)
    f1 = float(metrics.get("f1") or 0.0)
    if roc_auc > 0.90 and f1 > 0.80:
        return "ready"
    if roc_auc >= 0.75 and f1 >= 0.30:
        return "candidate"
    if roc_auc >= 0.60 and f1 >= 0.10:
        return "prototype"
    return "experimental"


def maturity_label(level: str, lang: str) -> str:
    key_map = {
        "experimental": "maturity_experimental",
        "prototype": "maturity_prototype",
        "candidate": "maturity_candidate",
        "ready": "maturity_ready",
    }
    return _text(lang, key_map.get(level, "maturity_experimental"))


def maturity_text(level: str, lang: str) -> str:
    key_map = {
        "experimental": "maturity_experimental_text",
        "prototype": "maturity_prototype_text",
        "candidate": "maturity_candidate_text",
        "ready": "maturity_ready_text",
    }
    return _text(lang, key_map.get(level, "maturity_experimental_text"))


def build_feature_summary(diagnostics: dict, feature_labels: dict[str, Any], lang: str) -> str:
    top_features = diagnostics.get("top_features") or []
    if not top_features:
        return _text(lang, "feature_empty")
    feature_list = ", ".join(_feature_label(str(item.get("name") or ""), feature_labels, lang) for item in top_features[:5])
    return _text(lang, "feature_intro", feature_list=feature_list)


def build_kind_distribution_commentary(diagnostics: dict, lang: str) -> str:
    counts_by_kind = diagnostics.get("counts_by_kind") or {}
    non_zero_counts = [int(value) for value in counts_by_kind.values() if int(value) > 0]
    if len(non_zero_counts) <= 1 or sum(non_zero_counts) < 8:
        return _text(lang, "kind_distribution_low")
    return _text(lang, "kind_distribution_balanced")


def build_metric_interpretations(diagnostics: dict, lang: str) -> list[str]:
    metrics = diagnostics.get("metrics") or {}
    accuracy = float(metrics.get("accuracy") or 0.0)
    precision = float(metrics.get("precision") or 0.0)
    recall = float(metrics.get("recall") or 0.0)
    f1 = float(metrics.get("f1") or 0.0)
    roc_auc = float(metrics.get("roc_auc") or 0.0)
    positive_rate = float(diagnostics.get("positive_rate") or 0.0)
    imbalance = positive_rate < 0.10
    degenerate = precision == 0.0 and recall == 0.0 and f1 == 0.0 and int(diagnostics.get("test_count") or 0) > 0

    lines: list[str] = []
    if imbalance:
        lines.append(_text(lang, "imbalance_warning"))
    if imbalance and accuracy >= 0.90:
        lines.append(_text(lang, "accuracy_paradox", accuracy=_pct(accuracy), positive_rate=_pct(positive_rate)))
    else:
        lines.append(_text(lang, "accuracy_normal", accuracy=_pct(accuracy)))

    if precision == 0.0:
        lines.append(_text(lang, "precision_zero"))
    elif precision < 0.40:
        lines.append(_text(lang, "precision_low"))
    else:
        lines.append(_text(lang, "precision_good"))

    if recall == 0.0:
        lines.append(_text(lang, "recall_zero"))
    elif recall < 0.40:
        lines.append(_text(lang, "recall_low"))
    else:
        lines.append(_text(lang, "recall_good"))

    if f1 == 0.0:
        lines.append(_text(lang, "f1_zero"))
    elif f1 < 0.30:
        lines.append(_text(lang, "f1_low"))
    else:
        lines.append(_text(lang, "f1_good"))

    if roc_auc > 0.90:
        auc_text = _text(lang, "auc_excellent")
    elif roc_auc >= 0.75:
        auc_text = _text(lang, "auc_strong")
    elif roc_auc >= 0.60:
        auc_text = _text(lang, "auc_moderate")
    elif roc_auc >= 0.50:
        auc_text = _text(lang, "auc_weak")
    else:
        auc_text = _text(lang, "auc_bad")
    lines.append(_text(lang, "auc_with_value", text=auc_text, roc_auc=_pct(roc_auc)))

    if degenerate:
        lines.append(_text(lang, "degenerate_warning"))
    return lines


def build_confusion_interpretation(diagnostics: dict, lang: str) -> str:
    metrics = diagnostics.get("metrics") or {}
    tp = int(round(float(metrics.get("tp") or 0.0)))
    tn = int(round(float(metrics.get("tn") or 0.0)))
    fp = int(round(float(metrics.get("fp") or 0.0)))
    fn = int(round(float(metrics.get("fn") or 0.0)))
    degenerate = float(metrics.get("precision") or 0.0) == 0.0 and float(metrics.get("recall") or 0.0) == 0.0 and float(metrics.get("f1") or 0.0) == 0.0 and int(diagnostics.get("test_count") or 0) > 0
    key = "confusion_degenerate" if degenerate else "confusion_normal"
    return _text(lang, key, tn=tn, fn=fn, fp=fp, tp=tp)


def build_limitations(diagnostics: dict, lang: str, extra_limitations: list[str]) -> list[str]:
    positive_rate = float(diagnostics.get("positive_rate") or 0.0)
    metrics = diagnostics.get("metrics") or {}
    degenerate = float(metrics.get("precision") or 0.0) == 0.0 and float(metrics.get("recall") or 0.0) == 0.0 and float(metrics.get("f1") or 0.0) == 0.0 and int(diagnostics.get("test_count") or 0) > 0
    lines = [
        _text(lang, "generic_limit_1"),
        _text(lang, "generic_limit_2"),
    ]
    if positive_rate < 0.10:
        lines.append(_text(lang, "generic_limit_imbalance"))
    if degenerate:
        lines.append(_text(lang, "generic_limit_degenerate"))
    else:
        lines.append(_text(lang, "generic_limit_next"))
    lines.extend(extra_limitations)
    return lines


def build_final_conclusion(diagnostics: dict, lang: str) -> str:
    metrics = diagnostics.get("metrics") or {}
    level = assess_maturity_level(metrics)
    maturity = maturity_label(level, lang)
    degenerate = float(metrics.get("precision") or 0.0) == 0.0 and float(metrics.get("recall") or 0.0) == 0.0 and float(metrics.get("f1") or 0.0) == 0.0 and int(diagnostics.get("test_count") or 0) > 0
    if degenerate:
        return _text(lang, "conclusion_degenerate")
    return _text(lang, "conclusion_normal", maturity=maturity, roc_auc=_pct(float(metrics.get("roc_auc") or 0.0)), f1=_pct(float(metrics.get("f1") or 0.0)))


def write_classio_ml_report(docx_path: Path, data: dict, charts: dict[str, Path], meta: dict[str, Any], *, lang: str) -> None:
    diagnostics = data["diagnostics"]
    metrics = diagnostics["metrics"]
    snapshot = data["snapshot"]
    sample_count = int(diagnostics.get("sample_count") or 0)
    positive_rate = float(diagnostics.get("positive_rate") or 0.0)
    scope = str(data.get("scope") or "single")
    generated_at = datetime.now().astimezone().strftime("%B %d, %Y")
    level = assess_maturity_level(metrics)
    maturity = maturity_label(level, lang)

    model_name = _localized(meta.get("model_name"), lang)
    scope_label = _text(lang, "multi_scope" if str(meta.get("scope_kind") or "").startswith("multi") or "multi" in scope else "single_scope")
    doc = Document()
    _set_doc_defaults(doc)
    _add_title(
        doc,
        _text(lang, "report_title", model_name=model_name, scope_label=scope_label),
        _text(lang, "report_subtitle", generated_at=generated_at, source_label=_text(lang, "source_label")),
    )

    doc.add_heading(_text(lang, "problem_heading"), level=1)
    doc.add_paragraph(_localized(meta["problem_multi"] if "multi" in scope else meta["problem_single"], lang))

    doc.add_heading(_text(lang, "tool_heading"), level=1)
    doc.add_paragraph(_localized(meta["tool_summary"], lang))

    doc.add_heading(_text(lang, "data_heading"), level=1)
    doc.add_paragraph(_localized(meta["data_multi"] if "multi" in scope else meta["data_single"], lang))
    doc.add_paragraph(_localized(meta["preparation_summary"], lang))
    summary_rows = [
        [_text(lang, "execution_mode"), _text(lang, "live_mode")],
        [_text(lang, "total_samples"), str(sample_count)],
        [_text(lang, "train_count"), str(int(diagnostics.get("train_count") or 0))],
        [_text(lang, "test_count"), str(int(diagnostics.get("test_count") or 0))],
        [_text(lang, "positive_rate"), _pct(positive_rate)],
    ]
    if "multi" in scope:
        summary_rows.insert(1, [_localized(meta.get("aggregate_entity_label"), lang), str(int(snapshot.get(meta.get("aggregate_count_key")) or 0))])
    else:
        summary_rows.insert(1, [_localized(meta.get("single_entity_label"), lang), str(data.get(meta.get("entity_id_key")) or "n/a")])
    _add_table(doc, [_text(lang, "summary_indicator"), _text(lang, "summary_value")], summary_rows)
    doc.add_paragraph(build_kind_distribution_commentary(diagnostics, lang))
    if sample_count < 12:
        doc.add_paragraph(_text(lang, "small_sample_warning"))
    _add_picture_with_caption(doc, charts["kinds"], _text(lang, "distribution_chart_caption"))

    doc.add_heading(_text(lang, "features_heading"), level=1)
    doc.add_paragraph(build_feature_summary(diagnostics, meta.get("feature_labels") or {}, lang))
    feature_rows = [[_feature_label(str(item["name"]), meta.get("feature_labels") or {}, lang), f"{float(item['weight']):.3f}"] for item in diagnostics.get("top_features") or []]
    if feature_rows:
        _add_table(doc, [_localized(meta.get("feature_column_label") or {"en": "Variable", "es": "Variable", "tr": "Degisken"}, lang), _localized(meta.get("weight_column_label") or {"en": "Learned weight", "es": "Peso aprendido", "tr": "Ogrenilen agirlik"}, lang)], feature_rows)
    _add_picture_with_caption(doc, charts["features"], _text(lang, "features_chart_caption"))

    doc.add_heading(_text(lang, "config_heading"), level=1)
    doc.add_paragraph(_localized(meta["config_summary"], lang))
    integration_summary = _localized(meta.get("integration_summary") or "", lang).strip()
    if integration_summary:
        doc.add_paragraph(integration_summary)

    doc.add_heading(_text(lang, "validation_heading"), level=1)
    _add_table(
        doc,
        [_text(lang, "metric"), _text(lang, "result")],
        [
            ["Accuracy", _pct(metrics.get("accuracy", 0.0))],
            ["Precision", _pct(metrics.get("precision", 0.0))],
            ["Recall", _pct(metrics.get("recall", 0.0))],
            ["F1-score", _pct(metrics.get("f1", 0.0))],
            ["ROC-AUC", _pct(metrics.get("roc_auc", 0.0))],
            [_text(lang, "ml_weight"), f"{float(diagnostics.get('blend_weight', 0.0)):.2f}"],
            [_text(lang, "maturity"), maturity],
        ],
    )
    _add_picture_with_caption(doc, charts["metrics"], _text(lang, "metrics_chart_caption"))
    _add_picture_with_caption(doc, charts["confusion"], _text(lang, "confusion_chart_caption"))
    doc.add_paragraph(_text(lang, "metric_interpretation_intro"))
    for line in build_metric_interpretations(diagnostics, lang):
        doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph(build_confusion_interpretation(diagnostics, lang))
    objective_diagnostics = diagnostics.get("objective_diagnostics") or {}
    if int(objective_diagnostics.get("sample_count") or 0) > 0:
        objective_metrics = objective_diagnostics.get("metrics") or {}
        doc.add_paragraph(_text(lang, "objective_selector_heading"))
        doc.add_paragraph(
            _text(
                lang,
                "objective_selector_line",
                sample_count=int(objective_diagnostics.get("sample_count") or 0),
                f1=_pct(objective_metrics.get("f1", 0.0)),
                roc_auc=_pct(objective_metrics.get("roc_auc", 0.0)),
            )
        )
    segment_rows = diagnostics.get("segment_rows") or []
    if segment_rows:
        doc.add_paragraph(_text(lang, "segment_intro"))
        _add_table(
            doc,
            [_text(lang, "segment"), _text(lang, "samples"), _text(lang, "positive_rate"), _text(lang, "f1"), _text(lang, "recommended_weight")],
            [
                [str(item["segment"]), str(int(item["sample_count"])), _pct(item["positive_rate"]), _pct(item["f1"]), f"{float(item['recommended_blend_weight']):.2f}"]
                for item in segment_rows
            ],
        )
        if "segments" in charts:
            _add_picture_with_caption(doc, charts["segments"], _text(lang, "segments_chart_caption"))

    doc.add_heading(_text(lang, "analysis_heading"), level=1)
    analysis_lines = [
        _text(lang, "analysis_sample_line", sample_count=sample_count, f1=_pct(metrics.get("f1", 0.0))),
        _text(lang, "analysis_positive_line", positive_rate=_pct(positive_rate)),
        _text(lang, "analysis_maturity_line", maturity=maturity),
        maturity_text(level, lang),
    ]
    for extra in meta.get("analysis_extras", []):
        analysis_lines.append(_localized(extra, lang))
    for line in analysis_lines:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph(_text(lang, "limitations_heading"))
    for line in build_limitations(diagnostics, lang, [_localized(item, lang) for item in meta.get("extra_limitations", [])]):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph(build_final_conclusion(diagnostics, lang))
    doc.save(docx_path)
